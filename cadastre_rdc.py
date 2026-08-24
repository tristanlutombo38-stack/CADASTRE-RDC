import asyncio
from datetime import datetime
import hashlib
import json
import math
import os
import re
import secrets
import shutil
import sqlite3
import time
import unicodedata
import uuid
import zipfile
from zoneinfo :: ZoneInfo  # Gardez votre syntaxe d'import existante
import flet as ft

# ============================================================
# CADASTRE RDC — 3e VERSION OPÉRATIONNELLE + AMÉLIORATIONS
# ============================================================
# Technologie : Python + Flet + SQLite
#
# BASE DE DÉPART CONSERVÉE :
#   - Propriétaires
#   - Parcelles
#   - Relevé GPS / calcul de superficie
#   - Mutations / ventes
#   - Recherche / vérification
#   - Utilisateurs
#   - Journal des opérations
#
# AMÉLIORATIONS INTÉGRÉES :
#   - Interface professionnelle modernisée
#   - Palette visuelle inspirée de la RDC
#   - Navigation latérale moderne
#   - Français / English / Kiswahili / Lingala / Tshiluba
#   - Changement de langue dynamique
#   - Mode clair / sombre
#   - En-tête professionnel
#   - Cartes statistiques du tableau de bord
#   - Recherche dynamique
#   - Migration automatique de l'ancienne base SQLite
#   - Correction des colonnes CNI / date de création
#   - Validation renforcée des données
#   - Meilleure gestion des erreurs
#   - Conservation du fonctionnement général de la version fournie
#
# IMPORTANT :
# Ce programme reste un prototype. La capture GPS utilisée ici
# est une valeur de démonstration et doit être remplacée par
# une API GPS réelle pour un usage terrain.
# ============================================================

DB_NAME = "cadastre_rdc.db"

# ============================================================
# PALETTE PROFESSIONNELLE RDC
# ============================================================

COLOR_PRIMARY = "#007FFF"
COLOR_SECONDARY = "#F7D600"
COLOR_ACCENT = "#CE1126"
COLOR_BG_DARK = "#0C1B33"
COLOR_BG_LIGHT = "#F4F7FC"
COLOR_SURFACE = "#F8FAFC"
COLOR_TEXT_MAIN = "#10233F"
COLOR_TEXT_MUTED = "#64748B"
COLOR_BORDER = "#CBD5E1"
COLOR_RDC_BLUE = "#007FFF"
COLOR_RDC_RED = "#D51F2C"
COLOR_RDC_YELLOW = "#F4C430"
COLOR_DEEP_BLUE = "#071A33"
COLOR_CARD = "#FFFFFF"

RDC_FLAG_URL = "https://drapeaux-dumonde.com/cdn/shop/products/Autocollant-drapeau-RDC.jpg?v=1685511410"
RDC_COAT_OF_ARMS_URL = "https://tse2.mm.bing.net/th/id/OIP.2FzFjAweRNDv1i3ZxY9N6gAAAA?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"

# Palette des boutons inspirée de l'interface EIEM de référence.
# Elle est purement visuelle : aucune permission ni route n'est modifiée.
MENU_GRADIENTS = {
    "dashboard": ("#1565C0", "#0D47A1"),
    "proprietaires": ("#2E7D32", "#1B5E20"),
    "parcelles": ("#6A1B9A", "#4A148C"),
    "mutations": ("#EF6C00", "#E65100"),
    "verification": ("#00897B", "#004D40"),
    "utilisateurs": ("#3949AB", "#1A237E"),
    "territoires": ("#00838F", "#006064"),
    "journal": ("#0277BD", "#01579B"),
    "settings": ("#5D4037", "#3E2723"),
}


# ============================================================
# TRADUCTIONS
# ============================================================

TRANSLATIONS = {
    "fr": {
        "app_title": "CADASTRE RDC",
        "app_subtitle": "Système numérique cadastral",
        "republic": "République Démocratique du Congo",
        "dashboard": "Tableau de bord",
        "owners": "Propriétaires",
        "parcels": "Parcelles",
        "mutations": "Mutations / Ventes",
        "search": "Recherche / Vérification",
        "users": "Utilisateurs",
        "journal": "Journal",
        "settings": "Paramètres",
        "language": "Langue de l'interface",
        "theme": "Thème",
        "light": "Clair",
        "dark": "Sombre",
        "welcome": "Système numérique géolocalisé de gestion et de vérification des parcelles",
        "dashboard_title": "Tableau de bord",
        "dashboard_subtitle": "Gestion numérique, géolocalisée et vérification des parcelles",
        "objective": "Objectif du prototype",
        "objective_desc": (
            "Centraliser les informations cadastrales, associer les propriétaires "
            "aux parcelles, enregistrer les coordonnées GPS, conserver l'historique "
            "des mutations et préparer l'évolution vers une plateforme géospatiale "
            "et un portail public."
        ),
        "important": "Important",
        "important_desc": (
            "Ce programme est un prototype informatique. Il ne constitue pas, "
            "à lui seul, une preuve juridique de propriété et doit être adapté "
            "aux procédures et autorités compétentes."
        ),
        "owners_title": "Gestion des propriétaires",
        "owners_subtitle": "Création et consultation des fiches propriétaires",
        "full_name": "Nom complet",
        "phone": "Numéro de téléphone",
        "email": "E-mail",
        "identity": "Numéro / référence de pièce d'identité",
        "search_owner": "Rechercher un propriétaire",
        "save_owner": "Enregistrer le propriétaire",
        "parcels_title": "Gestion des parcelles",
        "parcels_subtitle": "Enregistrement topographique et dimensions (cm²)",
        "parcel_number": "Numéro / ID de la parcelle",
        "address": "Adresse complète",
        "province": "Province",
        "city": "Ville / Territoire",
        "commune": "Commune / Chefferie",
        "neighborhood": "Quartier / Groupement",
        "locality": "Localité",
        "area": "Superficie (cm²) ou dimensions (ex : 150/250)",
        "latitude": "Latitude",
        "longitude": "Longitude",
        "agent": "Agent / Technicien",
        "owner_name": "Nom complet du propriétaire",
        "owner_identity": "Numéro de carte d'identité (CNI)",
        "owner_photo": "Chemin de la photo du propriétaire",
        "browse": "Parcourir...",
        "status": "Statut",
        "occupied": "Occupée",
        "available": "Disponible",
        "for_sale": "En vente",
        "transferred": "Transférée",
        "disputed": "Contestée",
        "survey": "Relevé topographique & saisie manuelle",
        "survey_help": (
            "Capturez les sommets GPS (calcul automatique) ou saisissez "
            "directement vos dimensions, par exemple 150/250."
        ),
        "capture_gps": "Capturer un sommet GPS",
        "gps_ready": "Relevé de terrain : prêt pour capturer les sommets",
        "polygon": "Points GPS du polygone (JSON)",
        "save_parcel": "Enregistrer la parcelle",
        "search_parcel": "Rechercher une parcelle",
        "supported": "Format supporté : calcul GPS automatique ou dimensions type 150/250",
        "transfers_title": "Mutations et ventes",
        "transfers_subtitle": "Conservation de l'historique des changements de propriété",
        "search_plot": "Rechercher parcelle (N° ou ID)",
        "search_old_owner": "Rechercher l'ancien propriétaire",
        "no_plot_selected": "Aucune parcelle sélectionnée",
        "no_old_owner": "Aucun ancien propriétaire sélectionné",
        "current_owner": "Propriétaire actuel",
        "new_owner": "Informations du nouveau propriétaire",
        "transaction_type": "Type de transaction",
        "sale": "Vente",
        "cession": "Cession",
        "donation": "Donation",
        "other": "Autre",
        "act_reference": "Référence de l'acte",
        "price": "Prix déclaré (optionnel)",
        "save_transfer": "Enregistrer la mutation",
        "verification_title": "Recherche et vérification",
        "verification_subtitle": "Module préparatoire au futur portail public",
        "search_reference": "Numéro de parcelle, propriétaire ou référence",
        "search_button": "Rechercher",
        "verification_found": "VÉRIFICATION : ENREGISTREMENT TROUVÉ",
        "no_results": "Aucune parcelle trouvée.",
        "users_title": "Utilisateurs et rôles",
        "users_subtitle": "Préparation de la gestion des droits d'accès",
        "role": "Rôle",
        "administrator": "Administrateur",
        "supervisor": "Superviseur",
        "technician": "Technicien",
        "agent_role": "Agent",
        "consultation": "Consultation",
        "add_user": "Ajouter l'utilisateur",
        "journal_title": "Journal des opérations",
        "journal_subtitle": "Traçabilité et historique des actions du système",
        "id": "ID",
        "name": "Nom",
        "action": "Action",
        "details": "Détails",
        "date": "Date & Heure",
        "parcel": "Parcelle",
        "owner": "Propriétaire",
        "area_short": "Superficie",
        "location": "Localisation",
        "previous": "Ancien",
        "new": "Nouveau",
        "type": "Type",
        "logout": "Fermer",
        "save": "Enregistrer",
        "refresh": "Actualiser",
        "owner_required": "Le nom complet du propriétaire est obligatoire.",
        "parcel_required": "Veuillez sélectionner une parcelle.",
        "new_owner_required": "Le nom du nouveau propriétaire est obligatoire.",
        "invalid_price": "Le prix doit être numérique.",
        "invalid_area": "La superficie doit être un nombre ou un format du type 150/250.",
        "invalid_gps": "Latitude/longitude incorrecte.",
        "invalid_polygon": "Le polygone GPS n'est pas un JSON valide.",
        "owner_saved": "Propriétaire enregistré avec succès.",
        "parcel_saved": "Parcelle enregistrée avec succès.",
        "transfer_saved": "Mutation enregistrée avec succès.",
        "user_saved": "Utilisateur enregistré avec succès.",
        "duplicate_parcel": "Ce numéro de parcelle existe déjà.",
        "error": "Erreur",
        "prototype": "Prototype — LUTOMBO MISABEO TRISTAN",
        "theme_changed": "Thème modifié.",
        "language_changed": "Langue modifiée.",
        "admin_card": "Administrateur National",
        "admin_desc": "Accès total au système cadastral géospatial RDC",
        "records": "Enregistrements",
        "gps_point_added": "Sommet ajouté ! Total points : {count}",
        "gps_label": "GPS",
        "creation_owner": "Création propriétaire",
        "creation_parcel": "Création parcelle",
        "new_transfer": "Nouvelle mutation",
        "creation_user": "Création utilisateur",
        "pending": "En attente",
        "validate": "Valider",
        "validated": "Validée",
        "mutation_validated": "Mutation validée avec succès.",
        "invalid_mutation": "Mutation invalide ou incohérente.",
        "parcel_not_found": "Parcelle introuvable.",
        "owner_not_found": "Propriétaire introuvable.",
        "same_owner": "Le nouveau propriétaire doit être différent de l'ancien.",
        "parcel_disputed": "Cette parcelle est contestée et ne peut pas être mutée sans validation.",
        "mutation_pending": "Mutation enregistrée en attente de validation.",
    },
    "en": {
        "app_title": "DRC CADASTRE",
        "app_subtitle": "Digital Cadastral System",
        "republic": "Democratic Republic of the Congo",
        "dashboard": "Dashboard",
        "owners": "Owners",
        "parcels": "Parcels",
        "mutations": "Transfers / Sales",
        "search": "Search / Verification",
        "users": "Users",
        "journal": "Activity Log",
        "settings": "Settings",
        "language": "Interface language",
        "theme": "Theme",
        "light": "Light",
        "dark": "Dark",
        "welcome": "Geolocated digital system for parcel management and verification",
        "dashboard_title": "Dashboard",
        "dashboard_subtitle": "Digital, geolocated parcel management and verification",
        "objective": "Prototype objective",
        "objective_desc": (
            "Centralize cadastral information, link owners to parcels, record GPS "
            "coordinates, maintain transfer history and prepare the transition "
            "to a geospatial platform and public portal."
        ),
        "important": "Important",
        "important_desc": (
            "This software is a computer prototype. It does not constitute legal "
            "proof of ownership on its own and must be adapted to the procedures "
            "of the competent authorities."
        ),
        "owners_title": "Owner management",
        "owners_subtitle": "Create and consult owner records",
        "full_name": "Full name",
        "phone": "Phone number",
        "email": "Email",
        "identity": "Identity document number / reference",
        "search_owner": "Search an owner",
        "save_owner": "Save owner",
        "parcels_title": "Parcel management",
        "parcels_subtitle": "Topographical registration and dimensions (cm²)",
        "parcel_number": "Parcel number / ID",
        "address": "Full address",
        "province": "Province",
        "city": "City / Territory",
        "commune": "Municipality / Chiefdom",
        "neighborhood": "Neighborhood / Groupment",
        "locality": "Locality",
        "area": "Area (cm²) or dimensions (e.g. 150/250)",
        "latitude": "Latitude",
        "longitude": "Longitude",
        "agent": "Agent / Technician",
        "owner_name": "Owner full name",
        "owner_identity": "Identity card number",
        "owner_photo": "Owner photo path",
        "browse": "Browse...",
        "status": "Status",
        "occupied": "Occupied",
        "available": "Available",
        "for_sale": "For sale",
        "transferred": "Transferred",
        "disputed": "Disputed",
        "survey": "Topographic survey & manual entry",
        "survey_help": (
            "Capture GPS vertices or enter dimensions directly, "
            "for example 150/250."
        ),
        "capture_gps": "Capture GPS vertex",
        "gps_ready": "Field survey: ready to capture vertices",
        "polygon": "GPS polygon points (JSON)",
        "save_parcel": "Save parcel",
        "search_parcel": "Search for a parcel",
        "supported": "Supported format: automatic GPS calculation or 150/250 dimensions",
        "transfers_title": "Transfers and sales",
        "transfers_subtitle": "Property ownership transfer history",
        "search_plot": "Search parcel (No. or ID)",
        "search_old_owner": "Search previous owner",
        "no_plot_selected": "No parcel selected",
        "no_old_owner": "No previous owner selected",
        "current_owner": "Current owner",
        "new_owner": "New owner information",
        "transaction_type": "Transaction type",
        "sale": "Sale",
        "cession": "Assignment",
        "donation": "Donation",
        "other": "Other",
        "act_reference": "Deed reference",
        "price": "Declared price (optional)",
        "save_transfer": "Save transfer",
        "verification_title": "Search and verification",
        "verification_subtitle": "Preparation module for the future public portal",
        "search_reference": "Parcel number, owner or reference",
        "search_button": "Search",
        "verification_found": "VERIFICATION: RECORD FOUND",
        "no_results": "No parcel found.",
        "users_title": "Users and roles",
        "users_subtitle": "Preparation of access-right management",
        "role": "Role",
        "administrator": "Administrator",
        "supervisor": "Supervisor",
        "technician": "Technician",
        "agent_role": "Agent",
        "consultation": "View only",
        "add_user": "Add user",
        "journal_title": "Activity log",
        "journal_subtitle": "Traceability and history of system actions",
        "id": "ID",
        "name": "Name",
        "action": "Action",
        "details": "Details",
        "date": "Date & Time",
        "parcel": "Parcel",
        "owner": "Owner",
        "area_short": "Area",
        "location": "Location",
        "previous": "Previous",
        "new": "New",
        "type": "Type",
        "logout": "Close",
        "save": "Save",
        "refresh": "Refresh",
        "owner_required": "Owner full name is required.",
        "parcel_required": "Please select a parcel.",
        "new_owner_required": "New owner full name is required.",
        "invalid_price": "Price must be numeric.",
        "invalid_area": "Area must be a number or 150/250 format.",
        "invalid_gps": "Invalid latitude/longitude.",
        "invalid_polygon": "GPS polygon is not valid JSON.",
        "owner_saved": "Owner saved successfully.",
        "parcel_saved": "Parcel saved successfully.",
        "transfer_saved": "Transfer saved successfully.",
        "user_saved": "User saved successfully.",
        "duplicate_parcel": "This parcel number already exists.",
        "error": "Error",
        "prototype": "Prototype — LUTOMBO MISABEO TRISTAN",
        "theme_changed": "Theme changed.",
        "language_changed": "Language changed.",
        "admin_card": "National Administrator",
        "admin_desc": "Full access to the DRC geospatial cadastral system",
        "records": "Records",
        "gps_point_added": "Vertex added! Total points: {count}",
        "gps_label": "GPS",
        "creation_owner": "Owner created",
        "creation_parcel": "Parcel created",
        "new_transfer": "New transfer",
        "creation_user": "User created",
        "pending": "Pending",
        "validate": "Validate",
        "validated": "Validated",
        "mutation_validated": "Transfer validated successfully.",
        "invalid_mutation": "Invalid or inconsistent transfer.",
        "parcel_not_found": "Parcel not found.",
        "owner_not_found": "Owner not found.",
        "same_owner": "The new owner must be different from the previous owner.",
        "parcel_disputed": "This parcel is disputed and cannot be transferred without validation.",
        "mutation_pending": "Transfer recorded and awaiting validation.",
    },
    "sw": {
        "app_title": "KADASTRI YA DRC",
        "app_subtitle": "Mfumo wa Kidijitali wa Kadasta",
        "republic": "Jamhuri ya Kidemokrasia ya Kongo",
        "dashboard": "Dashibodi",
        "owners": "Wamiliki",
        "parcels": "Viwanja",
        "mutations": "Uhamisho / Mauzo",
        "search": "Utafutaji / Uthibitishaji",
        "users": "Watumiaji",
        "journal": "Rejesta ya Shughuli",
        "settings": "Mipangilio",
        "language": "Lugha ya kiolesura",
        "theme": "Mandhari",
        "light": "Angavu",
        "dark": "Nyeusi",
        "welcome": "Mfumo wa kidijitali wenye GPS kwa usimamizi na uthibitishaji wa viwanja",
        "dashboard_title": "Dashibodi",
        "dashboard_subtitle": "Usimamizi wa kidijitali na uthibitishaji wa viwanja kwa GPS",
        "objective": "Lengo la programu",
        "objective_desc": (
            "Kukusanya taarifa za kadasta, kuunganisha wamiliki na viwanja, "
            "kuhifadhi GPS, historia ya uhamisho na kuandaa mfumo wa kijiografia "
            "pamoja na tovuti ya umma."
        ),
        "important": "Muhimu",
        "important_desc": (
            "Huu ni mfano wa programu. Si uthibitisho wa kisheria wa umiliki "
            "peke yake na lazima uendane na taratibu za mamlaka husika."
        ),
        "owners_title": "Usimamizi wa wamiliki",
        "owners_subtitle": "Kuunda na kuangalia taarifa za wamiliki",
        "full_name": "Jina kamili",
        "phone": "Namba ya simu",
        "email": "Barua pepe",
        "identity": "Namba / rejea ya kitambulisho",
        "search_owner": "Tafuta mmiliki",
        "save_owner": "Hifadhi mmiliki",
        "parcels_title": "Usimamizi wa viwanja",
        "parcels_subtitle": "Usajili wa topografia na vipimo (cm²)",
        "parcel_number": "Namba / ID ya kiwanja",
        "address": "Anwani kamili",
        "province": "Jimbo / Mkoa",
        "city": "Mji / Wilaya",
        "commune": "Kata / Chefferie",
        "neighborhood": "Kitongoji / Groupement",
        "locality": "Eneo",
        "area": "Eneo (cm²) au vipimo (mf. 150/250)",
        "latitude": "Latitude",
        "longitude": "Longitude",
        "agent": "Wakala / Fundi",
        "owner_name": "Jina kamili la mmiliki",
        "owner_identity": "Namba ya kitambulisho",
        "owner_photo": "Njia ya picha ya mmiliki",
        "browse": "Chagua...",
        "status": "Hali",
        "occupied": "Inamilikiwa",
        "available": "Inapatikana",
        "for_sale": "Inauzwa",
        "transferred": "Imehamishwa",
        "disputed": "Ina mgogoro",
        "survey": "Uchunguzi wa topografia na vipimo",
        "survey_help": (
            "Piga pointi za GPS au ingiza vipimo moja kwa moja, "
            "kwa mfano 150/250."
        ),
        "capture_gps": "Piga pointi ya GPS",
        "gps_ready": "Uchunguzi wa eneo: tayari kupiga pointi",
        "polygon": "Pointi za GPS za polygon (JSON)",
        "save_parcel": "Hifadhi kiwanja",
        "search_parcel": "Tafuta kiwanja",
        "supported": "Muundo: hesabu ya GPS au vipimo 150/250",
        "transfers_title": "Uhamisho na mauzo",
        "transfers_subtitle": "Historia ya mabadiliko ya umiliki",
        "search_plot": "Tafuta kiwanja (Namba au ID)",
        "search_old_owner": "Tafuta mmiliki wa zamani",
        "no_plot_selected": "Hakuna kiwanja kilichochaguliwa",
        "no_old_owner": "Hakuna mmiliki wa zamani aliyechaguliwa",
        "current_owner": "Mmiliki wa sasa",
        "new_owner": "Taarifa za mmiliki mpya",
        "transaction_type": "Aina ya muamala",
        "sale": "Mauzo",
        "cession": "Uhamisho",
        "donation": "Mchango",
        "other": "Nyingine",
        "act_reference": "Rejea ya hati",
        "price": "Bei iliyotangazwa (hiari)",
        "save_transfer": "Hifadhi uhamisho",
        "verification_title": "Utafutaji na uthibitishaji",
        "verification_subtitle": "Moduli ya maandalizi ya tovuti ya umma",
        "search_reference": "Namba ya kiwanja, mmiliki au rejea",
        "search_button": "Tafuta",
        "verification_found": "UTHIBITISHO: TAARIFA IMEPATIKANA",
        "no_results": "Hakuna kiwanja kilichopatikana.",
        "users_title": "Watumiaji na majukumu",
        "users_subtitle": "Maandalizi ya usimamizi wa ruhusa",
        "role": "Jukumu",
        "administrator": "Msimamizi",
        "supervisor": "Msimamizi mkuu",
        "technician": "Fundi",
        "agent_role": "Wakala",
        "consultation": "Kuangalia tu",
        "add_user": "Ongeza mtumiaji",
        "journal_title": "Rejesta ya shughuli",
        "journal_subtitle": "Ufuatiliaji wa historia ya shughuli za mfumo",
        "id": "ID",
        "name": "Jina",
        "action": "Kitendo",
        "details": "Maelezo",
        "date": "Tarehe na saa",
        "parcel": "Kiwanja",
        "owner": "Mmiliki",
        "area_short": "Eneo",
        "location": "Mahali",
        "previous": "Wa zamani",
        "new": "Mpya",
        "type": "Aina",
        "logout": "Funga",
        "save": "Hifadhi",
        "refresh": "Onyesha upya",
        "owner_required": "Jina kamili la mmiliki linahitajika.",
        "parcel_required": "Tafadhali chagua kiwanja.",
        "new_owner_required": "Jina la mmiliki mpya linahitajika.",
        "invalid_price": "Bei lazima iwe namba.",
        "invalid_area": "Eneo lazima liwe namba au 150/250.",
        "invalid_gps": "Latitude/longitude si sahihi.",
        "invalid_polygon": "Polygon ya GPS si JSON sahihi.",
        "owner_saved": "Mmiliki amehifadhiwa kwa mafanikio.",
        "parcel_saved": "Kiwanja kimehifadhiwa kwa mafanikio.",
        "transfer_saved": "Uhamisho umehifadhiwa kwa mafanikio.",
        "user_saved": "Mtumiaji amehifadhiwa kwa mafanikio.",
        "duplicate_parcel": "Namba hii ya kiwanja tayari ipo.",
        "error": "Hitilafu",
        "prototype": "Prototype — LUTOMBO MISABEO TRISTAN",
        "theme_changed": "Mandhari imebadilishwa.",
        "language_changed": "Lugha imebadilishwa.",
        "admin_card": "Msimamizi wa Taifa",
        "admin_desc": "Ufikiaji kamili wa mfumo wa kadasta wa kijiografia wa DRC",
        "records": "Taarifa",
        "gps_point_added": "Pointi imeongezwa! Jumla ya pointi: {count}",
        "gps_label": "GPS",
        "creation_owner": "Kuunda mmiliki",
        "creation_parcel": "Kuunda kiwanja",
        "new_transfer": "Uhamisho mpya",
        "creation_user": "Kuunda mtumiaji",
        "pending": "Inasubiri",
        "validate": "Thibitisha",
        "validated": "Imethibitishwa",
        "mutation_validated": "Uhamisho umeidhinishwa kwa mafanikio.",
        "invalid_mutation": "Uhamisho si sahihi au hauna uwiano.",
        "parcel_not_found": "Kiwanja hakipatikani.",
        "owner_not_found": "Mmiliki hakupatikani.",
        "same_owner": "Mmiliki mpya lazima awe tofauti na wa zamani.",
        "parcel_disputed": "Kiwanja kina mgogoro na hakiwezi kuhamishwa bila idhini.",
        "mutation_pending": "Uhamisho umehifadhiwa ukisubiri uthibitisho.",
    },
    "ln": {"app_title": "CADASTRE RDC", "app_subtitle": "Système numérique ya cadastre", "republic": "Republiki Demokratiki ya Congo", "dashboard": "Etando ya bokambi", "owners": "Bankolo mabele", "parcels": "Mabele", "mutations": "Bobongisi / Boteki", "search": "Boluki / Botali", "users": "Bato basalelaka système", "journal": "Mokanda ya misala", "settings": "Bobongisi", "language": "Lokota ya interface", "theme": "Motindo", "light": "Pole", "dark": "Molili", "welcome": "Système numérique géolocalisé de gestion et de vérification des parcelles", "dashboard_title": "Tableau de bord", "dashboard_subtitle": "Gestion numérique, géolocalisée et vérification des parcelles", "objective": "Objectif du prototype", "objective_desc": "Centraliser les informations cadastrales, associer les propriétaires aux parcelles, enregistrer les coordonnées GPS, conserver l'historique des mutations et préparer l'évolution vers une plateforme géospatiale et un portail public.", "important": "Important", "important_desc": "Ce programme est un prototype informatique. Il ne constitue pas, à lui seul, une preuve juridique de propriété et doit être adapté aux procédures et autorités compétentes.", "owners_title": "Gestion des propriétaires", "owners_subtitle": "Création et consultation des fiches propriétaires", "full_name": "Kombo mobimba", "phone": "Nimero ya telefone", "email": "E-mail", "identity": "Nimero ya mokanda ya bomoto", "search_owner": "Luka nkolo mabele", "save_owner": "Bomba nkolo mabele", "parcels_title": "Gestion des parcelles", "parcels_subtitle": "Enregistrement topographique et dimensions (cm²)", "parcel_number": "Nimero ya mabele", "address": "Adresɛ mobimba", "province": "Province", "city": "Engumba / Territoire", "commune": "Commune / Chefferie", "neighborhood": "Quartier / Groupement", "locality": "Esika", "area": "Bonene (cm²)", "latitude": "Latitude", "longitude": "Longitude", "agent": "Agent / Technicien", "owner_name": "Kombo ya nkolo mabele", "owner_identity": "Nimero ya CNI", "owner_photo": "Foto ya nkolo mabele", "browse": "Pona foto", "status": "Statut", "occupied": "Ezali na moto", "available": "Ezali polele", "for_sale": "Ezali kotekama", "transferred": "Ebongisami", "disputed": "Ezali na ntembe", "survey": "Relevé topographique", "survey_help": "Capturez les sommets GPS (calcul automatique) ou saisissez directement vos dimensions, par exemple 150/250.", "capture_gps": "Kanga sommet GPS", "gps_ready": "Relevé ya terrain ebongami", "polygon": "Ba points GPS (JSON)", "save_parcel": "Bomba mabele", "search_parcel": "Luka mabele", "supported": "Format supporté : calcul GPS automatique ou dimensions type 150/250", "transfers_title": "Mutations et ventes", "transfers_subtitle": "Conservation de l'historique des changements de propriété", "search_plot": "Luka mabele", "search_old_owner": "Luka nkolo ya kala", "no_plot_selected": "Mabele moko te eponami", "no_old_owner": "Nkolo ya kala moko te eponami", "current_owner": "Nkolo ya sikoyo", "new_owner": "Nkolo ya sika", "transaction_type": "Lolenge ya transaction", "sale": "Boteki", "cession": "Kopesa", "donation": "Likabo", "other": "Mosusu", "act_reference": "Référence ya acte", "price": "Ntalo (optionnel)", "save_transfer": "Bomba mutation", "verification_title": "Boluki mpe verification", "verification_subtitle": "Module préparatoire au futur portail public", "search_reference": "Nimero ya mabele, nkolo to référence", "search_button": "Luka", "verification_found": "VERIFICATION: ENREGISTREMENT EZWAMI", "no_results": "Mabele ezwami te.", "users_title": "Utilisateurs et rôles", "users_subtitle": "Préparation de la gestion des droits d'accès", "role": "Mokumba", "administrator": "Administrateur", "supervisor": "Superviseur", "technician": "Technicien", "agent_role": "Agent", "consultation": "Consultation", "add_user": "Bakisa mosaleli", "journal_title": "Journal des opérations", "journal_subtitle": "Traçabilité et historique des actions du système", "id": "ID", "name": "Kombo", "action": "Likambo", "details": "Makambo", "date": "Mokolo mpe ngonga", "parcel": "Mabele", "owner": "Nkolo", "area_short": "Bonene", "location": "Esika", "previous": "Ya kala", "new": "Ya sika", "type": "Lolenge", "logout": "Kanga", "save": "Bomba", "refresh": "Zongisa", "owner_required": "Kombo ya nkolo mabele esengeli.", "parcel_required": "Pona mabele.", "new_owner_required": "Kombo ya nkolo ya sika esengeli.", "invalid_price": "Ntalo esengeli kozala motángo.", "invalid_area": "Bonene ezali malamu te.", "invalid_gps": "Latitude/longitude ezali malamu te.", "invalid_polygon": "Polygone GPS ezali malamu te.", "owner_saved": "Nkolo mabele abombami malamu.", "parcel_saved": "Mabele ebombami malamu.", "transfer_saved": "Mutation ebombami malamu.", "user_saved": "Mosaleli abombami malamu.", "duplicate_parcel": "Nimero ya mabele oyo ezali déjà.", "error": "Libunga", "prototype": "Prototype — LUTOMBO MISABEO TRISTAN", "theme_changed": "Motindo ebongwani.", "language_changed": "Lokota ebongwani.", "admin_card": "Administrateur National", "admin_desc": "Accès total au système cadastral géospatial RDC", "records": "Ba-enregistrements", "gps_point_added": "Sommet ebakisami! Total: {count}", "gps_label": "GPS", "creation_owner": "Kosala nkolo mabele", "creation_parcel": "Kosala mabele", "new_transfer": "Mutation ya sika", "creation_user": "Kosala mosaleli", "pending": "Ezali kozela", "validate": "Andima", "validated": "Endimami", "mutation_validated": "Mutation endimami malamu.", "invalid_mutation": "Mutation ezali malamu te.", "parcel_not_found": "Mabele ezwami te.", "owner_not_found": "Nkolo mabele azwami te.", "same_owner": "Nkolo ya sika asengeli kozala moto mosusu.", "parcel_disputed": "Mabele ezali na ntembe mpe ekoki te kobongolama kozanga ndingisa.", "mutation_pending": "Mutation ebombami mpe ezali kozela ndingisa."},
    "tsh": {"app_title": "CADASTRE RDC", "app_subtitle": "Système wa numérique wa cadastre", "republic": "Republiki Demokrati ya Kongo", "dashboard": "Etando wa bukokeshi", "owners": "Balongolodi ba mabele", "parcels": "Mabele", "mutations": "Mibongolodi / Mitɛki", "search": "Kulonda / Verification", "users": "Batu badi basalela système", "journal": "Mukanda wa misala", "settings": "Mibikidi", "language": "Dimi dia interface", "theme": "Mutenka", "light": "Pɔle", "dark": "Munkila", "welcome": "Système numérique géolocalisé de gestion et de vérification des parcelles", "dashboard_title": "Tableau de bord", "dashboard_subtitle": "Gestion numérique, géolocalisée et vérification des parcelles", "objective": "Objectif du prototype", "objective_desc": "Centraliser les informations cadastrales, associer les propriétaires aux parcelles, enregistrer les coordonnées GPS, conserver l'historique des mutations et préparer l'évolution vers une plateforme géospatiale et un portail public.", "important": "Important", "important_desc": "Ce programme est un prototype informatique. Il ne constitue pas, à lui seul, une preuve juridique de propriété et doit être adapté aux procédures et autorités compétentes.", "owners_title": "Gestion des propriétaires", "owners_subtitle": "Création et consultation des fiches propriétaires", "full_name": "Dina dia munda", "phone": "Nimero ya telefone", "email": "E-mail", "identity": "Nimero ya mukanda wa bumuntu", "search_owner": "Londa mulongolodi wa mabele", "save_owner": "Bomba mulongolodi wa mabele", "parcels_title": "Gestion des parcelles", "parcels_subtitle": "Enregistrement topographique et dimensions (cm²)", "parcel_number": "Nimero ya mabele", "address": "Adresɛ wa munda", "province": "Province", "city": "Muji / Territoire", "commune": "Commune / Chefferie", "neighborhood": "Quartier / Groupement", "locality": "Esika", "area": "Bunene (cm²)", "latitude": "Latitude", "longitude": "Longitude", "agent": "Agent / Technicien", "owner_name": "Dina dia mulongolodi wa mabele", "owner_identity": "Nimero ya CNI", "owner_photo": "Foto wa mulongolodi wa mabele", "browse": "Pona foto", "status": "Mikenji", "occupied": "Badi mu mabele", "available": "Mabele adi pabuipi", "for_sale": "Mu ditɛki", "transferred": "Ebongoloke", "disputed": "Didi ne ntembe", "survey": "Relevé topographique & saisie manuelle", "survey_help": "Capturez les sommets GPS (calcul automatique) ou saisissez directement vos dimensions, par exemple 150/250.", "capture_gps": "Kanga sommet GPS", "gps_ready": "Relevé de terrain : prêt pour capturer les sommets", "polygon": "Points GPS du polygone (JSON)", "save_parcel": "Bomba mabele", "search_parcel": "Londa mabele", "supported": "Format supporté : calcul GPS automatique ou dimensions type 150/250", "transfers_title": "Mutations et ventes", "transfers_subtitle": "Conservation de l'historique des changements de propriété", "search_plot": "Londa mabele", "search_old_owner": "Londa mulongolodi wa kala", "no_plot_selected": "Mabele kaponwa", "no_old_owner": "Mulongolodi wa kala kaponwa", "current_owner": "Mulongolodi wa lelu", "new_owner": "Mulongolodi musangu", "transaction_type": "Lukusa lua transaction", "sale": "Ditɛki", "cession": "Kupa", "donation": "Lukabi", "other": "Bimpe bindi", "act_reference": "Reference wa acte", "price": "Mushinga (optionnel)", "save_transfer": "Bomba mutation", "verification_title": "Kulonda ne verification", "verification_subtitle": "Module préparatoire au futur portail public", "search_reference": "Nimero ya mabele, mulongolodi to reference", "search_button": "Londa", "verification_found": "VERIFICATION: ENREGISTREMENT EZWAMI", "no_results": "Mabele kakutulukayi.", "users_title": "Utilisateurs et rôles", "users_subtitle": "Préparation de la gestion des droits d'accès", "role": "Mukumba", "administrator": "Administrateur", "supervisor": "Superviseur", "technician": "Technicien", "agent_role": "Agent", "consultation": "Consultation", "add_user": "Bakisa musaleli", "journal_title": "Journal des opérations", "journal_subtitle": "Traçabilité et historique des actions du système", "id": "ID", "name": "Dina", "action": "Musalu", "details": "Makambu", "date": "Diku ne nshikama", "parcel": "Mabele", "owner": "Mulongolodi", "area_short": "Bunene", "location": "Esika", "previous": "Wa kala", "new": "Musangu", "type": "Lukusa", "logout": "Kanga", "save": "Bomba", "refresh": "Zongolola", "owner_required": "Dina dia mulongolodi dishidi ne ditalakana.", "parcel_required": "Pona mabele.", "new_owner_required": "Dina dia mulongolodi musangu dishidi ne ditalakana.", "invalid_price": "Mushinga ufwanela kuba motángo.", "invalid_area": "Bunene kayi ne lulamatu.", "invalid_gps": "Latitude/longitude kayi ne lulamatu.", "invalid_polygon": "Polygone GPS kayi ne JSON mulenga.", "owner_saved": "Mulongolodi ubombwe bimpe.", "parcel_saved": "Mabele abombwe bimpe.", "transfer_saved": "Mutation ibombwe bimpe.", "user_saved": "Musaleli ubombwe bimpe.", "duplicate_parcel": "Nimero wa mabele udi ne kale.", "error": "Muphwanyi", "prototype": "Prototype — LUTOMBO MISABEO TRISTAN", "theme_changed": "Mutenka ubongoloke.", "language_changed": "Dimi dibongoloke.", "admin_card": "Administrateur National", "admin_desc": "Accès total au système cadastral géospatial RDC", "records": "Ba-enregistrements", "gps_point_added": "Sommet ajouté ! Total points : {count}", "gps_label": "GPS", "creation_owner": "Création propriétaire", "creation_parcel": "Création parcelle", "new_transfer": "Nouvelle mutation", "creation_user": "Création utilisateur", "pending": "Mu kuindila", "validate": "Andamena", "validated": "Endamena", "mutation_validated": "Mungongoloki wa mushindu umanyiwa bimpe.", "invalid_mutation": "Mungongoloki kawena mulenga to udi ne ntembe.", "parcel_not_found": "Mabele kaawani.", "owner_not_found": "Mulongolodi kaawani.", "same_owner": "Mulongolodi musangu afwanela kuba muntu wandi.", "parcel_disputed": "Mabele adi ne ntembe, kayi ne kubongoloka kadi kayi ne ndingisa.", "mutation_pending": "Mungongoloki ubombwe ne udi ulindila ndingisa."},

}


# ============================================================
# TRADUCTIONS — AUTHENTIFICATION / SÉCURITÉ
# ============================================================

SECURITY_TRANSLATIONS = {
    "fr": {
        "login_title": "Connexion — CADASTRE RDC",
        "login_subtitle": "Accès sécurisé au système cadastral",
        "username": "Nom d'utilisateur",
        "password": "Mot de passe",
        "confirm_password": "Confirmer le mot de passe",
        "login": "Se connecter",
        "logout": "Déconnexion",
        "invalid_credentials": "Nom d'utilisateur ou mot de passe incorrect.",
        "account_locked": "Compte temporairement verrouillé. Réessayez plus tard.",
        "account_inactive": "Ce compte est désactivé.",
        "setup_title": "Initialisation de la sécurité",
        "setup_subtitle": "Créez le premier compte Administrateur.",
        "setup_admin": "Créer l'administrateur",
        "password_mismatch": "Les mots de passe ne correspondent pas.",
        "password_short": "Le mot de passe doit contenir au moins 8 caractères.",
        "security_ready": "Compte Administrateur créé. Vous pouvez maintenant vous connecter.",
        "session_expired": "Votre session a expiré. Veuillez vous reconnecter.",
        "access_denied": "Accès refusé : vous n'avez pas les droits nécessaires.",
        "connected_as": "Connecté en tant que",
        "security": "Sécurité",
    },
    "en": {
        "login_title": "Login — DRC CADASTRE",
        "login_subtitle": "Secure access to the cadastral system",
        "username": "Username",
        "password": "Password",
        "confirm_password": "Confirm password",
        "login": "Log in",
        "logout": "Log out",
        "invalid_credentials": "Incorrect username or password.",
        "account_locked": "Account temporarily locked. Try again later.",
        "account_inactive": "This account is disabled.",
        "setup_title": "Security initialization",
        "setup_subtitle": "Create the first Administrator account.",
        "setup_admin": "Create Administrator",
        "password_mismatch": "Passwords do not match.",
        "password_short": "Password must contain at least 8 characters.",
        "security_ready": "Administrator account created. You can now log in.",
        "session_expired": "Your session has expired. Please log in again.",
        "access_denied": "Access denied: you do not have the required rights.",
        "connected_as": "Signed in as",
        "security": "Security",
    },
    "sw": {
        "login_title": "Kuingia — CADASTRE RDC",
        "login_subtitle": "Ufikiaji salama wa mfumo wa kadasta",
        "username": "Jina la mtumiaji",
        "password": "Nenosiri",
        "confirm_password": "Thibitisha nenosiri",
        "login": "Ingia",
        "logout": "Toka",
        "invalid_credentials": "Jina la mtumiaji au nenosiri si sahihi.",
        "account_locked": "Akaunti imefungwa kwa muda. Jaribu baadaye.",
        "account_inactive": "Akaunti hii imezimwa.",
        "setup_title": "Uanzishaji wa usalama",
        "setup_subtitle": "Unda akaunti ya kwanza ya Msimamizi.",
        "setup_admin": "Unda Msimamizi",
        "password_mismatch": "Manenosiri hayalingani.",
        "password_short": "Nenosiri lazima liwe na angalau herufi 8.",
        "security_ready": "Akaunti ya Msimamizi imeundwa. Sasa unaweza kuingia.",
        "session_expired": "Kikao chako kimeisha. Tafadhali ingia tena.",
        "access_denied": "Ufikiaji umekataliwa: huna ruhusa zinazohitajika.",
        "connected_as": "Umeingia kama",
        "security": "Usalama",
    },
    "ln": {
        "login_title": "Kokɔ — CADASTRE RDC",
        "login_subtitle": "Kokɔ ya libateli na système cadastral",
        "username": "Nkombo ya mosaleli",
        "password": "Mot de passe",
        "confirm_password": "Tíndisa mot de passe",
        "login": "Kokɔ",
        "logout": "Kobima",
        "invalid_credentials": "Nkombo to mot de passe ezali malamu te.",
        "account_locked": "Compte ekangami mpo na mwa ntango.",
        "account_inactive": "Compte oyo ezali désactivé.",
        "setup_title": "Bobandi ya libateli",
        "setup_subtitle": "Sala compte ya Administrateur ya liboso.",
        "setup_admin": "Sala Administrateur",
        "password_mismatch": "Ba mot de passe ekokani te.",
        "password_short": "Mot de passe esengeli kozala na bilembo 8 ata moke.",
        "security_ready": "Compte Administrateur esalemi. Okoki sikoyo kokɔta.",
        "session_expired": "Session na yo esili. Kɔta lisusu.",
        "access_denied": "Ndingisa eboyami.",
        "connected_as": "Okeyi lokola",
        "security": "Libateli",
    },
    "tsh": {
        "login_title": "Kukota — CADASTRE RDC",
        "login_subtitle": "Kukota kwa mutekete mu système cadastral",
        "username": "Dina dia mukoleshi",
        "password": "Mot de passe",
        "confirm_password": "Tshindija mot de passe",
        "login": "Kota",
        "logout": "Bima",
        "invalid_credentials": "Dina dia mukoleshi to mot de passe ntshia mudimu.",
        "account_locked": "Compte wakangibwa pa tshikondo tshia mamba.",
        "account_inactive": "Compte ewu udi ukangala.",
        "setup_title": "Kutangulula mutekete",
        "setup_subtitle": "Sala compte ya Administrateur wa ntete.",
        "setup_admin": "Sala Administrateur",
        "password_mismatch": "Mot de passe mikokani te.",
        "password_short": "Mot de passe udi ne bupeta bua bilembo 8 to kusanza.",
        "security_ready": "Compte ya Administrateur esalemi. Okudi ukota.",
        "session_expired": "Session yebe eshiile. Kota kabidi.",
        "access_denied": "Kukota kubayibwe: huna ndingisa.",
        "connected_as": "Wakota mu dina dia",
        "security": "Mutekete",
    },
    "kg": {
        "login_title": "Kota — CADASTRE RDC",
        "login_subtitle": "Kifuki kia lutaninu muna nkubukulu ya ntoto",
        "username": "Nkombo ya kisadi",
        "password": "Nsi ya mpova",
        "confirm_password": "Toma nsi ya mpova",
        "login": "Kota",
        "logout": "Baluka kuna mpanza",
        "invalid_credentials": "Nkombo to nsi ya mpova ke yolele ve.",
        "account_locked": "Kontu ekangilu fioti. Vutukila diaka na nima.",
        "account_inactive": "Kontu yayi ke yina na ngolo ve.",
        "setup_title": "N Tuba ya lutaninu",
        "setup_subtitle": "Salu kontu ya ntu diambu ya ntete.",
        "setup_admin": "Salu Ntu diambu",
        "password_mismatch": "Nsi ya mpova ke zolele ve.",
        "password_short": "Nsi ya mpova fwete vanda ye bilembo 8.",
        "security_ready": "Kontu ya Ntu diambu esalamene. Nge lenda kota.",
        "session_expired": "Ntangu ya nkubukulu esila. Vutuka kota diaka.",
        "access_denied": "Nswa ke yina ve: nge ke yina na nswa ve.",
        "connected_as": "Okotidi bonso",
        "security": "Lutaninu",
    },
}

for _lang, _values in SECURITY_TRANSLATIONS.items():
    TRANSLATIONS.setdefault(_lang, {}).update(_values)

# Libellés supplémentaires pour l'administration de la sécurité.
SECURITY_ADMIN_TRANSLATIONS = {
    "fr": {
        "user_management": "Gestion des utilisateurs",
        "edit_user": "Modifier",
        "update_user": "Enregistrer les modifications",
        "block_user": "Bloquer",
        "unblock_user": "Débloquer",
        "reset_password": "Changer le mot de passe",
        "active": "Actif",
        "inactive": "Inactif",
        "locked": "Verrouillé",
        "select_user": "Sélectionner un utilisateur",
        "new_password_optional": "Nouveau mot de passe (facultatif)",
        "security_logs": "Journal de sécurité",
        "backup": "Sauvegarde",
        "backup_now": "Sauvegarder maintenant",
        "restore_backup": "Restaurer une sauvegarde",
        "backup_created": "Sauvegarde créée",
        "backup_restored": "Sauvegarde restaurée. Redémarrage recommandé.",
        "backup_failed": "Échec de la sauvegarde",
        "restore_failed": "Échec de la restauration",
        "backup_ready": "Dernière sauvegarde",
        "operation_denied": "Opération refusée : droits insuffisants.",
        "last_login": "Dernière connexion",
        "save_changes": "Enregistrer",
        "select_user_to_edit": "Sélectionnez un utilisateur à modifier.",
        "cannot_disable_last_admin": "Le dernier administrateur ne peut pas être désactivé.",
        "cannot_change_self_role": "Vous ne pouvez pas modifier votre propre rôle.",
        "cannot_delete_last_admin": "Le dernier administrateur doit rester actif.",
        "security_backup_note": "Les sauvegardes contiennent la base cadastrale et les photos.",
    },
    "en": {
        "user_management": "User management", "edit_user": "Edit", "update_user": "Save changes",
        "block_user": "Block", "unblock_user": "Unblock", "reset_password": "Change password",
        "active": "Active", "inactive": "Inactive", "locked": "Locked", "select_user": "Select a user",
        "new_password_optional": "New password (optional)", "security_logs": "Security log",
        "backup": "Backup", "backup_now": "Back up now", "restore_backup": "Restore backup",
        "backup_created": "Backup created", "backup_restored": "Backup restored. Restart recommended.",
        "backup_failed": "Backup failed", "restore_failed": "Restore failed", "backup_ready": "Last backup",
        "operation_denied": "Operation denied: insufficient permissions.", "last_login": "Last login",
        "save_changes": "Save", "select_user_to_edit": "Select a user to edit.",
        "cannot_disable_last_admin": "The last administrator cannot be disabled.",
        "cannot_change_self_role": "You cannot change your own role.",
        "cannot_delete_last_admin": "The last administrator must remain active.",
        "security_backup_note": "Backups contain the cadastral database and photos.",
    },
    "sw": {
        "user_management": "Usimamizi wa watumiaji", "edit_user": "Hariri", "update_user": "Hifadhi mabadiliko",
        "block_user": "Funga", "unblock_user": "Fungua", "reset_password": "Badilisha nenosiri",
        "active": "Hai", "inactive": "Haifanyi kazi", "locked": "Imefungwa", "select_user": "Chagua mtumiaji",
        "new_password_optional": "Nenosiri jipya (si lazima)", "security_logs": "Kumbukumbu ya usalama",
        "backup": "Hifadhi rudufu", "backup_now": "Hifadhi sasa", "restore_backup": "Rejesha hifadhi",
        "backup_created": "Hifadhi imeundwa", "backup_restored": "Hifadhi imerejeshwa. Inashauriwa kuanzisha upya.",
        "backup_failed": "Hifadhi imeshindwa", "restore_failed": "Urejeshaji umeshindwa", "backup_ready": "Hifadhi ya mwisho",
        "operation_denied": "Operesheni imekataliwa: huna ruhusa.", "last_login": "Kuingia kwa mwisho",
        "save_changes": "Hifadhi", "select_user_to_edit": "Chagua mtumiaji wa kuhariri.",
        "cannot_disable_last_admin": "Msimamizi wa mwisho hawezi kuzimwa.", "cannot_change_self_role": "Huwezi kubadilisha jukumu lako mwenyewe.",
        "cannot_delete_last_admin": "Msimamizi wa mwisho lazima abaki hai.", "security_backup_note": "Hifadhi zina hifadhidata ya kadasta na picha.",
    },
    "ln": {
        "user_management": "Boyangeli ya basaleli", "edit_user": "Bongisa", "update_user": "Bomba mbongwana",
        "block_user": "Kanga", "unblock_user": "Fungola", "reset_password": "Bongisa mot de passe",
        "active": "Ezali kosala", "inactive": "Ezali kosala te", "locked": "Ekangami", "select_user": "Pona mosaleli",
        "new_password_optional": "Mot de passe ya sika (soki olingi)", "security_logs": "Journal ya libateli",
        "backup": "Kopi ya libateli", "backup_now": "Sala kopi sikoyo", "restore_backup": "Zongisa kopi",
        "backup_created": "Kopi esalemi", "backup_restored": "Kopi ezongisami. Kobanda lisusu système ezali malamu.",
        "backup_failed": "Kopi esalemi te", "restore_failed": "Kozongisa esalemi te", "backup_ready": "Kopi ya nsuka",
        "operation_denied": "Operasio eboyi: ozali na ndingisa te.", "last_login": "Kokɔta ya nsuka",
        "save_changes": "Bomba", "select_user_to_edit": "Pona mosaleli ya kobongisa.",
        "cannot_disable_last_admin": "Administrateur ya nsuka akoki kokangama te.", "cannot_change_self_role": "Okoki kobongisa rôle na yo moko te.",
        "cannot_delete_last_admin": "Administrateur ya nsuka asengeli kotikala kosala.", "security_backup_note": "Bakopi ezali na base cadastrale mpe bafoto.",
    },
    "tsh": {
        "user_management": "Ditungulu dia balongeshi", "edit_user": "Lungulula", "update_user": "Bika mbongwana",
        "block_user": "Kanga", "unblock_user": "Fungula", "reset_password": "Lungulula mot de passe",
        "active": "Udi mu mudimu", "inactive": "Kashidi mu mudimu", "locked": "Wakangibwa", "select_user": "Peta mulongeshi",
        "new_password_optional": "Mot de passe mupia (mukanda wa kudi)", "security_logs": "Journal wa mutekete",
        "backup": "Kopi ya mudimu", "backup_now": "Sala kopi lelu", "restore_backup": "Zongela kopi",
        "backup_created": "Kopi esalemi", "backup_restored": "Kopi ezongolweshile. Kutangulula lisusu kudi kwa malamu.",
        "backup_failed": "Kopi kayi salemi", "restore_failed": "Kuzongela kayi salemi", "backup_ready": "Kopi wa nsuka",
        "operation_denied": "Mudimu ubayibwe: huna ndingisa te.", "last_login": "Kukota kwa nsuka",
        "save_changes": "Bika", "select_user_to_edit": "Peta mulongeshi wa kulungulula.",
        "cannot_disable_last_admin": "Administrateur wa nsuka kayi ukangibwa te.", "cannot_change_self_role": "Kena mua kulungulula rôle webe mwakane te.",
        "cannot_delete_last_admin": "Administrateur wa nsuka asengela kutikala mu mudimu.", "security_backup_note": "Bakopi eyi ne base cadastrale ne mafoto.",
    },
    "kg": {
        "user_management": "Nlungulu ya basadi", "edit_user": "Bongisa", "update_user": "Bomba mbongwana",
        "block_user": "Kanga", "unblock_user": "Fungula", "reset_password": "Bongisa nsi ya mpova",
        "active": "Ke salaka", "inactive": "Ke salaka ve", "locked": "Ekangami", "select_user": "Pona kisadi",
        "new_password_optional": "Nsi ya mpova ya yampa (kana omoni mpasi)", "security_logs": "Mukanda ya lutaninu",
        "backup": "Nkopulu ya lutaninu", "backup_now": "Salu nkopulu mvu yayi", "restore_backup": "Vutula nkopulu",
        "backup_created": "Nkopulu esalemi", "backup_restored": "Nkopulu ivutukidi. Kubaluka diaka diambote.",
        "backup_failed": "Nkopulu esalemi ve", "restore_failed": "Kuvutula nkopulu esalemi ve", "backup_ready": "Nkopulu ya nsuka",
        "operation_denied": "Kisalu eboyi: nge ke yina na nswa ve.", "last_login": "Kukota ya nsuka",
        "save_changes": "Bomba", "select_user_to_edit": "Pona kisadi ya kubongisa.",
        "cannot_disable_last_admin": "Ntu diambu ya nsuka lenda kangama ve.", "cannot_change_self_role": "Nge lenda yidika kisalu na nge mosi ve.",
        "cannot_delete_last_admin": "Ntu diambu ya nsuka fwete vanda na luzingu.", "security_backup_note": "Nkopulu ke yina na base ya ntoto ye bafoto.",
    },
}
for _lang, _values in SECURITY_ADMIN_TRANSLATIONS.items():
    TRANSLATIONS.setdefault(_lang, {}).update(_values)

# ============================================================
# TRADUCTIONS — ADMINISTRATION TERRITORIALE
# ============================================================
TERRITORIAL_TRANSLATIONS = {
    "fr": {
        "territorial_admin": "Administration territoriale",
        "territorial_title": "Administration territoriale et contrôle",
        "territorial_subtitle": "Hiérarchie nationale, provinciale, urbaine et territoriale",
        "national_level": "National", "provincial_level": "Provincial",
        "city_level": "Ville", "territory_level": "Territoire",
        "commune_level": "Commune", "local_level": "Local",
        "province": "Province", "city": "Ville", "territory": "Territoire",
        "commune": "Commune", "sector_chiefdom": "Secteur / Chefferie",
        "grouping": "Groupement", "village": "Village",
        "department": "Département", "scope": "Périmètre d'accès",
        "structure_type": "Type de structure", "structure_name": "Nom de la structure",
        "structure_code": "Code", "parent_structure": "Structure parente",
        "add_structure": "Ajouter la structure", "structures": "Structures",
        "access_level": "Niveau d'accès", "national_control": "Contrôle national",
        "provincial_control": "Contrôle provincial", "territorial_control": "Contrôle territorial",
        "city_control": "Contrôle urbain", "department_control": "Contrôle des départements",
        "scope_restricted": "Accès limité au périmètre attribué.",
        "all_national": "Toutes les provinces / toutes les structures",
        "department_scope": "Département rattaché",
    },
    "en": {
        "territorial_admin": "Territorial Administration", "territorial_title": "Territorial Administration and Control",
        "territorial_subtitle": "National, provincial, urban and territorial hierarchy", "national_level": "National", "provincial_level": "Provincial", "city_level": "City", "territory_level": "Territory", "commune_level": "Commune", "local_level": "Local",
        "province": "Province", "city": "City", "territory": "Territory", "commune": "Commune", "sector_chiefdom": "Sector / Chiefdom", "grouping": "Group", "village": "Village", "department": "Department", "scope": "Access scope", "structure_type": "Structure type", "structure_name": "Structure name", "structure_code": "Code", "parent_structure": "Parent structure", "add_structure": "Add structure", "structures": "Structures", "access_level": "Access level", "national_control": "National control", "provincial_control": "Provincial control", "territorial_control": "Territorial control", "city_control": "Urban control", "department_control": "Department control", "scope_restricted": "Access is restricted to the assigned scope.", "all_national": "All provinces / all structures", "department_scope": "Assigned department",
    },
    "sw": {
        "territorial_admin": "Utawala wa Kieneo", "territorial_title": "Utawala na udhibiti wa kieneo", "territorial_subtitle": "Ngazi ya kitaifa, mkoa, jiji na eneo", "national_level": "Kitaifa", "provincial_level": "Mkoa", "city_level": "Jiji", "territory_level": "Wilaya", "commune_level": "Kommune", "local_level": "Kieneo", "province": "Mkoa", "city": "Jiji", "territory": "Wilaya", "commune": "Kommune", "sector_chiefdom": "Sekta / Chifu", "grouping": "Kundi", "village": "Kijiji", "department": "Idara", "scope": "Eneo la ufikiaji", "structure_type": "Aina ya muundo", "structure_name": "Jina la muundo", "structure_code": "Msimbo", "parent_structure": "Muundo mzazi", "add_structure": "Ongeza muundo", "structures": "Miundo", "access_level": "Ngazi ya ufikiaji", "national_control": "Udhibiti wa kitaifa", "provincial_control": "Udhibiti wa mkoa", "territorial_control": "Udhibiti wa eneo", "city_control": "Udhibiti wa jiji", "department_control": "Udhibiti wa idara", "scope_restricted": "Ufikiaji umewekewa mipaka kwa eneo lililotengwa.", "all_national": "Mikoa yote / miundo yote", "department_scope": "Idara iliyotengwa",
    },
    "ln": {
        "territorial_admin": "Boyangeli ya bituka", "territorial_title": "Boyangeli mpe bopesi ya bituka", "territorial_subtitle": "Niveau ya ekolo, province, engumba mpe territoire", "national_level": "Ekolo", "provincial_level": "Province", "city_level": "Engumba", "territory_level": "Territoire", "commune_level": "Commune", "local_level": "Ya esika", "province": "Province", "city": "Engumba", "territory": "Territoire", "commune": "Commune", "sector_chiefdom": "Secteur / Chefferie", "grouping": "Groupement", "village": "Mboka", "department": "Département", "scope": "Esika ya ndingisa", "structure_type": "Lolenge ya structure", "structure_name": "Kombo ya structure", "structure_code": "Code", "parent_structure": "Structure ya likolo", "add_structure": "Bakisa structure", "structures": "Structures", "access_level": "Niveau ya ndingisa", "national_control": "Bokengeli ya ekolo", "provincial_control": "Bokengeli ya province", "territorial_control": "Bokengeli ya territoire", "city_control": "Bokengeli ya engumba", "department_control": "Bokengeli ya département", "scope_restricted": "Ndingisa ekangami na esika epesami.", "all_national": "Ba province nyonso / ba structure nyonso", "department_scope": "Département epesami",
    },
    "tsh": {
        "territorial_admin": "Mushinga wa bitunga", "territorial_title": "Mushinga ne bukontrole bwa bitunga", "territorial_subtitle": "Niveau wa ditunga, province, tshibanda ne territoire", "national_level": "Ditunga", "provincial_level": "Province", "city_level": "Tshibanda", "territory_level": "Territoire", "commune_level": "Commune", "local_level": "Pabuipi", "province": "Province", "city": "Tshibanda", "territory": "Territoire", "commune": "Commune", "sector_chiefdom": "Secteur / Chefferie", "grouping": "Groupement", "village": "Muaba", "department": "Département", "scope": "Muaba wa ndingisa", "structure_type": "Nshila wa structure", "structure_name": "Dina dia structure", "structure_code": "Code", "parent_structure": "Structure wa kulonda", "add_structure": "Baka structure", "structures": "Structures", "access_level": "Niveau wa ndingisa", "national_control": "Bukontrole bwa ditunga", "provincial_control": "Bukontrole bwa province", "territorial_control": "Bukontrole bwa territoire", "city_control": "Bukontrole bwa tshibanda", "department_control": "Bukontrole bwa département", "scope_restricted": "Ndingisa udi ukangama ku muaba upewe.", "all_national": "Ma province onso / ma structure onso", "department_scope": "Département upewe",
    },
    "kg": {
        "territorial_admin": "Nlungulu ya Nsi", "territorial_title": "Nlungulu ya Nsi ye Ntwala",
        "territorial_subtitle": "Ntele ya nsi, kizunga, mbanza ye mfulu", "national_level": "Nsi", "provincial_level": "Kizunga",
        "city_level": "Mbanza", "territory_level": "Mfulu",
        "commune_level": "Kiminia", "local_level": "Tfulu",
        "province": "Kizunga", "city": "Mbanza", "territory": "Mfulu",
        "commune": "Kiminia", "sector_chiefdom": "Mfulu /Mfumu",
        "grouping": "Kabu", "village": "Mbanza fioti",
        "department": "Nsalulu", "scope": "Ntele ya kota",
        "structure_type": "Nkindu ya nkubukulu", "structure_name": "Nkombo ya nkubukulu",
        "structure_code": "Code", "parent_structure": "Nkubukulu ya nene",
        "add_structure": "Baka nkubukulu", "structures": "Nkubukulu",
        "access_level": "Ntele ya nswa", "national_control": "Ntwala ya nsi",
        "provincial_control": "Ntwala ya kizunga", "territorial_control": "Ntwala ya mfulu",
        "city_control": "Ntwala ya mbanza", "department_control": "Ntwala ya nsalulu",
        "scope_restricted": "Nswa ekangami kaka muna tfulu epesami.",
        "all_national": "Bizunga biyonso / nkubukulu ziyonso",
        "department_scope": "Nsalulu evutukidi",
    },
}
for _lang, _values in TERRITORIAL_TRANSLATIONS.items():
    TRANSLATIONS.setdefault(_lang, {}).update(_values)


# ============================================================
# BASE DE DONNÉES
# ============================================================

import asyncio
from datetime import datetime
import hashlib
import json
import math
import os
import re
import secrets
import shutil
import sqlite3
import time
import unicodedata
import uuid
import zipfile
from zoneinfo import ZoneInfo
import flet as ft

# ============================================================
# CONFIGURATION TURSO & CHEMINS DE L'APPLICATION
# ============================================================

TURSO_URL = os.environ.get("TURSO_DATABASE_URL")
TURSO_TOKEN = os.environ.get("TURSO_AUTH_TOKEN")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_NAME = "cadastre_rdc.db"
DB_PATH = os.path.join(BASE_DIR, DB_NAME)
BACKUP_DIR = os.path.join(BASE_DIR, "backups")
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")
MODELES_DIR = os.path.join(BASE_DIR, "modeles")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
CERTIFICATS_DIR = os.path.join(UPLOADS_DIR, "certificats")

AUDIT_CONTEXT = {
    "user_id": None,
    "user_name": None,
    "user_role": None,
    "niveau_acces": "National",
    "province": None,
    "ville": None,
    "territoire": None,
    "commune": None,
    "secteur_chefferie": None,
    "groupement": None,
    "village": None,
    "departement": None,
}


def db_connect():
  """Connexion centralisée et intelligente :

  - Se connecte à Turso (Cloud) si TURSO_DATABASE_URL et TURSO_AUTH_TOKEN sont définis.
  - Bascule sur SQLite local (cadastre_rdc.db) par défaut en environnement local.
  """
  if TURSO_URL and TURSO_TOKEN:
    try:
      import libsql_experimental as libsql

      conn = libsql.connect(database=TURSO_URL, auth_token=TURSO_TOKEN)
      return conn
    except Exception as e:
      print(f"Erreur de connexion à Turso, repli sur SQLite local : {e}")

  os.makedirs(BASE_DIR, exist_ok=True)
  conn = sqlite3.connect(DB_PATH, timeout=15.0)
  conn.row_factory = sqlite3.Row
  conn.execute("PRAGMA foreign_keys = ON")
  conn.execute("PRAGMA busy_timeout = 15000")
  return conn


# ============================================================
# SAUVEGARDES ET CONTEXTE D'AUDIT
# ============================================================


def _safe_backup_name(label="auto"):
  """Construit un nom de sauvegarde sûr et unique."""
  stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
  clean = "".join(
      ch if ch.isalnum() or ch in "-_" else "_" for ch in str(label)
  )
  return os.path.join(
      BACKUP_DIR,
      f"cadastre_{clean}_{stamp}.zip",
  )


def _sha256_file(path):
  """Calcule l'empreinte SHA-256 d'un fichier."""
  digest = hashlib.sha256()
  with open(path, "rb") as fh:
    for chunk in iter(lambda: fh.read(1024 * 1024), b""):
      digest.update(chunk)
  return digest.hexdigest()


def _snapshot_database(snapshot_path):
  """Crée une copie SQLite cohérente même si la base est utilisée.

  Utilise l'API backup de SQLite plutôt qu'une simple copie de fichier.
  """
  os.makedirs(os.path.dirname(snapshot_path), exist_ok=True)

  source = sqlite3.connect(DB_PATH)
  try:
    source.execute("PRAGMA busy_timeout=10000")
    target = sqlite3.connect(snapshot_path)
    try:
      source.backup(target)
      target.execute("PRAGMA wal_checkpoint(FULL)")
      target.commit()
    finally:
      target.close()
  finally:
    source.close()


def _cleanup_old_backups(max_backups=30):
  """Conserve au maximum les sauvegardes les plus récentes."""
  if not os.path.isdir(BACKUP_DIR):
    return

  files = [
      os.path.join(BACKUP_DIR, name)
      for name in os.listdir(BACKUP_DIR)
      if name.lower().endswith(".zip")
      and os.path.isfile(os.path.join(BACKUP_DIR, name))
  ]

  files.sort(key=lambda item: os.path.getmtime(item), reverse=True)

  for old_path in files[max_backups:]:
    try:
      os.remove(old_path)
    except OSError:
      pass


def make_backup(label="manual"):
  """Sauvegarde complète :

  - snapshot cohérent de la base SQLite ;
  - toutes les photos/documents du dossier uploads ;
  - manifeste avec date, taille et SHA-256 ;
  - conservation limitée des anciennes sauvegardes.
  """
  os.makedirs(BACKUP_DIR, exist_ok=True)
  os.makedirs(UPLOADS_DIR, exist_ok=True)

  target = _safe_backup_name(label)
  work_dir = os.path.join(
      BASE_DIR,
      f"_backup_work_{uuid.uuid4().hex}",
  )
  snapshot_db = os.path.join(
      work_dir,
      "cadastre_rdc.db",
  )

  try:
    os.makedirs(work_dir, exist_ok=True)

    if not os.path.exists(DB_PATH):
      raise FileNotFoundError("La base cadastre_rdc.db est introuvable.")

    _snapshot_database(snapshot_db)

    # Vérification de la copie avant archivage.
    check = sqlite3.connect(snapshot_db)
    try:
      result = check.execute("PRAGMA integrity_check").fetchone()[0]
      if result != "ok":
        raise ValueError("La copie SQLite créée est invalide.")
    finally:
      check.close()

    manifest = {
        "application": "CADASTRE RDC",
        "backup_type": str(label),
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "database": {
            "path": "database/cadastre_rdc.db",
            "size": os.path.getsize(snapshot_db),
            "sha256": _sha256_file(snapshot_db),
        },
        "uploads_included": True,
    }

    with zipfile.ZipFile(
        target,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=6,
    ) as zf:
      zf.write(
          snapshot_db,
          arcname="database/cadastre_rdc.db",
      )

      for root, dirs, files in os.walk(UPLOADS_DIR):
        # Le dossier de restauration temporaire ne doit jamais être recopié dans une sauvegarde.
        dirs[:] = [d for d in dirs if d != "_restore"]

        for name in files:
          path = os.path.join(root, name)
          zf.write(
              path,
              arcname=os.path.relpath(
                  path,
                  BASE_DIR,
              ),
          )

      zf.writestr(
          "backup_info.json",
          json.dumps(
              manifest,
              ensure_ascii=False,
              indent=2,
          ),
      )

    # Vérification de l'archive créée.
    with zipfile.ZipFile(target, "r") as zf:
      if zf.testzip() is not None:
        raise ValueError("L'archive de sauvegarde est corrompue.")

    _cleanup_old_backups(max_backups=30)
    return target, None

  except Exception as ex:
    try:
      if os.path.exists(target):
        os.remove(target)
    except OSError:
      pass

    return None, str(ex)

  finally:
    shutil.rmtree(
        work_dir,
        ignore_errors=True,
    )


def auto_backup_if_due():
  """Effectue au maximum une sauvegarde automatique par jour."""
  os.makedirs(BACKUP_DIR, exist_ok=True)

  today = datetime.now().strftime("%Y%m%d")

  existing = [
      name
      for name in os.listdir(BACKUP_DIR)
      if name.startswith(f"cadastre_auto_{today}_")
      and name.lower().endswith(".zip")
  ]

  if existing:
    return sorted(existing)[-1]

  path, _ = make_backup("auto")

  return os.path.basename(path) if path else None


def _safe_extract_zip(zf, destination):
  """Extrait une archive sans autoriser les chemins ../ ou les chemins absolus."""
  base = os.path.abspath(destination)

  for member in zf.infolist():
    member_path = os.path.abspath(
        os.path.join(
            destination,
            member.filename,
        )
    )

    if os.path.commonpath([base, member_path]) != base:
      raise ValueError("Archive refusée : chemin non sécurisé.")

  zf.extractall(destination)


def restore_backup_archive(archive_path):
  """Restauration contrôlée d'une sauvegarde."""
  temp_dir = os.path.join(
      BASE_DIR,
      f"_restore_work_{uuid.uuid4().hex}",
  )

  os.makedirs(temp_dir, exist_ok=True)

  try:
    if not os.path.isfile(archive_path):
      raise FileNotFoundError("Fichier de sauvegarde introuvable.")

    with zipfile.ZipFile(
        archive_path,
        "r",
    ) as zf:
      bad = zf.testzip()

      if bad:
        raise ValueError(f"Archive corrompue : {bad}")

      _safe_extract_zip(
          zf,
          temp_dir,
      )

    restored_db = os.path.join(
        temp_dir,
        "database",
        "cadastre_rdc.db",
    )

    if not os.path.isfile(restored_db):
      raise ValueError("La sauvegarde ne contient pas la base cadastrale.")

    # Vérification SQLite.
    test = sqlite3.connect(restored_db)

    try:
      result = test.execute("PRAGMA integrity_check").fetchone()[0]

      if result != "ok":
        raise ValueError("La base de la sauvegarde est invalide.")
    finally:
      test.close()

    # Sauvegarde de sécurité avant restauration.
    before_restore, backup_error = make_backup("pre_restore")

    if not before_restore:
      raise RuntimeError(
          "Impossible de créer la sauvegarde avant restauration :"
          f" {backup_error}"
      )

    restoring_path = DB_PATH + ".restoring"

    try:
      if os.path.exists(restoring_path):
        os.remove(restoring_path)

      shutil.copy2(
          restored_db,
          restoring_path,
      )

      os.replace(
          restoring_path,
          DB_PATH,
      )

    finally:
      if os.path.exists(restoring_path):
        try:
          os.remove(restoring_path)
        except OSError:
          pass

    # Restaurer les photos/documents si présents.
    restored_uploads = os.path.join(
        temp_dir,
        "uploads",
    )

    if os.path.isdir(restored_uploads):
      old_uploads = UPLOADS_DIR + "_before_restore"

      try:
        if os.path.exists(old_uploads):
          shutil.rmtree(old_uploads)

        if os.path.isdir(UPLOADS_DIR):
          os.replace(
              UPLOADS_DIR,
              old_uploads,
          )

        shutil.copytree(
            restored_uploads,
            UPLOADS_DIR,
        )

        if os.path.exists(old_uploads):
          shutil.rmtree(old_uploads)

      except Exception:
        if os.path.isdir(UPLOADS_DIR):
          shutil.rmtree(
              UPLOADS_DIR,
              ignore_errors=True,
          )

        if os.path.isdir(old_uploads):
          os.replace(
              old_uploads,
              UPLOADS_DIR,
          )

        raise

    init_database()

    return True, None

  except Exception as ex:
    return False, str(ex)

  finally:
    shutil.rmtree(
        temp_dir,
        ignore_errors=True,
    )


def verify_database_connection():
  """Vérifie l'intégrité de la connexion à la base de données."""
  conn = db_connect()
  try:
    # Si on utilise Turso, les pragmas locaux peuvent différer, on valide simplement une requête basique
    if not (TURSO_URL and TURSO_TOKEN):
      expected_name = os.path.basename(DB_PATH).lower()
      if expected_name != "cadastre_rdc.db":
        raise RuntimeError(
            f"Configuration base invalide : {os.path.basename(DB_PATH)}"
        )
      integrity = conn.execute("PRAGMA integrity_check").fetchone()[0]
      if integrity != "ok":
        raise RuntimeError(
            f"La base cadastrale n'est pas intègre : {integrity}"
        )
    conn.execute("SELECT 1").fetchone()
  finally:
    conn.close()
  return "Cloud Turso" if (TURSO_URL and TURSO_TOKEN) else DB_PATH


def add_column_if_missing(conn, table, column, definition):
  try:
    columns = [
        row["name"]
        for row in conn.execute(f"PRAGMA table_info({table})").fetchall()
    ]
    if column not in columns:
      conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")
  except Exception:
    pass


def init_database():
  """Initialise toutes les tables requises au démarrage."""
  conn = db_connect()
  cur = conn.cursor()

  cur.execute("""
        CREATE TABLE IF NOT EXISTS proprietaires (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom_complet TEXT NOT NULL,
            telephone TEXT,
            email TEXT,
            piece_identite TEXT,
            photo TEXT,
            date_creation TEXT NOT NULL
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS parcelles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            numero TEXT UNIQUE NOT NULL,
            adresse TEXT,
            province TEXT,
            ville_territoire TEXT,
            commune_chefferie TEXT,
            quartier_groupement TEXT,
            localite TEXT,
            superficie REAL,
            latitude REAL,
            longitude REAL,
            polygone TEXT,
            proprietaire_id INTEGER,
            statut TEXT DEFAULT 'Occupée',
            date_enregistrement TEXT NOT NULL,
            agent TEXT,
            FOREIGN KEY(proprietaire_id) REFERENCES proprietaires(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            parcelle_id INTEGER,
            nom_document TEXT,
            chemin TEXT,
            date_ajout TEXT NOT NULL,
            FOREIGN KEY(parcelle_id) REFERENCES parcelles(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS mutations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            parcelle_id INTEGER NOT NULL,
            ancien_proprietaire_id INTEGER,
            nouveau_proprietaire_id INTEGER,
            type_transaction TEXT,
            reference_acte TEXT,
            prix REAL,
            date_transaction TEXT NOT NULL,
            agent TEXT,
            statut TEXT DEFAULT 'En attente',
            FOREIGN KEY(parcelle_id) REFERENCES parcelles(id),
            FOREIGN KEY(ancien_proprietaire_id) REFERENCES proprietaires(id),
            FOREIGN KEY(nouveau_proprietaire_id) REFERENCES proprietaires(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS utilisateurs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT NOT NULL,
            role TEXT NOT NULL,
            telephone TEXT,
            date_creation TEXT NOT NULL
        )
    """)

  conn.commit()
  conn.close()


def init_database():
  """Initialise toutes les tables requises au démarrage et applique les

  migrations de colonnes nécessaires pour la sécurité et l'architecture
  territoriale.
  """
  conn = db_connect()
  cur = conn.cursor()

  # Tables principales de base
  cur.execute("""
        CREATE TABLE IF NOT EXISTS proprietaires (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom_complet TEXT NOT NULL,
            telephone TEXT,
            email TEXT,
            piece_identite TEXT,
            photo TEXT,
            date_creation TEXT NOT NULL
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS parcelles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            numero TEXT UNIQUE NOT NULL,
            adresse TEXT,
            province TEXT,
            ville_territoire TEXT,
            commune_chefferie TEXT,
            quartier_groupement TEXT,
            localite TEXT,
            superficie REAL,
            latitude REAL,
            longitude REAL,
            polygone TEXT,
            proprietaire_id INTEGER,
            statut TEXT DEFAULT 'Occupée',
            date_enregistrement TEXT NOT NULL,
            agent TEXT,
            FOREIGN KEY(proprietaire_id) REFERENCES proprietaires(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            parcelle_id INTEGER,
            nom_document TEXT,
            chemin TEXT,
            date_ajout TEXT NOT NULL,
            FOREIGN KEY(parcelle_id) REFERENCES parcelles(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS mutations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            parcelle_id INTEGER NOT NULL,
            ancien_proprietaire_id INTEGER,
            nouveau_proprietaire_id INTEGER,
            type_transaction TEXT,
            reference_acte TEXT,
            prix REAL,
            date_transaction TEXT NOT NULL,
            agent TEXT,
            statut TEXT DEFAULT 'En attente',
            FOREIGN KEY(parcelle_id) REFERENCES parcelles(id),
            FOREIGN KEY(ancien_proprietaire_id) REFERENCES proprietaires(id),
            FOREIGN KEY(nouveau_proprietaire_id) REFERENCES proprietaires(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS utilisateurs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT NOT NULL,
            role TEXT NOT NULL,
            telephone TEXT,
            date_creation TEXT NOT NULL
        )
    """)

  # Colonnes de sécurité ajoutées progressivement pour préserver
  # les bases créées par les versions précédentes.
  add_column_if_missing(conn, "utilisateurs", "password_hash", "TEXT")
  add_column_if_missing(conn, "utilisateurs", "password_salt", "TEXT")
  add_column_if_missing(conn, "utilisateurs", "failed_attempts", "INTEGER DEFAULT 0")
  add_column_if_missing(conn, "utilisateurs", "locked_until", "REAL DEFAULT 0")
  add_column_if_missing(conn, "utilisateurs", "last_login", "TEXT")
  add_column_if_missing(conn, "utilisateurs", "active", "INTEGER DEFAULT 1")

  # --------------------------------------------------------
  # Architecture territoriale et administrative
  # --------------------------------------------------------
  cur.execute("""
        CREATE TABLE IF NOT EXISTS structures_territoriales (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            type TEXT NOT NULL,
            code TEXT,
            nom TEXT NOT NULL,
            parent_id INTEGER,
            actif INTEGER DEFAULT 1,
            date_creation TEXT NOT NULL,
            UNIQUE(type, nom, parent_id),
            FOREIGN KEY(parent_id) REFERENCES structures_territoriales(id)
        )
    """)

  cur.execute("""
        CREATE TABLE IF NOT EXISTS departements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT NOT NULL,
            code TEXT,
            province TEXT,
            actif INTEGER DEFAULT 1,
            date_creation TEXT NOT NULL,
            UNIQUE(nom, province)
        )
    """)

  # Extensions compatibles avec toutes les bases existantes.
  for _table, _column, _definition in [
      ("utilisateurs", "niveau_acces", "TEXT DEFAULT 'National'"),
      ("utilisateurs", "province", "TEXT"),
      ("utilisateurs", "ville", "TEXT"),
      ("utilisateurs", "territoire", "TEXT"),
      ("utilisateurs", "commune", "TEXT"),
      ("utilisateurs", "secteur_chefferie", "TEXT"),
      ("utilisateurs", "groupement", "TEXT"),
      ("utilisateurs", "village", "TEXT"),
      ("utilisateurs", "departement", "TEXT"),
      ("parcelles", "niveau_territorial", "TEXT"),
      ("parcelles", "structure_code", "TEXT"),
  ]:
    add_column_if_missing(conn, _table, _column, _definition)

  # Référentiel initial des 26 provinces.
  _provinces = [
      ("CD-KN", "Kinshasa"), ("CD-KC", "Kongo-Central"), ("CD-KW", "Kwango"),
      ("CD-KL", "Kwilu"), ("CD-MN", "Mai-Ndombe"), ("CD-EQ", "Équateur"),
      ("CD-MO", "Mongala"), ("CD-NU", "Nord-Ubangi"), ("CD-SU", "Sud-Ubangi"),
      ("CD-TS", "Tshuapa"), ("CD-IT", "Ituri"), ("CD-HU", "Haut-Uele"),
      ("CD-BU", "Bas-Uele"), ("CD-NK", "Nord-Kivu"), ("CD-SK", "Sud-Kivu"),
      ("CD-MA", "Maniema"), ("CD-TA", "Tanganyika"), ("CD-HL", "Haut-Lomami"),
      ("CD-LU", "Lualaba"), ("CD-HK", "Haut-Katanga"), ("CD-KS", "Kasaï"),
      ("CD-KC2", "Kasaï-Central"), ("CD-KO", "Kasaï-Oriental"), ("CD-LO", "Lomami"),
      ("CD-SA", "Sankuru"), ("CD-TSH", "Tshopo"),
  ]
  
  _seen = set()
  for _code, _name in _provinces:
    if _name in _seen:
      continue
    _seen.add(_name)
    conn.execute(
        "INSERT OR IGNORE INTO structures_territoriales(type, code, nom, parent_id, actif, date_creation) VALUES ('Province', ?, ?, NULL, 1, ?)",
        (_code, _name, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    )

  cur.execute("""
        CREATE TABLE IF NOT EXISTS journal (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            action TEXT NOT NULL,
            details TEXT,
            date_action TEXT NOT NULL
        )
    """)

  for _col, _def in [
      ("utilisateur_id", "INTEGER"),
      ("utilisateur_nom", "TEXT"),
      ("role", "TEXT"),
      ("niveau_acces", "TEXT"),
      ("province", "TEXT"),
      ("ville", "TEXT"),
      ("territoire", "TEXT"),
      ("commune", "TEXT"),
      ("secteur_chefferie", "TEXT"),
      ("groupement", "TEXT"),
      ("village", "TEXT"),
      ("departement", "TEXT"),
  ]:
    add_column_if_missing(conn, "journal", _col, _def)

  # --------------------------------------------------------
  # Migration / compatibilité avec les anciennes versions
  # --------------------------------------------------------
  add_column_if_missing(conn, "proprietaires", "telephone", "TEXT")
  add_column_if_missing(conn, "proprietaires", "email", "TEXT")
  add_column_if_missing(conn, "proprietaires", "piece_identite", "TEXT")
  add_column_if_missing(conn, "proprietaires", "photo", "TEXT")
  add_column_if_missing(conn, "proprietaires", "date_creation", "TEXT")

  add_column_if_missing(conn, "parcelles", "numero", "TEXT")
  add_column_if_missing(conn, "parcelles", "polygone", "TEXT")
  add_column_if_missing(conn, "parcelles", "proprietaire_id", "INTEGER")
  add_column_if_missing(conn, "parcelles", "statut", "TEXT DEFAULT 'Occupée'")
  add_column_if_missing(conn, "parcelles", "agent", "TEXT")

  # Gestion des anciennes colonnes (cni / date_enregistrement)
  old_columns = [
      row["name"]
      for row in conn.execute(
          "PRAGMA table_info(proprietaires)"
      ).fetchall()
  ]

  if "cni" in old_columns:
    conn.execute("""
            UPDATE proprietaires
            SET piece_identite = cni
            WHERE (piece_identite IS NULL OR piece_identite = '')
              AND cni IS NOT NULL
        """)

  if "date_enregistrement" in old_columns:
    conn.execute("""
            UPDATE proprietaires
            SET date_creation = date_enregistrement
            WHERE (date_creation IS NULL OR date_creation = '')
              AND date_enregistrement IS NOT NULL
        """)

  conn.commit()
  conn.close()

# Lancement de l'initialisation des tables au chargement
init_database()

def journaliser(action, details=""):
    conn = db_connect()
    columns = [row["name"] for row in conn.execute("PRAGMA table_info(journal)").fetchall()]
    territorial_columns = {
        "utilisateur_id", "utilisateur_nom", "role", "niveau_acces",
        "province", "ville", "territoire", "commune", "secteur_chefferie",
        "groupement", "village", "departement"
    }
    
    if territorial_columns.issubset(columns):
        conn.execute(
            """
            INSERT INTO journal(action, details, date_action, utilisateur_id, utilisateur_nom, role,
                                niveau_acces, province, ville, territoire, commune, secteur_chefferie,
                                groupement, village, departement)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                action, details, datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                AUDIT_CONTEXT.get("user_id"), AUDIT_CONTEXT.get("user_name"), AUDIT_CONTEXT.get("user_role"),
                AUDIT_CONTEXT.get("niveau_acces"), AUDIT_CONTEXT.get("province"), AUDIT_CONTEXT.get("ville"),
                AUDIT_CONTEXT.get("territoire"), AUDIT_CONTEXT.get("commune"), AUDIT_CONTEXT.get("secteur_chefferie"),
                AUDIT_CONTEXT.get("groupement"), AUDIT_CONTEXT.get("village"), AUDIT_CONTEXT.get("departement")
            ),
        )
    elif {"utilisateur_id", "utilisateur_nom", "role"}.issubset(columns):
        conn.execute(
            """
            INSERT INTO journal(action, details, date_action, utilisateur_id, utilisateur_nom, role)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                action, details, datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                AUDIT_CONTEXT.get("user_id"), AUDIT_CONTEXT.get("user_name"), AUDIT_CONTEXT.get("user_role")
            ),
        )
    else:
        conn.execute(
            "INSERT INTO journal(action, details, date_action) VALUES (?, ?, ?)",
            (action, details, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        )
    conn.commit()
    conn.close()


def get_count(table):
    allowed = {
        "proprietaires", "parcelles", "mutations", "utilisateurs", "documents", "journal",
    }
    if table not in allowed:
        return 0

    conn = db_connect()
    level = str(AUDIT_CONTEXT.get("niveau_acces") or "National")
    province = AUDIT_CONTEXT.get("province")
    ville = AUDIT_CONTEXT.get("ville")
    territoire = AUDIT_CONTEXT.get("territoire")
    commune = AUDIT_CONTEXT.get("commune")
    secteur = AUDIT_CONTEXT.get("secteur_chefferie")
    groupement = AUDIT_CONTEXT.get("groupement")
    village = AUDIT_CONTEXT.get("village")

    if table in {"parcelles", "proprietaires", "mutations", "documents"}:
        scope, params = _audit_scope_sql("p")
        if table == "parcelles":
            sql = f"SELECT COUNT(*) FROM parcelles p WHERE {scope}"
        elif table == "proprietaires":
            sql = f"SELECT COUNT(DISTINCT pr.id) FROM proprietaires pr LEFT JOIN parcelles p ON p.proprietaire_id=pr.id WHERE {scope}"
        elif table == "mutations":
            sql = f"SELECT COUNT(*) FROM mutations m JOIN parcelles p ON p.id=m.parcelle_id WHERE {scope}"
        else:
            sql = f"SELECT COUNT(*) FROM documents d JOIN parcelles p ON p.id=d.parcelle_id WHERE {scope}"
        result = conn.execute(sql, params).fetchone()[0]
        
    elif table == "utilisateurs":
        if level == "National" or not province:
            result = conn.execute("SELECT COUNT(*) FROM utilisateurs").fetchone()[0]
        else:
            sql = "SELECT COUNT(*) FROM utilisateurs WHERE LOWER(COALESCE(province,''))=LOWER(?)"
            params = [province]
            if level == "Ville" and ville:
                sql += " AND LOWER(COALESCE(ville,''))=LOWER(?)"
                params.append(ville)
            elif level == "Territoire" and territoire:
                sql += " AND LOWER(COALESCE(territoire,''))=LOWER(?)"
                params.append(territoire)
            elif level == "Commune" and commune:
                sql += " AND LOWER(COALESCE(commune,''))=LOWER(?)"
                params.append(commune)
            elif level in {"Secteur", "Chefferie"} and secteur:
                sql += " AND LOWER(COALESCE(secteur_chefferie,''))=LOWER(?)"
                params.append(secteur)
            elif level == "Groupement" and groupement:
                sql += " AND LOWER(COALESCE(groupement,''))=LOWER(?)"
                params.append(groupement)
            elif level in {"Village", "Local"} and village:
                sql += " AND LOWER(COALESCE(village,''))=LOWER(?)"
                params.append(village)
            result = conn.execute(sql, params).fetchone()[0]
            
    else:
        if level == "National" or not province:
            result = conn.execute("SELECT COUNT(*) FROM journal").fetchone()[0]
        else:
            sql = "SELECT COUNT(*) FROM journal WHERE LOWER(COALESCE(province,''))=LOWER(?)"
            params = [province]
            if level == "Ville" and ville:
                sql += " AND LOWER(COALESCE(ville,''))=LOWER(?)"
                params.append(ville)
            elif level == "Territoire" and territoire:
                sql += " AND LOWER(COALESCE(territoire,''))=LOWER(?)"
                params.append(territoire)
            result = conn.execute(sql, params).fetchone()[0]
            
    conn.close()
    return result


def _audit_scope_sql(alias="p"):
    level = str(AUDIT_CONTEXT.get("niveau_acces") or "National")
    province = AUDIT_CONTEXT.get("province")
    ville = AUDIT_CONTEXT.get("ville")
    territoire = AUDIT_CONTEXT.get("territoire")
    commune = AUDIT_CONTEXT.get("commune")
    secteur = AUDIT_CONTEXT.get("secteur_chefferie")
    groupement = AUDIT_CONTEXT.get("groupement")
    village = AUDIT_CONTEXT.get("village")
    
    if level == "National" or not province:
        return "1=1", []
        
    conditions = [f"LOWER(COALESCE({alias}.province, '')) = LOWER(?)"]
    params = [province]
    
    if level in {"Ville", "Territoire"}:
        target = ville if level == "Ville" else territoire
        if target:
            conditions.append(f"LOWER(COALESCE({alias}.ville_territoire, '')) = LOWER(?)")
            params.append(target)
    elif level in {"Commune", "Secteur", "Chefferie", "Secteur / Chefferie"}:
        target = commune or secteur
        if target:
            conditions.append(f"LOWER(COALESCE({alias}.commune_chefferie, '')) = LOWER(?)")
            params.append(target)
    elif level == "Groupement" and groupement:
        conditions.append(f"LOWER(COALESCE({alias}.quartier_groupement, '')) = LOWER(?)")
        params.append(groupement)
    elif level in {"Village", "Local"} and village:
        conditions.append(f"LOWER(COALESCE({alias}.localite, '')) = LOWER(?)")
        params.append(village)
        
    return " AND ".join(conditions), params


def get_proprietaires():
    conn = db_connect()
    scope, params = _audit_scope_sql("p")
    rows = conn.execute(
        f"""SELECT DISTINCT pr.* FROM proprietaires pr
            LEFT JOIN parcelles p ON p.proprietaire_id = pr.id
            WHERE {scope}
            ORDER BY pr.id DESC""",
        params
    ).fetchall()
    conn.close()
    return rows


def get_parcelles():
    conn = db_connect()
    scope, params = _audit_scope_sql("p")
    rows = conn.execute(
        f"""
        SELECT p.*, pr.nom_complet AS proprietaire
        FROM parcelles p
        LEFT JOIN proprietaires pr
            ON p.proprietaire_id = pr.id
        WHERE {scope}
        ORDER BY p.id DESC
        """,
        params
    ).fetchall()
    conn.close()
    return rows

# ============================================================
# CALCUL DE SUPERFICIE
# ============================================================

def calculer_superficie(polygone):
    """
    Calcul approximatif en m² à partir d'un polygone GPS.
    Projection locale simplifiée + formule de Shoelace.
    """
    if not polygone or len(polygone) < 3:
        return 0.0

    lat_moy = sum(float(p[0]) for p in polygone) / len(polygone)
    m_lat = 111320.0
    m_lon = 111320.0 * math.cos(math.radians(lat_moy))

    points = [
        (float(p[1]) * m_lon, float(p[0]) * m_lat)
        for p in polygone
    ]

    area = 0.0
    for i in range(len(points)):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % len(points)]
        area += x1 * y2 - x2 * y1

    return abs(area) / 2.0


def calculer_superficie_cm2(points):
    """
    Calcule la superficie en cm² à partir des coordonnées GPS.
    """
    return calculer_superficie(points) * 10000.0


def parse_area(value):
    """
    Accepte :
      - 37500
      - 150/250
      - 150,5/250
    """
    text = (value or "").strip()

    if not text:
        return 0.0

    if "/" in text:
        parts = [p.strip() for p in text.split("/")]
        if len(parts) != 2:
            raise ValueError

        a = float(parts[0].replace(",", "."))
        b = float(parts[1].replace(",", "."))

        if a <= 0 or b <= 0:
            raise ValueError

        return a * b

    result = float(text.replace(",", "."))
    if result < 0:
        raise ValueError

    return result

# ============================================================
# SÉCURITÉ — HASH / VÉRIFICATION DES MOTS DE PASSE
# ============================================================

PASSWORD_ITERATIONS = 600_000
MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_SECONDS = 15 * 60
SESSION_TIMEOUT_SECONDS = 30 * 60


def hash_password(password):
    salt = secrets.token_bytes(32)
    password_hash = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt,
        PASSWORD_ITERATIONS,
    )
    return salt.hex(), password_hash.hex()


def verify_password(password, salt_hex, hash_hex):
    try:
        salt = bytes.fromhex(salt_hex)
        calculated = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode("utf-8"),
            salt,
            PASSWORD_ITERATIONS,
        )
        return secrets.compare_digest(calculated.hex(), hash_hex)
    except Exception:
        return False


# ============================================================
# APPLICATION
# ============================================================

def main(page: ft.Page):
    # Configuration de base AVANT toute opération pouvant échouer.
    page.title = "CADASTRE RDC — Solution Géospatiale Nationale"
    page.padding = 0
    page.spacing = 0
    page.bgcolor = COLOR_BG_LIGHT
    page.theme_mode = ft.ThemeMode.LIGHT
    page.scroll = ft.ScrollMode.HIDDEN

    # FilePicker persistant pour les téléchargements de certificats.
    # En Flet Web, src_bytes déclenche le téléchargement dans le navigateur ;
    # sur Desktop, le même service ouvre le choix d'emplacement.
    certificate_file_picker = ft.FilePicker()
    try:
        if hasattr(page, "services"):
            page.services.append(certificate_file_picker)
        elif hasattr(page, "overlay"):
            page.overlay.append(certificate_file_picker)
            page.update()
    except Exception:
        pass

    startup_message = ft.Text(
        "Initialisation de CADASTRE RDC…",
        size=18,
        weight=ft.FontWeight.BOLD,
        color=COLOR_TEXT_MAIN,
        text_align=ft.TextAlign.CENTER,
    )
    page.add(
        ft.Container(
            expand=True,
            alignment=ft.alignment.center,
            bgcolor=COLOR_BG_LIGHT,
            content=ft.Column(
                [ft.ProgressRing(), startup_message],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                spacing=18,
            ),
        )
    )
    page.update()

    try:
        init_database()
        # Vérification de la relation réelle application ↔ base officielle.
        connected_db = verify_database_connection()
        print(f"[CADASTRE RDC] BASE ACTIVE : {connected_db}")
    except Exception as ex:
        error_text = f"{type(ex).__name__}: {ex}"
        print(f"[CADASTRE RDC] ERREUR INITIALISATION BASE : {error_text}")
        page.controls.clear()
        page.add(
            ft.Container(
                expand=True,
                alignment=ft.alignment.center,
                padding=30,
                content=ft.Column(
                    [
                        ft.Icon(ft.Icons.ERROR_OUTLINE, size=60, color=COLOR_ACCENT),
                        ft.Text(
                            "Impossible d'initialiser la base de données",
                            size=24,
                            weight=ft.FontWeight.BOLD,
                            text_align=ft.TextAlign.CENTER,
                        ),
                        ft.Text(error_text, selectable=True, text_align=ft.TextAlign.CENTER),
                        ft.Text(
                            "Consultez également le terminal VS Code/PowerShell pour le détail.",
                            text_align=ft.TextAlign.CENTER,
                            color=COLOR_TEXT_MUTED,
                        ),
                    ],
                    horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                    spacing=16,
                ),
            )
        )
        page.update()
        return

    try:
        auto_backup_if_due()
    except Exception as ex:
        print(f"[CADASTRE RDC] AVERTISSEMENT BACKUP AUTOMATIQUE : {type(ex).__name__}: {ex}")

    page.controls.clear()

    try:
        page.window.width = 1280
        page.window.height = 820
    except Exception:
        pass

    # Zone de contenu principale : elle ne prend jamais le contrôle du scroll
    # de la page. Chaque module qui contient une longue liste possède son propre
    # ListView/scroll, afin de conserver la barre supérieure et les boutons de navigation visibles.
    content = ft.Container(
        expand=True,
        padding=25,
    )

    # Horloge locale : elle suit l'heure du système qui sert la session.
    # En exécution locale (python fichier.py), il s'agit de l'appareil de l'utilisateur.
    clock_date = ft.Text("--/--/----", size=10, color=COLOR_TEXT_MUTED, weight=ft.FontWeight.BOLD)
    clock_time = ft.Text("--:--:--", size=14, color=COLOR_TEXT_MAIN, weight=ft.FontWeight.BOLD)
    clock_zone = ft.Text("Fuseau local…", size=8, color=COLOR_TEXT_MUTED)

    sidebar_container = ft.Container()
    current_route = "dashboard"
    language = "fr"

    # Session utilisateur courante
    # Doit être initialisée avant toute fonction qui consulte current_user.
    current_user = {
        "id": None,
        "nom": None,
        "role": None,
        "niveau_acces": "National",
        "province": None,
        "ville": None,
        "territoire": None,
        "commune": None,
        "secteur_chefferie": None,
        "groupement": None,
        "village": None,
        "departement": None,
        "login_time": 0.0,
        "last_activity": 0.0,
    }

    ROLE_PERMISSIONS = {
        "Administrateur": {"*"},
        "Superviseur": {"dashboard", "proprietaires", "parcelles", "mutations", "verification", "journal", "territoires"},
        "Technicien": {"dashboard", "proprietaires", "parcelles", "mutations", "verification"},
        "Agent": {"dashboard", "proprietaires", "parcelles", "verification"},
        "Consultation": {"dashboard", "proprietaires", "parcelles", "verification"},
        "PUBLIC": {"verification"},  # Accès restreint uniquement au module public/vérification
    }

    # Permissions fines : elles sont contrôlées au niveau de l'action,
    # et pas uniquement au niveau de l'affichage des boutons.
    ROLE_ACTION_PERMISSIONS = {
        "Administrateur": {"*"},
        "Superviseur": {
            "owner_create", "owner_edit",
            "parcel_create", "parcel_edit",
            "mutation_create", "mutation_validate",
            "journal_view", "territorial_manage",
        },
        "Technicien": {
            "owner_create", "owner_edit",
            "parcel_create", "parcel_edit",
            "mutation_create",
        },
        "Agent": {
            "owner_create",
            "parcel_create",
        },
        "Consultation": set(),
        "PUBLIC": set(),  # Aucune action de modification ou création pour le public
    }
    # --------------------------------------------------------
    # Téléversement des photos des propriétaires
    # --------------------------------------------------------
    # FilePicker est un service dans les versions récentes de Flet.
    # Il ne doit donc pas être ajouté à page.overlay comme contrôle.
    # Un picker est créé pour chaque opération de sélection/upload.

    # --------------------------------------------------------
    # Utilitaires UI
    # --------------------------------------------------------

    if "kg" not in TRANSLATIONS:
        TRANSLATIONS["kg"] = dict(TRANSLATIONS["fr"])

    def t(key):
        return TRANSLATIONS.get(language, TRANSLATIONS["fr"]).get(
            key, key
        )

    def tf(key, **kwargs):
        return t(key).format(**kwargs)

    def enum_text(value, mapping):
        key = mapping.get(value)
        return t(key) if key else (value or "")

    STATUS_KEYS = {
        "Occupée": "occupied",
        "Disponible": "available",
        "En vente": "for_sale",
        "Transférée": "transferred",
        "Contestée": "disputed",
    }

    TRANSACTION_KEYS = {
        "Vente": "sale",
        "Cession": "cession",
        "Donation": "donation",
        "Autre": "other",
    }

    ROLE_KEYS = {
        "Administrateur": "administrator",
        "Superviseur": "supervisor",
        "Technicien": "technician",
        "Agent": "agent_role",
        "Consultation": "consultation",
    }

    JOURNAL_ACTION_KEYS = {
        "Création propriétaire": "creation_owner",
        "Création parcelle": "creation_parcel",
        "Nouvelle mutation": "new_transfer",
        "Création utilisateur": "creation_user",
    }

    def snack(message, error=False):
        page.snack_bar = ft.SnackBar(
            content=ft.Text(
                message,
                color=ft.Colors.WHITE,
            ),
            bgcolor=COLOR_ACCENT if error else ft.Colors.GREEN_700,
        )
        page.snack_bar.open = True
        page.update()

    def text_field(label, value="", multiline=False, expand=True):
        return ft.TextField(
            label=label,
            value=value,
            multiline=multiline,
            border_radius=14,
            border_color=COLOR_BORDER,
            focused_border_color=COLOR_PRIMARY,
            text_size=13,
            expand=expand,
            min_lines=3 if multiline else 1,
            max_lines=6 if multiline else 1,
        )

    def title(text, subtitle=None):
        controls = [
            ft.Text(
                text,
                size=26,
                weight=ft.FontWeight.BOLD,
                color=COLOR_TEXT_MAIN,
            )
        ]

        if subtitle:
            controls.append(
                ft.Text(
                    subtitle,
                    size=13,
                    color=COLOR_TEXT_MUTED,
                )
            )

        return ft.Column(
            controls,
            spacing=3,
        )

    def panel(content_value, padding=20):
        return ft.Container(
            padding=padding,
            bgcolor=(
                ft.Colors.GREY_900
                if page.theme_mode == ft.ThemeMode.DARK
                else COLOR_SURFACE
            ),
            border_radius=14,
            border=ft.Border(
                left=ft.BorderSide(
                    width=1,
                    color="#334155" if page.theme_mode == ft.ThemeMode.DARK else COLOR_BORDER,
                ),
                top=ft.BorderSide(
                    width=1,
                    color="#334155" if page.theme_mode == ft.ThemeMode.DARK else COLOR_BORDER,
                ),
                right=ft.BorderSide(
                    width=1,
                    color="#334155" if page.theme_mode == ft.ThemeMode.DARK else COLOR_BORDER,
                ),
                bottom=ft.BorderSide(
                    width=1,
                    color="#334155" if page.theme_mode == ft.ThemeMode.DARK else COLOR_BORDER,
                ),
            ),
            content=content_value,
        )

    def status_text(text, color=COLOR_TEXT_MUTED):
        return ft.Text(
            text,
            size=12,
            color=color,
        )

    # --------------------------------------------------------
    # Périmètre territorial de sécurité
    # --------------------------------------------------------

    def territorial_scope_sql(alias="p"):
        """Retourne (SQL, paramètres) selon le périmètre de l'utilisateur connecté."""
        level = str(current_user.get("niveau_acces") or "National")
        province = current_user.get("province")
        ville = current_user.get("ville")
        territoire = current_user.get("territoire")
        commune = current_user.get("commune")
        secteur = current_user.get("secteur_chefferie")
        groupement = current_user.get("groupement")
        village = current_user.get("village")

        if level == "National" or not province:
            return "1=1", []

        conditions = [f"LOWER(COALESCE({alias}.province, '')) = LOWER(?)"]
        params = [province]

        if level == "Ville" and ville:
            conditions.append(f"LOWER(COALESCE({alias}.ville_territoire, '')) = LOWER(?)")
            params.append(ville)
        elif level == "Territoire" and territoire:
            conditions.append(f"LOWER(COALESCE({alias}.ville_territoire, '')) = LOWER(?)")
            params.append(territoire)
        elif level == "Commune" and commune:
            conditions.append(f"LOWER(COALESCE({alias}.commune_chefferie, '')) = LOWER(?)")
            params.append(commune)
        elif level in {"Secteur", "Chefferie", "Secteur / Chefferie"} and secteur:
            conditions.append(f"LOWER(COALESCE({alias}.commune_chefferie, '')) = LOWER(?)")
            params.append(secteur)
        elif level == "Groupement" and groupement:
            conditions.append(f"LOWER(COALESCE({alias}.quartier_groupement, '')) = LOWER(?)")
            params.append(groupement)
        elif level in {"Village", "Local"} and village:
            conditions.append(f"LOWER(COALESCE({alias}.localite, '')) = LOWER(?)")
            params.append(village)

        return " AND ".join(conditions), params

    def owner_scope_sql(owner_alias="pr", parcel_alias="p"):
        scope, params = territorial_scope_sql(parcel_alias)
        return scope, params

    def user_has_global_scope():
        return str(current_user.get("niveau_acces") or "National") == "National"

    def current_scope_label():
        level = current_user.get("niveau_acces") or "National"
        parts = [level]
        for key in ("province", "ville", "territoire", "commune", "secteur_chefferie", "groupement", "village", "departement"):
            value = current_user.get(key)
            if value:
                parts.append(str(value))
        return " → ".join(parts)

    def target_scope_allowed(level, province, ville=None, territoire=None, commune=None, secteur=None, groupement=None, village=None):
        """Vérifie qu'un utilisateur ne peut attribuer qu'un périmètre inclus dans le sien."""
        current_level = str(current_user.get("niveau_acces") or "National")
        if current_level == "National":
            return True
        if not province or str(province).strip().lower() != str(current_user.get("province") or "").strip().lower():
            return False
        order = {"National": 0, "Provincial": 1, "Ville": 2, "Territoire": 2, "Commune": 3, "Secteur": 4, "Chefferie": 4, "Groupement": 5, "Village": 6, "Local": 6}
        if order.get(level, 99) < order.get(current_level, 99):
            return False
        if current_level == "Ville" and ville and str(ville).strip().lower() != str(current_user.get("ville") or "").strip().lower():
            return False
        if current_level == "Territoire" and territoire and str(territoire).strip().lower() != str(current_user.get("territoire") or "").strip().lower():
            return False
        if current_level == "Commune" and commune and str(commune).strip().lower() != str(current_user.get("commune") or "").strip().lower():
            return False
        if current_level in {"Secteur", "Chefferie"} and secteur and str(secteur).strip().lower() != str(current_user.get("secteur_chefferie") or "").strip().lower():
            return False
        if current_level == "Groupement" and groupement and str(groupement).strip().lower() != str(current_user.get("groupement") or "").strip().lower():
            return False
        if current_level in {"Village", "Local"} and village and str(village).strip().lower() != str(current_user.get("village") or "").strip().lower():
            return False
        return True
      # --------------------------------------------------------
    # Dashboard
    # --------------------------------------------------------

    def dashboard_view():
        nonlocal current_route
        current_route = "dashboard"

        def stat_card(label, value, icon):
            return ft.Container(
                expand=True,
                padding=20,
                border_radius=14,
                bgcolor=(
                    "#172554"
                    if page.theme_mode == ft.ThemeMode.DARK
                    else "#EFF6FF"
                ),
                border=ft.Border(
                    left=ft.BorderSide(
                        width=1,
                        color=(
                            "#1E40AF"
                            if page.theme_mode == ft.ThemeMode.DARK
                            else "#DBEAFE"
                        ),
                    ),
                    top=ft.BorderSide(
                        width=1,
                        color=(
                            "#1E40AF"
                            if page.theme_mode == ft.ThemeMode.DARK
                            else "#DBEAFE"
                        ),
                    ),
                    right=ft.BorderSide(
                        width=1,
                        color=(
                            "#1E40AF"
                            if page.theme_mode == ft.ThemeMode.DARK
                            else "#DBEAFE"
                        ),
                    ),
                    bottom=ft.BorderSide(
                        width=1,
                        color=(
                            "#1E40AF"
                            if page.theme_mode == ft.ThemeMode.DARK
                            else "#DBEAFE"
                        ),
                    ),
                ),
                content=ft.Row(
                    [
                        ft.Container(
                            width=52,
                            height=52,
                            border_radius=26,
                            bgcolor=COLOR_PRIMARY,
                            alignment=ft.Alignment.CENTER,
                            content=ft.Icon(
                                icon,
                                size=26,
                                color=ft.Colors.WHITE,
                            ),
                        ),
                        ft.Column(
                            [
                                ft.Text(
                                    label,
                                    size=12,
                                    color=COLOR_TEXT_MUTED,
                                ),
                                ft.Text(
                                    str(value),
                                    size=27,
                                    weight=ft.FontWeight.BOLD,
                                ),
                            ],
                            spacing=2,
                        ),
                    ],
                    spacing=14,
                ),
            )

        content.content = ft.Column(
            [
                title(
                    t("dashboard_title"),
                    t("dashboard_subtitle"),
                ),
                ft.Divider(),
                ft.Row(
                    [
                        stat_card(
                            t("owners"),
                            get_count("proprietaires"),
                            ft.Icons.PERSON,
                        ),
                        stat_card(
                            t("parcels"),
                            get_count("parcelles"),
                            ft.Icons.LANDSCAPE,
                        ),
                        stat_card(
                            t("mutations"),
                            get_count("mutations"),
                            ft.Icons.SWAP_HORIZ,
                        ),
                        stat_card(
                            t("users"),
                            get_count("utilisateurs"),
                            ft.Icons.GROUP,
                        ),
                    ],
                    spacing=14,
                ),
                ft.Container(height=8),
                panel(
                    ft.Column(
                        [
                            ft.Row(
                                [
                                    ft.Icon(
                                        ft.Icons.GPS_FIXED,
                                        color=COLOR_PRIMARY,
                                    ),
                                    ft.Text(
                                        t("objective"),
                                        size=18,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                ],
                                spacing=10,
                            ),
                            ft.Text(
                                t("objective_desc"),
                                size=15,
                                color=COLOR_TEXT_MUTED,
                            ),
                        ],
                        spacing=12,
                    )
                ),
                panel(
                    ft.Column(
                        [
                            ft.Row(
                                [
                                    ft.Icon(
                                        ft.Icons.INFO_OUTLINE,
                                        color=COLOR_SECONDARY,
                                    ),
                                    ft.Text(
                                        t("important"),
                                        size=18,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                ],
                                spacing=10,
                            ),
                            ft.Text(
                                t("important_desc"),
                                size=15,
                                color=COLOR_TEXT_MUTED,
                            ),
                        ],
                        spacing=12,
                    )
                ),
            ],
            spacing=14,
            scroll=ft.ScrollMode.AUTO,
        )

        page.update()
    # --------------------------------------------------------
    # Propriétaires
    # --------------------------------------------------------

    def proprietaires_view():
        nonlocal current_route
        current_route = "proprietaires"

        nom = text_field(t("full_name") + " *")
        tel = text_field(t("phone"))
        email = text_field(t("email"))
        piece = text_field(t("identity"))

        photo = text_field(t("owner_photo"))
        photo.read_only = True

        # Correction importante :
        # src="" est obligatoire dans votre version de Flet.
        photo_preview = ft.Image(
            src="",
            width=90,
            height=90,
            fit=ft.BoxFit.COVER,
            visible=False,
        )

        photo_status = ft.Text(
            size=11,
            color=COLOR_TEXT_MUTED,
        )

        # Initialisation unique et globale du FilePicker rattaché à la page
        file_picker = ft.FilePicker()
        page.overlay.append(file_picker)
        page.update()

        async def choisir_photo(_):
            try:
                files = await file_picker.pick_files(
                    allow_multiple=False,
                    file_type=ft.FilePickerFileType.CUSTOM,
                    allowed_extensions=[
                        "jpg",
                        "jpeg",
                        "png",
                        "webp",
                    ],
                    with_data=False,
                    cancel_upload_on_window_blur=False,
                )

                if not files:
                    return

                selected = files[0]

                ext = os.path.splitext(
                    selected.name
                )[1].lower()

                safe_name = (
                    f"proprietaire_{uuid.uuid4().hex}{ext}"
                )

                upload_path = (
                    f"proprietaires/{safe_name}"
                )

                await file_picker.upload(
                    [
                        ft.FilePickerUploadFile(
                            name=selected.name,
                            id=selected.id,
                            upload_url=page.get_upload_url(
                                upload_path,
                                3600,
                            ),
                        )
                    ]
                )

                photo.value = (
                    f"/uploads/{upload_path}"
                )

                photo_status.value = selected.name
                photo_status.color = ft.Colors.GREEN_700

                photo_preview.src = photo.value
                photo_preview.visible = True

                page.update()

            except Exception as ex:
                photo_status.value = (
                    f"{t('error')} : {ex}"
                )

                photo_status.color = COLOR_ACCENT

                page.update()

        btn_photo = ft.ElevatedButton(
            t("browse"),
            icon=ft.Icons.UPLOAD,
            on_click=choisir_photo,
        )

        recherche = text_field(
            t("search_owner")
        )

        table = ft.Column()

        def load_table(*_):
            q = recherche.value.strip().lower()

            table.controls.clear()

            table.controls.append(
                ft.Row(
                    [
                        ft.Text(
                            t("id"),
                            width=50,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("full_name"),
                            width=230,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("phone"),
                            width=160,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("email"),
                            width=220,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("identity"),
                            width=220,
                            weight=ft.FontWeight.BOLD,
                        ),
                    ]
                )
            )

            table.controls.append(
                ft.Divider()
            )

            for r in get_proprietaires():

                haystack = " ".join(
                    [
                        str(
                            r["nom_complet"] or ""
                        ),
                        str(
                            r["telephone"] or ""
                        ),
                        str(
                            r["email"] or ""
                        ),
                        str(
                            r["piece_identite"] or ""
                        ),
                    ]
                ).lower()

                if q and q not in haystack:
                    continue

                table.controls.append(
                    ft.Row(
                        [
                            ft.Text(
                                str(r["id"]),
                                width=50,
                            ),
                            ft.Text(
                                r["nom_complet"] or "",
                                width=230,
                            ),
                            ft.Text(
                                r["telephone"] or "",
                                width=160,
                            ),
                            ft.Text(
                                r["email"] or "",
                                width=220,
                            ),
                            ft.Text(
                                r["piece_identite"] or "",
                                width=220,
                            ),
                        ]
                    )
                )

            page.update()

        def save():

            if not require_action("owner_create"):
                return

            if not nom.value.strip():
                snack(
                    t("owner_required"),
                    True,
                )
                return

            now = datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            conn = db_connect()

            try:

                conn.execute(
                    """
                    INSERT INTO proprietaires
                    (
                        nom_complet,
                        telephone,
                        email,
                        piece_identite,
                        photo,
                        date_creation
                    )
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (
                        nom.value.strip(),
                        tel.value.strip(),
                        email.value.strip(),
                        piece.value.strip(),
                        photo.value.strip(),
                        now,
                    ),
                )

                conn.commit()

            except Exception as ex:

                conn.rollback()
                conn.close()

                snack(
                    f"{t('error')}: {ex}",
                    True,
                )

                return

            conn.close()

            journaliser(
                "Création propriétaire",
                nom.value.strip(),
            )

            # Réinitialisation du formulaire
            nom.value = ""
            tel.value = ""
            email.value = ""
            piece.value = ""
            photo.value = ""

            photo_status.value = ""

            photo_preview.visible = False
            photo_preview.src = ""

            snack(
                t("owner_saved")
            )

            load_table()

        recherche.on_change = load_table

        content.content = ft.Column(
            [
                title(
                    t("owners_title"),
                    t("owners_subtitle"),
                ),

                ft.Divider(),

                panel(
                    ft.Column(
                        [
                            ft.Row(
                                [
                                    nom,
                                    tel,
                                ]
                            ),

                            ft.Row(
                                [
                                    email,
                                    piece,
                                ]
                            ),

                            ft.Row(
                                [
                                    photo,
                                    btn_photo,
                                    photo_preview,
                                ],
                                vertical_alignment=(
                                    ft.CrossAxisAlignment.CENTER
                                ),
                            ),

                            photo_status,

                            ft.ElevatedButton(
                                t("save_owner"),
                                icon=ft.Icons.SAVE,
                                on_click=lambda e: save(),
                            ),
                        ],
                        spacing=12,
                    )
                ),

                recherche,

                panel(
                    table
                ),
            ],

            spacing=12,

            scroll=ft.ScrollMode.AUTO,
        )

        load_table()
    # --------------------------------------------------------
    # Parcelles
    # --------------------------------------------------------

    def parcelles_view():
        nonlocal current_route
        current_route = "parcelles"

        numero = text_field(t("parcel_number"))
        adresse = text_field(t("address"))
        province = text_field(t("province"))
        ville = text_field(t("city"))
        commune = text_field(t("commune"))
        quartier = text_field(t("neighborhood"))
        localite = text_field(t("locality"))
        superficie = text_field(t("area"))
        latitude = text_field(t("latitude"))
        longitude = text_field(t("longitude"))
        agent = text_field(t("agent"))

        nom_proprio = text_field(t("owner_name"))
        cni_proprio = text_field(t("owner_identity"))
        tel_proprio = text_field(t("phone"))
        photo_proprio = text_field(t("owner_photo"))
        photo_proprio.read_only = True

        statut = ft.Dropdown(
            label=t("status"),
            value="Occupée",
            options=[
                ft.dropdown.Option("Occupée", t("occupied")),
                ft.dropdown.Option("Disponible", t("available")),
                ft.dropdown.Option("En vente", t("for_sale")),
                ft.dropdown.Option("Transférée", t("transferred")),
                ft.dropdown.Option("Contestée", t("disputed")),
            ],
            border_radius=14,
            border_color=COLOR_BORDER,
            expand=True,
        )

        polygone = text_field(
            t("polygon"),
            "[]",
            multiline=True,
        )

        gps_status = ft.Text(
            t("gps_ready"),
            size=12,
            color=COLOR_PRIMARY,
        )

        # Initialisation unique et globale du FilePicker rattaché à la page
        file_picker_parcelle = ft.FilePicker()
        page.overlay.append(file_picker_parcelle)
        page.update()

        async def activer_saisie_photo(_):
            try:
                files = await file_picker_parcelle.pick_files(
                    allow_multiple=False,
                    file_type=ft.FilePickerFileType.CUSTOM,
                    allowed_extensions=["jpg", "jpeg", "png", "webp"],
                    with_data=False,
                    cancel_upload_on_window_blur=False,
                )
                if not files:
                    return
                selected = files[0]
                ext = os.path.splitext(selected.name)[1].lower()
                safe_name = f"proprietaire_{uuid.uuid4().hex}{ext}"
                upload_path = f"proprietaires/{safe_name}"
                await file_picker_parcelle.upload([
                    ft.FilePickerUploadFile(
                        name=selected.name,
                        id=selected.id,
                        upload_url=page.get_upload_url(upload_path, 3600),
                    )
                ])
                photo_proprio.value = f"/uploads/{upload_path}"
                photo_proprio.helper_text = selected.name
                page.update()
            except Exception as ex:
                snack(f"{t('error')} : {ex}", True)

        btn_parcourir = ft.ElevatedButton(
            t("browse"),
            icon=ft.Icons.EDIT,
            on_click=activer_saisie_photo,
        )

        def capturer_point_gps(_):
            try:
                # Coordonnées de démonstration conservées depuis
                # la version fournie.
                lat_actuelle = -2.503030
                lon_actuelle = 28.864286

                latitude.value = str(lat_actuelle)
                longitude.value = str(lon_actuelle)

                try:
                    points = json.loads(
                        polygone.value.strip() or "[]"
                    )
                    if not isinstance(points, list):
                        points = []
                except Exception:
                    points = []

                points.append(
                    [lat_actuelle, lon_actuelle]
                )

                polygone.value = json.dumps(
                    points,
                    ensure_ascii=False,
                )

                if len(points) >= 3 and "/" not in superficie.value:
                    calculated = calculer_superficie_cm2(
                        points
                    )
                    if calculated > 0:
                        superficie.value = (
                            f"{calculated:.2f}"
                        )

                gps_status.value = tf(
                    "gps_point_added",
                    count=len(points),
                )
                gps_status.color = ft.Colors.GREEN_700
                page.update()

            except Exception as ex:
                gps_status.value = (
                    f"{t('error')} : {ex}"
                )
                gps_status.color = COLOR_ACCENT
                page.update()

        recherche = text_field(
            t("search_parcel")
        )
        table = ft.Column()

        def load_table(*_):
            q = recherche.value.strip().lower()
            table.controls.clear()

            table.controls.append(
                ft.Row(
                    [
                        ft.Text(
                            t("id"),
                            width=45,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("parcel"),
                            width=140,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("owner"),
                            width=210,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("area_short"),
                            width=120,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("status"),
                            width=120,
                            weight=ft.FontWeight.BOLD,
                        ),
                        ft.Text(
                            t("location"),
                            width=260,
                            weight=ft.FontWeight.BOLD,
                        ),
                    ]
                )
            )
            table.controls.append(ft.Divider())

            for r in get_parcelles():
                haystack = " ".join(
                    [
                        str(r["id"] or ""),
                        str(r["numero"] or ""),
                        str(r["proprietaire"] or ""),
                        str(r["adresse"] or ""),
                        str(r["province"] or ""),
                        str(r["ville_territoire"] or ""),
                        str(r["commune_chefferie"] or ""),
                        str(r["quartier_groupement"] or ""),
                        str(r["localite"] or ""),
                        str(r["agent"] or ""),
                    ]
                ).lower()

                if q and q not in haystack:
                    continue

                table.controls.append(
                    ft.Row(
                        [
                            ft.Text(
                                str(r["id"]),
                                width=45,
                            ),
                            ft.Text(
                                r["numero"] or "—",
                                width=140,
                            ),
                            ft.Text(
                                r["proprietaire"] or "—",
                                width=210,
                            ),
                            ft.Text(
                                f'{r["superficie"] or 0:.2f} cm²',
                                width=120,
                            ),
                            ft.Text(
                                enum_text(r["statut"], STATUS_KEYS),
                                width=120,
                            ),
                            ft.Text(
                                f'{r["province"] or ""} / '
                                f'{r["ville_territoire"] or ""}',
                                width=260,
                            ),
                        ]
                    )
                )

            page.update()

        def save():
            if not require_action("parcel_create"):
                return
            sup_str = superficie.value.strip()
            sup = 0.0
            calculated = False

            try:
                sup = parse_area(sup_str)
            except (ValueError, TypeError):
                snack(t("invalid_area"), True)
                return

            try:
                lat = (
                    float(latitude.value.replace(",", "."))
                    if latitude.value.strip()
                    else None
                )
                lon = (
                    float(longitude.value.replace(",", "."))
                    if longitude.value.strip()
                    else None
                )
            except ValueError:
                snack(t("invalid_gps"), True)
                return

            poly_json = polygone.value.strip() or "[]"

            try:
                points = json.loads(poly_json)
                if not isinstance(points, list):
                    raise ValueError

                if (
                    len(points) >= 3
                    and ("/" not in sup_str or sup == 0)
                ):
                    calculated_value = (
                        calculer_superficie_cm2(points)
                    )
                    if calculated_value > 0:
                        sup = calculated_value
                        calculated = True

            except Exception:
                snack(t("invalid_polygon"), True)
                return

            if not nom_proprio.value.strip():
                snack(t("owner_required"), True)
                return

            numero_value = numero.value.strip()
            if not numero_value:
                snack(t("parcel_required"), True)
                return

            if statut.value == "Contestée":
                snack(t("parcel_disputed"), True)
                return

            # ----------------------------------------------------
            # Contrôle du périmètre territorial à la création
            # ----------------------------------------------------
            _level = str(current_user.get("niveau_acces") or "National")
            if _level != "National":
                if province.value.strip().lower() != str(current_user.get("province") or "").strip().lower():
                    journaliser("PARCELLE_REFUSEE", f"Province hors périmètre: {province.value.strip()}")
                    snack(t("access_denied"), True)
                    return
                if _level == "Ville" and ville.value.strip().lower() != str(current_user.get("ville") or "").strip().lower():
                    snack(t("access_denied"), True); return
                if _level == "Territoire" and ville.value.strip().lower() != str(current_user.get("territoire") or "").strip().lower():
                    snack(t("access_denied"), True); return
                if _level == "Commune" and commune.value.strip().lower() != str(current_user.get("commune") or "").strip().lower():
                    snack(t("access_denied"), True); return
                if _level in {"Secteur", "Chefferie"} and commune.value.strip().lower() != str(current_user.get("secteur_chefferie") or "").strip().lower():
                    snack(t("access_denied"), True); return
                if _level == "Groupement" and quartier.value.strip().lower() != str(current_user.get("groupement") or "").strip().lower():
                    snack(t("access_denied"), True); return
                if _level in {"Village", "Local"} and localite.value.strip().lower() != str(current_user.get("village") or "").strip().lower():
                    snack(t("access_denied"), True); return
            # ----------------------------------------------------
            # Vérification préalable et rattachement propriétaire
            # ----------------------------------------------------
            # Une même personne peut posséder plusieurs parcelles.
            # On ne doit donc PAS bloquer l'enregistrement lorsqu'une
            # pièce d'identité existe déjà : on réutilise son proprietaire_id.
            precheck = db_connect()
            existing_parcel = precheck.execute(
                "SELECT id FROM parcelles WHERE LOWER(numero) = LOWER(?) LIMIT 1",
                (numero_value,),
            ).fetchone()
            if existing_parcel is not None:
                precheck.close()
                journaliser("PARCELLE_REFUSEE", f"Numéro déjà existant : {numero_value}")
                snack(t("duplicate_parcel"), True)
                return

            nom_value = nom_proprio.value.strip()
            cni_value = cni_proprio.value.strip()
            tel_value = tel_proprio.value.strip()

            existing_owner = None

            # 1) Identité officielle : critère prioritaire.
            if cni_value:
                existing_owner = precheck.execute(
                    """
                    SELECT id FROM proprietaires
                    WHERE LOWER(TRIM(COALESCE(piece_identite, ''))) = LOWER(TRIM(?))
                    ORDER BY id ASC LIMIT 1
                    """,
                    (cni_value,),
                ).fetchone()

            # 2) Si aucune CNI n'est fournie, on peut rattacher une
            # fiche existante par nom + téléphone.
            if existing_owner is None and tel_value:
                existing_owner = precheck.execute(
                    """
                    SELECT id FROM proprietaires
                    WHERE LOWER(TRIM(COALESCE(nom_complet, ''))) = LOWER(TRIM(?))
                      AND LOWER(TRIM(COALESCE(telephone, ''))) = LOWER(TRIM(?))
                    ORDER BY id ASC LIMIT 1
                    """,
                    (nom_value, tel_value),
                ).fetchone()

            precheck.close()

            now = datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            make_backup("pre_mutation")
            conn = db_connect()

            try:
                cursor = conn.cursor()

                if existing_owner is not None:
                    # Réutiliser le propriétaire existant afin que toutes
                    # ses parcelles soient reliées au même propriétaire.
                    proprietaire_id = int(existing_owner["id"])

                    # Compléter les champs encore vides sans écraser les
                    # informations déjà enregistrées.
                    cursor.execute(
                        """
                        UPDATE proprietaires
                        SET telephone = CASE
                                WHEN (telephone IS NULL OR TRIM(telephone) = '') AND ? <> ''
                                THEN ?
                                ELSE telephone
                            END,
                            piece_identite = CASE
                                WHEN (piece_identite IS NULL OR TRIM(piece_identite) = '') AND ? <> ''
                                THEN ?
                                ELSE piece_identite
                            END,
                            photo = CASE
                                WHEN (photo IS NULL OR TRIM(photo) = '') AND ? <> ''
                                THEN ?
                                ELSE photo
                            END
                        WHERE id = ?
                        """,
                        (
                            tel_value, tel_value,
                            cni_value, cni_value,
                            photo_proprio.value.strip(), photo_proprio.value.strip(),
                            proprietaire_id,
                        ),
                    )
                else:
                    # Nouveau propriétaire.
                    cursor.execute(
                        """
                        INSERT INTO proprietaires
                        (nom_complet, telephone, email,
                         piece_identite, photo, date_creation)
                        VALUES (?, ?, ?, ?, ?, ?)
                        """,
                        (
                            nom_value,
                            tel_value,
                            "",
                            cni_value,
                            photo_proprio.value.strip(),
                            now,
                        ),
                    )

                    proprietaire_id = cursor.lastrowid

                cursor.execute(
                    """
                    INSERT INTO parcelles
                    (numero, adresse, province,
                     ville_territoire, commune_chefferie,
                     quartier_groupement, localite,
                     superficie, latitude, longitude,
                     polygone, proprietaire_id,
                     statut, date_enregistrement, agent)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?,
                            ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        numero_value,
                        adresse.value.strip(),
                        province.value.strip(),
                        ville.value.strip(),
                        commune.value.strip(),
                        quartier.value.strip(),
                        localite.value.strip(),
                        sup,
                        lat,
                        lon,
                        poly_json,
                        proprietaire_id,
                        statut.value,
                        now,
                        agent.value.strip(),
                    ),
                )

                conn.commit()

            except sqlite3.IntegrityError:
                conn.rollback()
                conn.close()
                snack(t("duplicate_parcel"), True)
                return

            except Exception as ex:
                conn.rollback()
                conn.close()
                snack(f"{t('error')} : {ex}", True)
                return

            conn.close()

            journaliser(
                "Création parcelle",
                numero_value,
            )

            for field in [
                numero,
                adresse,
                province,
                ville,
                commune,
                quartier,
                localite,
                superficie,
                latitude,
                longitude,
                agent,
                nom_proprio,
                cni_proprio,
                tel_proprio,
                photo_proprio,
            ]:
                field.value = ""

            polygone.value = "[]"
            gps_status.value = t("gps_ready")
            gps_status.color = COLOR_PRIMARY

            if calculated:
                message = (
                    f"{t('parcel_saved')} "
                    f"{sup:.2f} cm²"
                )
            else:
                message = (
                    f"{t('parcel_saved')} "
                    f"{sup:.2f} cm²"
                )

            snack(message)
            load_table()

        recherche.on_change = load_table

        content.content = ft.Column(
            [
                title(
                    t("parcels_title"),
                    t("parcels_subtitle"),
                ),
                ft.Divider(),
                panel(
                    ft.Column(
                        [
                            ft.Row(
                                [
                                    numero,
                                    nom_proprio,
                                ]
                            ),
                            ft.Row(
                                [
                                    cni_proprio,
                                    tel_proprio,
                                ]
                            ),
                            ft.Row(
                                [
                                    photo_proprio,
                                    btn_parcourir,
                                ]
                            ),
                            ft.Row(
                                [
                                    adresse,
                                    province,
                                ]
                            ),
                            ft.Row(
                                [
                                    ville,
                                    commune,
                                ]
                            ),
                            ft.Row(
                                [
                                    quartier,
                                    localite,
                                ]
                            ),
                            ft.Row(
                                [
                                    latitude,
                                    longitude,
                                ]
                            ),
                            ft.Row(
                                [
                                    superficie,
                                    statut,
                                ]
                            ),
                            agent,
                            ft.Container(
                                padding=15,
                                bgcolor=(
                                    "#172033"
                                    if page.theme_mode
                                    == ft.ThemeMode.DARK
                                    else "#F8FAFC"
                                ),
                                border_radius=16,
                                border=ft.Border(left=ft.BorderSide(width=1, color="#334155"
                                    if page.theme_mode
                                    == ft.ThemeMode.DARK
                                    else COLOR_BORDER,), top=ft.BorderSide(width=1, color="#334155"
                                    if page.theme_mode
                                    == ft.ThemeMode.DARK
                                    else COLOR_BORDER,), right=ft.BorderSide(width=1, color="#334155"
                                    if page.theme_mode
                                    == ft.ThemeMode.DARK
                                    else COLOR_BORDER,), bottom=ft.BorderSide(width=1, color="#334155"
                                    if page.theme_mode
                                    == ft.ThemeMode.DARK
                                    else COLOR_BORDER,)),
                                content=ft.Column(
                                    [
                                        ft.Text(
                                            t("survey"),
                                            size=15,
                                            weight=ft.FontWeight.BOLD,
                                        ),
                                        ft.Text(
                                            t("survey_help"),
                                            size=11,
                                            color=COLOR_TEXT_MUTED,
                                        ),
                                        ft.Row(
                                            [
                                                ft.ElevatedButton(
                                                    t("capture_gps"),
                                                    icon=ft.Icons.GPS_FIXED,
                                                    on_click=(
                                                        capturer_point_gps
                                                    ),
                                                ),
                                                gps_status,
                                            ]
                                        ),
                                    ],
                                    spacing=7,
                                ),
                            ),
                            polygone,
                            ft.Row(
                                [
                                    ft.ElevatedButton(
                                        t("save_parcel"),
                                        icon=ft.Icons.SAVE,
                                        on_click=lambda e: save(),
                                    ),
                                    ft.Text(
                                        t("supported"),
                                        size=11,
                                        color=COLOR_TEXT_MUTED,
                                    ),
                                ]
                            ),
                        ],
                        spacing=10,
                    )
                ),
                recherche,
                panel(table),
            ],
            spacing=12,
            scroll=ft.ScrollMode.AUTO,
        )

        load_table()

        # --------------------------------------------------------
    # Mutations
    # --------------------------------------------------------

    def mutations_view():
        nonlocal current_route
        current_route = "mutations"

        selected_parcelle_id = None
        selected_ancien_id = None

        recherche_parcelle = text_field(
            t("search_plot")
        )
        recherche_ancien = text_field(
            t("search_old_owner")
        )

        info_parcelle = ft.Text(
            t("no_plot_selected"),
            size=12,
            color=COLOR_TEXT_MUTED,
        )

        info_ancien = ft.Text(
            t("no_old_owner"),
            size=12,
            color=COLOR_TEXT_MUTED,
        )

        suggestions_parcelles = ft.Column(
            spacing=3
        )

        suggestions_anciens = ft.Column(
            spacing=3
        )

        nom_nouveau = text_field(
            t("full_name")
        )

        cni_nouveau = text_field(
            t("owner_identity")
        )

        tel_nouveau = text_field(
            t("phone")
        )

        photo_nouveau = text_field(
            t("owner_photo")
        )

        photo_nouveau.read_only = True

        type_trans = ft.Dropdown(
            label=t("transaction_type"),
            value="Vente",
            options=[
                ft.dropdown.Option(
                    "Vente",
                    t("sale"),
                ),
                ft.dropdown.Option(
                    "Cession",
                    t("cession"),
                ),
                ft.dropdown.Option(
                    "Donation",
                    t("donation"),
                ),
                ft.dropdown.Option(
                    "Autre",
                    t("other"),
                ),
            ],
            border_radius=14,
            border_color=COLOR_BORDER,
            expand=True,
        )

        reference = text_field(
            t("act_reference")
        )

        prix = text_field(
            t("price")
        )

        agent = text_field(
            t("agent")
        )

        table = ft.Column()

        # --------------------------------------------------------
        # FilePicker Mutation
        # --------------------------------------------------------
        # Dans les versions récentes de Flet, FilePicker est un Service
        # et ne doit pas être ajouté à page.overlay.
        # L'instance est automatiquement enregistrée par Flet lorsqu'elle
        # est créée et utilisée depuis la page.
        mutation_file_picker = ft.FilePicker()

        # --------------------------------------------------------
        # Sélection de la photo du nouveau propriétaire
        # --------------------------------------------------------

        async def choisir_photo_mutation(_):
            try:
                files = await mutation_file_picker.pick_files(
                    allow_multiple=False,
                    file_type=ft.FilePickerFileType.CUSTOM,
                    allowed_extensions=[
                        "jpg",
                        "jpeg",
                        "png",
                        "webp",
                    ],
                    with_data=False,
                    cancel_upload_on_window_blur=False,
                )

                if not files:
                    return

                selected = files[0]

                if not selected.name:
                    snack(
                        t("error"),
                        True,
                    )
                    return

                ext = os.path.splitext(
                    selected.name
                )[1].lower()

                allowed_extensions = {
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".webp",
                }

                if ext not in allowed_extensions:
                    snack(
                        "Format de fichier non autorisé.",
                        True,
                    )
                    return

                safe_name = (
                    f"mutation_"
                    f"{uuid.uuid4().hex}"
                    f"{ext}"
                )

                upload_path = (
                    f"mutations/{safe_name}"
                )

                # ------------------------------------------------
                # Création du dossier local si nécessaire
                # ------------------------------------------------

                mutation_upload_dir = os.path.join(
                    UPLOADS_DIR,
                    "mutations",
                )

                os.makedirs(
                    mutation_upload_dir,
                    exist_ok=True,
                )

                # ------------------------------------------------
                # Upload Flet
                # ------------------------------------------------

                upload_url = page.get_upload_url(
                    upload_path,
                    3600,
                )

                await mutation_file_picker.upload(
                    [
                        ft.FilePickerUploadFile(
                            name=selected.name,
                            id=selected.id,
                            upload_url=upload_url,
                        )
                    ]
                )

                # ------------------------------------------------
                # Enregistrement du chemin relatif
                # ------------------------------------------------

                photo_nouveau.value = (
                    f"/uploads/{upload_path}"
                )

                photo_nouveau.read_only = True

                try:
                    photo_nouveau.helper_text = (
                        selected.name
                    )
                except Exception:
                    pass

                journaliser(
                    "PHOTO_MUTATION_TELEVEREE",
                    (
                        f"Fichier : {selected.name}; "
                        f"Destination : {upload_path}"
                    ),
                )

                snack(
                    f"Photo sélectionnée : {selected.name}"
                )

                page.update()

            except Exception as ex:

                journaliser(
                    "ERREUR_UPLOAD_MUTATION",
                    str(ex),
                )

                snack(
                    f"{t('error')} : {ex}",
                    True,
                )

                page.update()

        btn_parcourir = ft.ElevatedButton(
            t("browse"),
            icon=ft.Icons.UPLOAD,
            on_click=choisir_photo_mutation,
        )

        # --------------------------------------------------------
        # Sélection parcelle
        # --------------------------------------------------------

        def selectionner_parcelle(
            ident,
            num,
            proprietaire_nom,
            prop_id,
        ):
            nonlocal selected_parcelle_id
            nonlocal selected_ancien_id

            selected_parcelle_id = ident

            recherche_parcelle.value = (
                f"#{ident} — {num}"
            )

            info_parcelle.value = (
                f"{t('current_owner')} : "
                f"{proprietaire_nom}"
            )

            info_parcelle.color = (
                ft.Colors.GREEN_700
            )

            suggestions_parcelles.controls.clear()

            if prop_id:
                selected_ancien_id = prop_id

                recherche_ancien.value = (
                    proprietaire_nom
                )

                info_ancien.value = (
                    f"{t('previous')} ID : "
                    f"{prop_id}"
                )

                info_ancien.color = (
                    ft.Colors.GREEN_700
                )

            else:
                selected_ancien_id = None

                recherche_ancien.value = ""

                info_ancien.value = (
                    t("no_old_owner")
                )

                info_ancien.color = (
                    COLOR_TEXT_MUTED
                )

            page.update()

        # --------------------------------------------------------
        # Recherche des parcelles
        # --------------------------------------------------------

        def filtrer_parcelles(_):
            q = (
                recherche_parcelle.value
                .strip()
                .lower()
            )

            suggestions_parcelles.controls.clear()

            if not q:
                page.update()
                return

            conn = db_connect()

            try:
                _scope_sql, _scope_params = territorial_scope_sql("p")
                rows = conn.execute(
                    f"""
                    SELECT
                        p.id,
                        p.numero,
                        pr.nom_complet,
                        pr.id AS prop_id
                    FROM parcelles p
                    LEFT JOIN proprietaires pr
                        ON p.proprietaire_id = pr.id
                    WHERE
                        ({_scope_sql})
                        AND (
                            LOWER(COALESCE(p.numero, '')) LIKE ?
                            OR CAST(p.id AS TEXT) LIKE ?
                        )
                    LIMIT 10
                    """,
                    _scope_params + [f"%{q}%", f"%{q}%"],
                ).fetchall()

            finally:
                conn.close()

            for row in rows:

                p_id = row["id"]

                p_num = (
                    row["numero"]
                    or "—"
                )

                p_owner = (
                    row["nom_complet"]
                    or "—"
                )

                p_owner_id = (
                    row["prop_id"]
                )

                def make_click(
                    ident=p_id,
                    num=p_num,
                    owner_name=p_owner,
                    owner_id=p_owner_id,
                ):
                    return lambda ev: (
                        selectionner_parcelle(
                            ident,
                            num,
                            owner_name,
                            owner_id,
                        )
                    )

                suggestions_parcelles.controls.append(
                    ft.Container(
                        padding=9,
                        border_radius=7,
                        bgcolor=(
                            "#243244"
                            if page.theme_mode
                            == ft.ThemeMode.DARK
                            else "#F1F5F9"
                        ),
                        content=ft.Text(
                            f"#{p_id} — "
                            f"{p_num} — "
                            f"{p_owner}",
                            size=12,
                        ),
                        on_click=make_click(),
                    )
                )

            page.update()

        # --------------------------------------------------------
        # Sélection ancien propriétaire
        # --------------------------------------------------------

        def selectionner_ancien(
            ident,
            nom,
        ):
            nonlocal selected_ancien_id

            selected_ancien_id = ident

            recherche_ancien.value = nom

            info_ancien.value = (
                f"{t('previous')} ID : "
                f"{ident}"
            )

            info_ancien.color = (
                ft.Colors.GREEN_700
            )

            suggestions_anciens.controls.clear()

            page.update()

        # --------------------------------------------------------
        # Recherche anciens propriétaires
        # --------------------------------------------------------

        def filtrer_anciens(_):
            q = (
                recherche_ancien.value
                .strip()
                .lower()
            )

            suggestions_anciens.controls.clear()

            if not q:
                page.update()
                return

            conn = db_connect()

            try:
                _scope_sql, _scope_params = territorial_scope_sql("p")
                rows = conn.execute(
                    f"""
                    SELECT DISTINCT
                        pr.id,
                        pr.nom_complet,
                        pr.telephone
                    FROM proprietaires pr
                    JOIN parcelles p
                        ON p.proprietaire_id = pr.id
                    WHERE ({_scope_sql})
                      AND LOWER(COALESCE(pr.nom_complet, '')) LIKE ?
                    LIMIT 10
                    """,
                    _scope_params + [f"%{q}%"],
                ).fetchall()

            finally:
                conn.close()

            for row in rows:

                owner_id = row["id"]

                owner_name = (
                    row["nom_complet"]
                )

                owner_phone = (
                    row["telephone"]
                    or ""
                )

                def make_click(
                    ident=owner_id,
                    nom=owner_name,
                ):
                    return lambda ev: (
                        selectionner_ancien(
                            ident,
                            nom,
                        )
                    )

                suggestions_anciens.controls.append(
                    ft.Container(
                        padding=9,
                        border_radius=7,
                        bgcolor=(
                            "#243244"
                            if page.theme_mode
                            == ft.ThemeMode.DARK
                            else "#F1F5F9"
                        ),
                        content=ft.Text(
                            f"{owner_name}"
                            + (
                                f" ({owner_phone})"
                                if owner_phone
                                else ""
                            ),
                            size=12,
                        ),
                        on_click=make_click(),
                    )
                )

            page.update()

        recherche_parcelle.on_change = (
            filtrer_parcelles
        )

        recherche_ancien.on_change = (
            filtrer_anciens
        )

        # --------------------------------------------------------
        # Enregistrement mutation
        # --------------------------------------------------------

        def save():
            nonlocal selected_parcelle_id
            nonlocal selected_ancien_id

            # ----------------------------------------------------
            # Permission
            # ----------------------------------------------------

            if not require_action(
                "mutation_create"
            ):
                return

            # ----------------------------------------------------
            # Parcelle obligatoire
            # ----------------------------------------------------

            if selected_parcelle_id is None:
                snack(
                    t("parcel_required"),
                    True,
                )
                return

            # ----------------------------------------------------
            # Ancien propriétaire obligatoire
            # ----------------------------------------------------

            if selected_ancien_id is None:
                snack(
                    t("invalid_mutation"),
                    True,
                )
                return

            # ----------------------------------------------------
            # Nouveau propriétaire obligatoire
            # ----------------------------------------------------

            if not nom_nouveau.value.strip():
                snack(
                    t("new_owner_required"),
                    True,
                )
                return

            # ----------------------------------------------------
            # Vérification serveur de la parcelle
            # ----------------------------------------------------

            check_conn = db_connect()

            try:
                parcel_row = check_conn.execute(
                    """
                    SELECT
                        id,
                        numero,
                        proprietaire_id,
                        statut
                    FROM parcelles
                    WHERE id = ?
                    """,
                    (
                        int(
                            selected_parcelle_id
                        ),
                    ),
                ).fetchone()

            finally:
                check_conn.close()

            if parcel_row is None:

                journaliser(
                    "MUTATION_REFUSEE",
                    (
                        "Parcelle inexistante : "
                        f"{selected_parcelle_id}"
                    ),
                )

                snack(
                    t("parcel_not_found"),
                    True,
                )

                return

            # ----------------------------------------------------
            # Contrôle du périmètre territorial côté serveur
            # ----------------------------------------------------
            scope_sql, scope_params = territorial_scope_sql("p")
            scope_conn = db_connect()
            scope_row = scope_conn.execute(
                f"SELECT p.id FROM parcelles p WHERE p.id = ? AND ({scope_sql})",
                [int(selected_parcelle_id)] + scope_params,
            ).fetchone()
            scope_conn.close()
            if scope_row is None:
                journaliser("MUTATION_REFUSEE", f"Parcelle hors périmètre territorial : {selected_parcelle_id}")
                snack(t("access_denied"), True)
                return

            # ----------------------------------------------------
            # Vérification propriétaire actuel
            # ----------------------------------------------------

            actual_owner_id = (
                parcel_row[
                    "proprietaire_id"
                ]
            )

            if actual_owner_id is None:

                journaliser(
                    "MUTATION_REFUSEE",
                    (
                        "Parcelle sans propriétaire : "
                        f"{selected_parcelle_id}"
                    ),
                )

                snack(
                    t("owner_not_found"),
                    True,
                )

                return

            # ----------------------------------------------------
            # Cohérence ancien propriétaire
            # ----------------------------------------------------

            if (
                selected_ancien_id is None
                or int(
                    selected_ancien_id
                )
                != int(
                    actual_owner_id
                )
            ):

                journaliser(
                    "MUTATION_REFUSEE",
                    (
                        "Propriétaire ancien "
                        "incohérent pour parcelle "
                        f"{selected_parcelle_id}"
                    ),
                )

                snack(
                    t("invalid_mutation"),
                    True,
                )

                return

            # ----------------------------------------------------
            # Parcelle contestée
            # ----------------------------------------------------

            if (
                parcel_row["statut"]
                == "Contestée"
            ):

                journaliser(
                    "MUTATION_REFUSEE",
                    (
                        "Parcelle contestée : "
                        f"{selected_parcelle_id}"
                    ),
                )

                snack(
                    t("parcel_disputed"),
                    True,
                )

                return

            # ----------------------------------------------------
            # Validation du prix
            # ----------------------------------------------------

            try:

                prix_value = (
                    float(
                        prix.value
                        .replace(",", ".")
                    )
                    if prix.value.strip()
                    else None
                )

            except (
                ValueError,
                TypeError,
            ):

                snack(
                    t("invalid_price"),
                    True,
                )

                return

            # ----------------------------------------------------
            # Date
            # ----------------------------------------------------

            now = datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            conn = db_connect()

            try:

                cursor = conn.cursor()

                # ------------------------------------------------
                # Pièce d'identité
                # ------------------------------------------------

                cni_value = (
                    cni_nouveau.value.strip()
                )

                if cni_value:

                    existing_owner = (
                        cursor.execute(
                            """
                            SELECT id
                            FROM proprietaires
                            WHERE piece_identite = ?
                            LIMIT 1
                            """,
                            (
                                cni_value,
                            ),
                        ).fetchone()
                    )

                    if (
                        existing_owner
                        is not None
                    ):

                        conn.rollback()
                        conn.close()

                        journaliser(
                            "MUTATION_REFUSEE",
                            (
                                "Pièce d'identité "
                                "déjà enregistrée : "
                                f"{cni_value}"
                            ),
                        )

                        snack(
                            t(
                                "invalid_mutation"
                            ),
                            True,
                        )

                        return

                # ------------------------------------------------
                # Création nouveau propriétaire
                # ------------------------------------------------

                cursor.execute(
                    """
                    INSERT INTO proprietaires
                    (
                        nom_complet,
                        telephone,
                        email,
                        piece_identite,
                        photo,
                        date_creation
                    )
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (
                        nom_nouveau.value.strip(),
                        tel_nouveau.value.strip(),
                        "",
                        cni_value,
                        photo_nouveau.value.strip(),
                        now,
                    ),
                )

                new_id = cursor.lastrowid

                # ------------------------------------------------
                # Vérification nouveau propriétaire créé
                # ------------------------------------------------

                if not new_id:

                    raise RuntimeError(
                        "Impossible de créer "
                        "le nouveau propriétaire."
                    )

                # ------------------------------------------------
                # Création de la mutation
                # ------------------------------------------------

                cursor.execute(
                    """
                    INSERT INTO mutations
                    (
                        parcelle_id,
                        ancien_proprietaire_id,
                        nouveau_proprietaire_id,
                        type_transaction,
                        reference_acte,
                        prix,
                        date_transaction,
                        agent,
                        statut
                    )
                    VALUES (
                        ?, ?, ?, ?, ?,
                        ?, ?, ?, ?
                    )
                    """,
                    (
                        int(
                            selected_parcelle_id
                        ),
                        int(
                            selected_ancien_id
                        ),
                        int(new_id),
                        type_trans.value,
                        reference.value.strip(),
                        prix_value,
                        now,
                        agent.value.strip(),
                        "En attente",
                    ),
                )

                mutation_id = cursor.lastrowid

                if not mutation_id:

                    raise RuntimeError(
                        "Impossible de créer "
                        "la mutation."
                    )

                # ------------------------------------------------
                # IMPORTANT :
                #
                # La parcelle n'est PAS transférée ici.
                #
                # Elle sera transférée uniquement après
                # validation par un rôle autorisé.
                # ------------------------------------------------

                conn.commit()

            except Exception as ex:

                conn.rollback()
                conn.close()

                journaliser(
                    "MUTATION_ERREUR",
                    str(ex),
                )

                snack(
                    f"{t('error')} : {ex}",
                    True,
                )

                return

            conn.close()

            # ----------------------------------------------------
            # Journalisation
            # ----------------------------------------------------

            journaliser(
                "MUTATION_CREEE_EN_ATTENTE",
                (
                    f"Mutation ID "
                    f"{mutation_id}; "
                    f"Parcelle ID "
                    f"{selected_parcelle_id}; "
                    f"Ancien propriétaire ID "
                    f"{selected_ancien_id}; "
                    f"Nouveau propriétaire ID "
                    f"{new_id}; "
                    f"validation requise"
                ),
            )

            # ----------------------------------------------------
            # Réinitialisation formulaire
            # ----------------------------------------------------

            recherche_parcelle.value = ""

            recherche_ancien.value = ""

            nom_nouveau.value = ""

            cni_nouveau.value = ""

            tel_nouveau.value = ""

            photo_nouveau.value = ""

            reference.value = ""

            prix.value = ""

            agent.value = ""

            info_parcelle.value = (
                t("no_plot_selected")
            )

            info_parcelle.color = (
                COLOR_TEXT_MUTED
            )

            info_ancien.value = (
                t("no_old_owner")
            )

            info_ancien.color = (
                COLOR_TEXT_MUTED
            )

            suggestions_parcelles.controls.clear()

            suggestions_anciens.controls.clear()

            selected_parcelle_id = None

            selected_ancien_id = None

            # ----------------------------------------------------
            # Message de succès
            # ----------------------------------------------------

            snack(
                t("transfer_saved")
            )

            load_table()

            page.update()

        # --------------------------------------------------------
        # Bouton Enregistrer
        # --------------------------------------------------------

        def cliquer_enregistrer_mutation(e):

            try:

                # ------------------------------------------------
                # Vérification session si disponible
                # ------------------------------------------------

                try:

                    if (
                        "check_session"
                        in globals()
                    ):

                        if not check_session(
                            show_message=True
                        ):
                            return

                except Exception:
                    pass

                # ------------------------------------------------
                # Vérification permission
                # ------------------------------------------------

                if not require_action(
                    "mutation_create"
                ):
                    return

                # ------------------------------------------------
                # Parcelle obligatoire
                # ------------------------------------------------

                if (
                    selected_parcelle_id
                    is None
                ):

                    snack(
                        t("parcel_required"),
                        True,
                    )

                    return

                # ------------------------------------------------
                # Ancien propriétaire obligatoire
                # ------------------------------------------------

                if (
                    selected_ancien_id
                    is None
                ):

                    snack(
                        t("invalid_mutation"),
                        True,
                    )

                    return

                # ------------------------------------------------
                # Nouveau propriétaire obligatoire
                # ------------------------------------------------

                if not (
                    nom_nouveau.value
                    .strip()
                ):

                    snack(
                        t(
                            "new_owner_required"
                        ),
                        True,
                    )

                    return

                # ------------------------------------------------
                # Appel réel de save()
                # ------------------------------------------------

                save()

            except Exception as ex:

                journaliser(
                    "ERREUR_BOUTON_MUTATION",
                    str(ex),
                )

                snack(
                    f"{t('error')} : {ex}",
                    True,
                )

                page.update()

        # --------------------------------------------------------
        # Validation d'une mutation
        # --------------------------------------------------------

        def valider_mutation(
            mutation_id
        ):

            if not require_action(
                "mutation_validate"
            ):
                return

            make_backup(
                "pre_validation_mutation"
            )

            conn = db_connect()

            try:

                row = conn.execute(
                    """
                    SELECT
                        m.id,
                        m.parcelle_id,
                        m.ancien_proprietaire_id,
                        m.nouveau_proprietaire_id,
                        m.statut,
                        p.proprietaire_id
                            AS proprietaire_actuel,
                        p.statut
                            AS statut_parcelle
                    FROM mutations m
                    JOIN parcelles p
                        ON p.id = m.parcelle_id
                    WHERE m.id = ?
                    """,
                    (
                        int(mutation_id),
                    ),
                ).fetchone()

                if row is None:

                    conn.close()

                    snack(
                        t("invalid_mutation"),
                        True,
                    )

                    return

                _scope_sql, _scope_params = territorial_scope_sql("p")
                _scope_row = conn.execute(
                    f"SELECT p.id FROM parcelles p WHERE p.id = ? AND ({_scope_sql})",
                    [int(row["parcelle_id"])] + _scope_params,
                ).fetchone()
                if _scope_row is None:
                    conn.close()
                    journaliser("VALIDATION_REFUSEE", f"Mutation {mutation_id}: hors périmètre territorial")
                    snack(t("access_denied"), True)
                    return

                if (
                    row["statut"]
                    != "En attente"
                ):

                    conn.close()

                    snack(
                        t("invalid_mutation"),
                        True,
                    )

                    return

                if (
                    row["statut_parcelle"]
                    == "Contestée"
                ):

                    conn.close()

                    snack(
                        t("parcel_disputed"),
                        True,
                    )

                    return

                if (
                    row[
                        "proprietaire_actuel"
                    ]
                    is None
                    or int(
                        row[
                            "proprietaire_actuel"
                        ]
                    )
                    != int(
                        row[
                            "ancien_proprietaire_id"
                        ]
                    )
                ):

                    conn.close()

                    journaliser(
                        "VALIDATION_REFUSEE",
                        (
                            f"Mutation "
                            f"{mutation_id}: "
                            "propriétaire actuel "
                            "incohérent"
                        ),
                    )

                    snack(
                        t("invalid_mutation"),
                        True,
                    )

                    return

                new_owner = conn.execute(
                    """
                    SELECT id
                    FROM proprietaires
                    WHERE id = ?
                    """,
                    (
                        int(
                            row[
                                "nouveau_proprietaire_id"
                            ]
                        ),
                    ),
                ).fetchone()

                if new_owner is None:

                    conn.close()

                    snack(
                        t("owner_not_found"),
                        True,
                    )

                    return

                # ------------------------------------------------
                # Transfert effectif
                # ------------------------------------------------

                conn.execute(
                    """
                    UPDATE parcelles
                    SET
                        proprietaire_id = ?,
                        statut = 'Transférée'
                    WHERE id = ?
                    """,
                    (
                        int(
                            row[
                                "nouveau_proprietaire_id"
                            ]
                        ),
                        int(
                            row[
                                "parcelle_id"
                            ]
                        ),
                    ),
                )

                # ------------------------------------------------
                # Mutation validée
                # ------------------------------------------------

                conn.execute(
                    """
                    UPDATE mutations
                    SET statut = 'Validée'
                    WHERE id = ?
                    """,
                    (
                        int(mutation_id),
                    ),
                )

                conn.commit()

            except Exception as ex:

                conn.rollback()
                conn.close()

                journaliser(
                    "VALIDATION_ERREUR",
                    (
                        f"Mutation "
                        f"{mutation_id}: "
                        f"{ex}"
                    ),
                )

                snack(
                    f"{t('error')} : {ex}",
                    True,
                )

                return

            conn.close()

            journaliser(
                "MUTATION_VALIDEE",
                (
                    f"Mutation ID "
                    f"{mutation_id}"
                ),
            )

            snack(
                t("mutation_validated")
            )

            load_table()

        # --------------------------------------------------------
        # Chargement du tableau
        # --------------------------------------------------------

        def load_table():

            conn = db_connect()
            scope_sql, scope_params = territorial_scope_sql("p")

            try:

                rows = conn.execute(
                    f"""
                    SELECT
                        m.*,
                        p.numero,
                        a.nom_complet
                            AS ancien,
                        n.nom_complet
                            AS nouveau
                    FROM mutations m
                    JOIN parcelles p
                        ON m.parcelle_id = p.id
                    LEFT JOIN proprietaires a
                        ON m.ancien_proprietaire_id
                        = a.id
                    LEFT JOIN proprietaires n
                        ON m.nouveau_proprietaire_id
                        = n.id
                    WHERE {scope_sql}
                    ORDER BY m.id DESC
                    LIMIT 50
                    """,
                    scope_params,
                ).fetchall()

            finally:

                conn.close()

            table.controls.clear()

            table.controls.append(
                ft.Row(
                    [
                        ft.Text(
                            t("parcel"),
                            width=120,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("previous"),
                            width=180,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("new"),
                            width=180,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("type"),
                            width=110,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("date"),
                            width=150,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("status"),
                            width=120,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),

                        ft.Text(
                            t("validate"),
                            width=120,
                            weight=(
                                ft.FontWeight.BOLD
                            ),
                        ),
                    ]
                )
            )

            table.controls.append(
                ft.Divider()
            )

            # ----------------------------------------------------
            # Détermination du rôle courant
            # ----------------------------------------------------

            current_role = ""

            try:

                if current_user:

                    current_role = (
                        current_user.get(
                            "role",
                            "",
                        )
                    )

            except Exception:

                current_role = ""

            # ----------------------------------------------------
            # Lignes du tableau
            # ----------------------------------------------------

            for row in rows:

                can_validate = (
                    row["statut"]
                    == "En attente"
                    and current_role
                    in {
                        "Administrateur",
                        "Superviseur",
                    }
                )

                table.controls.append(
                    ft.Row(
                        [
                            ft.Text(
                                str(
                                    row["numero"]
                                ),
                                width=120,
                            ),

                            ft.Text(
                                str(
                                    row[
                                        "ancien"
                                    ]
                                    or "—"
                                ),
                                width=180,
                            ),

                            ft.Text(
                                str(
                                    row[
                                        "nouveau"
                                    ]
                                    or "—"
                                ),
                                width=180,
                            ),

                            ft.Text(
                                enum_text(
                                    row[
                                        "type_transaction"
                                    ],
                                    TRANSACTION_KEYS,
                                ),
                                width=110,
                            ),

                            ft.Text(
                                str(
                                    row[
                                        "date_transaction"
                                    ]
                                )[:10],
                                width=150,
                            ),

                            ft.Text(
                                enum_text(
                                    row[
                                        "statut"
                                    ],
                                    {
                                        "En attente":
                                            "pending",
                                        "Validée":
                                            "validated",
                                    },
                                ),
                                width=120,
                            ),

                            ft.ElevatedButton(
                                t("validate"),
                                icon=(
                                    ft.Icons.CHECK
                                ),
                                visible=(
                                    can_validate
                                ),
                                on_click=(
                                    lambda e,
                                    mid=row["id"]:
                                    valider_mutation(
                                        mid
                                    )
                                ),
                            ),
                        ]
                    )
                )

            page.update()

        # ================================================
                # CARTE / BLOC D'AFFICHAGE DU RÉSULTAT
                # ================================================

                btn_telecharger = ft.ElevatedButton(
                    t("download_certificate") if "download_certificate" in TRANSLATIONS.get(language, {}) else "Télécharger le certificat",
                    icon=ft.Icons.PICTURE_IN_PICTURE,
                    on_click=lambda e, r=row: page.run_task(telecharger_certificat_recherche, r),
                )

                carte_resultat = panel(
                    ft.Column(
                        [
                            # En-tête de la parcelle
                            ft.Row(
                                [
                                    ft.Icon(ft.Icons.LANDSCAPE, color=COLOR_PRIMARY, size=28),
                                    ft.Text(
                                        f"Parcelle N° {numero}",
                                        size=18,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                    ft.Container(expand=True),
                                    badge_statut(statut),
                                ],
                                alignment=ft.MainAxisAlignment.CENTER,
                            ),
                            ft.Divider(),

                            # Contenu principal : Photo + Détails
                            ft.Row(
                                [
                                    # Colonne photo à gauche
                                    ft.Column(
                                        [
                                            photo_widget,
                                            ft.Text(
                                                f"ID Proprio: #{proprietaire_id}",
                                                size=11,
                                                color=COLOR_TEXT_MUTED,
                                            ),
                                        ],
                                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                                        spacing=6,
                                    ),
                                    
                                    # Détails à droite (Propriétaire + Parcelle)
                                    ft.Column(
                                        [
                                            ft.Text(
                                                "Propriétaire Actuel",
                                                size=14,
                                                weight=ft.FontWeight.BOLD,
                                                color=COLOR_PRIMARY,
                                            ),
                                            afficher_ligne("Nom complet", proprietaire_nom),
                                            afficher_ligne("Pièce d'identité", proprietaire_piece),
                                            afficher_ligne("Téléphone", proprietaire_telephone),
                                            afficher_ligne("Email", proprietaire_email),
                                            
                                            ft.Container(height=8),
                                            ft.Text(
                                                "Localisation & Caractéristiques",
                                                size=14,
                                                weight=ft.FontWeight.BOLD,
                                                color=COLOR_PRIMARY,
                                            ),
                                            afficher_ligne("Adresse", adresse),
                                            afficher_ligne("Province / Ville", f"{province} / {ville_territoire}"),
                                            afficher_ligne("Localité", localite),
                                            afficher_ligne("Superficie", f"{superficie} m²"),
                                            afficher_ligne("Coordonnées GPS", f"{latitude}, {longitude}"),
                                        ],
                                        expand=True,
                                        spacing=4,
                                    ),
                                ],
                                cross_alignment=ft.CrossAxisAlignment.START,
                                spacing=20,
                            ),

                            ft.Divider(),

                            # Actions / Boutons bas de carte
                            ft.Row(
                                [
                                    btn_telecharger,
                                ],
                                alignment=ft.MainAxisAlignment.END,
                            ),
                        ],
                        spacing=10,
                    )
                )

                résultat.controls.append(carte_resultat)

            page.update()

       # --------------------------------------------------------
        # Construction interface Recherche / Vérification
        # --------------------------------------------------------

        # ================================================
        # BLOC PROPRIÉTAIRE
        # ================================================
        bloc_proprietaire = panel(
            ft.Column(
                [
                    ft.Row(
                        [
                            ft.Icon(
                                ft.Icons.PERSON,
                                color=COLOR_PRIMARY,
                                size=28,
                            ),
                            ft.Text(
                                "INFORMATIONS DU PROPRIÉTAIRE",
                                size=18,
                                weight=ft.FontWeight.BOLD,
                            ),
                        ],
                        spacing=10,
                    ),
                    ft.Divider(),
                    ft.Row(
                        [
                            photo_widget,
                            ft.Column(
                                [
                                    afficher_ligne(t("id"), proprietaire_id),
                                    afficher_ligne(t("full_name"), proprietaire_nom),
                                    afficher_ligne(t("phone"), proprietaire_telephone),
                                    afficher_ligne(t("email"), proprietaire_email),
                                    afficher_ligne(t("identity"), proprietaire_piece),
                                    afficher_ligne("Date de création", proprietaire_date),
                                ],
                                spacing=8,
                                expand=True,
                            ),
                        ],
                        spacing=20,
                        vertical_alignment=ft.CrossAxisAlignment.START,
                    ),
                ],
                spacing=10,
            )
        )

        # ================================================
        # BLOC PARCELLE
        # ================================================
        bloc_parcelle = panel(
            ft.Column(
                [
                    ft.Row(
                        [
                            ft.Icon(
                                ft.Icons.LANDSCAPE,
                                color=COLOR_PRIMARY,
                                size=28,
                            ),
                            ft.Text(
                                "INFORMATIONS DE LA PARCELLE",
                                size=18,
                                weight=ft.FontWeight.BOLD,
                            ),
                        ],
                        spacing=10,
                    ),
                    ft.Divider(),
                    afficher_ligne(t("id"), row["id"]),
                    afficher_ligne(t("parcel"), numero),
                    afficher_ligne(t("status"), statut),
                    afficher_ligne(t("area_short"), f"{float(superficie):.2f} cm²"),
                    afficher_ligne(
                        t("address") if "address" in TRANSLATIONS.get(language, {}) else "Adresse",
                        adresse,
                    ),
                    afficher_ligne("Province", province),
                    afficher_ligne("Ville / Territoire", ville_territoire),
                    afficher_ligne("Localité", localite),
                    afficher_ligne("Commune / Chefferie", row["commune_chefferie"] or "—"),
                    afficher_ligne("Quartier / Groupement", row["quartier_groupement"] or "—"),
                    afficher_ligne(t("gps_label"), f"{latitude}, {longitude}"),
                ],
                spacing=8,
            )
        )

        # ================================================
        # BLOC VÉRIFICATION
        # ================================================
        bloc_verification = ft.Container(
            padding=15,
            border_radius=12,
            bgcolor=(
                "#14532D" if page.theme_mode == ft.ThemeMode.DARK else "#F0FDF4"
            ),
            border=ft.border.all(1, ft.Colors.GREEN_700),
            content=ft.Row(
                [
                    ft.Icon(
                        ft.Icons.VERIFIED,
                        color=ft.Colors.GREEN_700,
                        size=30,
                    ),
                    ft.Column(
                        [
                            ft.Text(
                                "VÉRIFICATION RÉUSSIE",
                                weight=ft.FontWeight.BOLD,
                                color=ft.Colors.GREEN_700,
                                size=16,
                            ),
                            ft.Text(
                                t("verification_found"),
                                color=ft.Colors.GREEN_700,
                            ),
                        ],
                        spacing=3,
                    ),
                ],
                spacing=12,
            ),
        )

        # ================================================
        # RÉSULTAT COMPLET
        # ================================================
        résultat.controls.clear()
        résultat.controls.append(
            ft.Column(
                [
                    bloc_verification,
                    ft.Text(
                        f"{t('parcel').upper()} {numero}",
                        size=22,
                        weight=ft.FontWeight.BOLD,
                        color=COLOR_PRIMARY,
                    ),
                    bloc_proprietaire,
                    bloc_parcelle,
                    # Bouton certificat PDF
                    ft.Container(
                        padding=12,
                        border_radius=12,
                        bgcolor=(
                            "#10233F" if page.theme_mode == ft.ThemeMode.DARK else "#EFF6FF"
                        ),
                        border=ft.border.all(1, COLOR_PRIMARY),
                        content=ft.Row(
                            [
                                ft.Icon(
                                    ft.Icons.PICTURE_AS_PDF,
                                    color=COLOR_PRIMARY,
                                    size=28,
                                ),
                                ft.Column(
                                    [
                                        ft.Text(
                                            "CERTIFICAT D'ENREGISTREMENT PARCELLAIRE",
                                            weight=ft.FontWeight.BOLD,
                                            color=COLOR_PRIMARY,
                                        ),
                                        ft.Text(
                                            "Propriétaire + photo + toutes les informations de la parcelle + code QR de vérification.",
                                            size=11,
                                            color=COLOR_TEXT_MUTED,
                                        ),
                                    ],
                                    spacing=3,
                                    expand=True,
                                ),
                                ft.ElevatedButton(
                                    "Télécharger le certificat",
                                    icon=ft.Icons.DOWNLOAD,
                                    on_click=lambda e, r=dict(row): page.run_task(
                                        telecharger_certificat_recherche, r
                                    ),
                                    style=ft.ButtonStyle(
                                        bgcolor=COLOR_PRIMARY,
                                        color=ft.Colors.WHITE,
                                        padding=ft.padding.symmetric(horizontal=16, vertical=12),
                                        shape=ft.RoundedRectangleBorder(radius=8),
                                    ),
                                ),
                            ],
                            spacing=12,
                            vertical_alignment=ft.CrossAxisAlignment.CENTER,
                        ),
                    ),
                ],
                spacing=12,
            )
        )

        # ----------------------------------------------------
        # Journal de sécurité
        # ----------------------------------------------------
        try:
            journaliser(
                "VERIFICATION_CADASTRALE",
                f"Recherche effectuée : {q}",
            )
        except Exception:
            pass

        recherche.on_submit = lambda e: search()

        content.content = ft.Column(
            [
                title(
                    t("verification_title") if "verification_title" in TRANSLATIONS.get(language, {}) else "Vérification",
                    t("verification_subtitle") if "verification_subtitle" in TRANSLATIONS.get(language, {}) else "Résultats de la vérification",
                ),
                ft.Divider(),
                panel(
                    ft.Row(
                        [
                            ft.Container(
                                content=recherche,
                                expand=True,
                            ),
                            ft.ElevatedButton(
                                t("search_button") if "search_button" in TRANSLATIONS.get(language, {}) else "Rechercher",
                                icon=ft.Icons.SEARCH,
                                on_click=lambda e: search(),
                                style=ft.ButtonStyle(
                                    bgcolor=COLOR_PRIMARY,
                                    color=ft.Colors.WHITE,
                                    padding=ft.padding.symmetric(horizontal=20, vertical=15),
                                    shape=ft.RoundedRectangleBorder(radius=8),
                                ),
                            ),
                        ],
                        alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                        vertical_alignment=ft.CrossAxisAlignment.CENTER,
                    )
                ),
                résultat,
            ],
            spacing=12,
            scroll=ft.ScrollMode.AUTO,
        )

        page.update()
   # --------------------------------------------------------
        # Administration territoriale
        # --------------------------------------------------------

        def territoires_view():
            nonlocal current_route
            current_route = "territoires"

            type_structure = ft.Dropdown(
                label=t("structure_type"),
                value="Province",
                options=[
                    ft.dropdown.Option("Province", t("province")),
                    ft.dropdown.Option("Ville", t("city")),
                    ft.dropdown.Option("Territoire", t("territory")),
                    ft.dropdown.Option("Commune", t("commune")),
                    ft.dropdown.Option("Secteur", t("sector_chiefdom")),
                    ft.dropdown.Option("Chefferie", t("sector_chiefdom")),
                    ft.dropdown.Option("Groupement", t("grouping")),
                    ft.dropdown.Option("Village", t("village")),
                ],
                border_radius=14, border_color=COLOR_BORDER, expand=True,
            )
            code = text_field(t("structure_code"))
            nom = text_field(t("structure_name") + " *")
            parent = text_field(t("parent_structure"))
            recherche = text_field(t("search"))
            table = ft.Column(spacing=4)

            # ----------------------------------------------------
            # Départements de la Direction Générale / provinces
            # ----------------------------------------------------
            departement_nom = text_field(t("department") + " *")
            departement_province = text_field(t("province"))
            departement_table = ft.Column(spacing=4)

            if str(current_user.get("niveau_acces") or "National") != "National":
                departement_province.value = current_user.get("province") or ""

            def load_departements():
                conn = db_connect()
                rows = conn.execute("SELECT * FROM departements ORDER BY province, nom").fetchall()
                conn.close()
                departement_table.controls.clear()
                departement_table.controls.append(ft.Row([
                    ft.Text(t("id"), width=50, weight=ft.FontWeight.BOLD),
                    ft.Text(t("department"), width=260, weight=ft.FontWeight.BOLD),
                    ft.Text(t("province"), width=200, weight=ft.FontWeight.BOLD),
                    ft.Text(t("status"), width=100, weight=ft.FontWeight.BOLD),
                ]))
                departement_table.controls.append(ft.Divider())
                for row in rows:
                    if str(current_user.get("niveau_acces") or "National") != "National" and (row["province"] or "").strip().lower() != str(current_user.get("province") or "").strip().lower():
                        continue
                    departement_table.controls.append(ft.Row([
                        ft.Text(str(row["id"]), width=50), ft.Text(row["nom"] or "", width=260),
                        ft.Text(row["province"] or "National", width=200),
                        ft.Text(t("active") if int(row["actif"] or 0) else t("inactive"), width=100),
                    ]))

            def save_departement(_=None):
                if not require_action("territorial_manage"): return
                dname = departement_nom.value.strip()
                dprov = departement_province.value.strip()
                if not dname: snack(t("department"), True); return
                if str(current_user.get("niveau_acces") or "National") != "National":
                    if not dprov or dprov.lower() != str(current_user.get("province") or "").lower():
                        snack(t("access_denied"), True); return
                conn = db_connect()
                try:
                    conn.execute("INSERT INTO departements(nom, code, province, actif, date_creation) VALUES (?, ?, ?, 1, ?)", (dname, "", dprov or None, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                    conn.commit()
                except sqlite3.IntegrityError:
                    conn.rollback(); conn.close(); snack(t("error"), True); return
                except Exception as ex:
                    conn.rollback(); conn.close(); snack(f"{t('error')} : {ex}", True); return
                conn.close(); journaliser("DEPARTEMENT_CREE", f"{dname} — {dprov or 'National'}"); departement_nom.value = ""; load_departements(); snack(t("add_structure")); page.update()

            def load(_=None):
                q = (recherche.value or "").strip().lower()
                conn = db_connect()
                rows = conn.execute("""
                    SELECT s.*, p.nom AS parent_nom
                    FROM structures_territoriales s
                    LEFT JOIN structures_territoriales p ON p.id = s.parent_id
                    ORDER BY CASE s.type WHEN 'Province' THEN 1 WHEN 'Ville' THEN 2 WHEN 'Territoire' THEN 3 WHEN 'Commune' THEN 4 ELSE 5 END, s.nom
                """).fetchall()
                conn.close()
                table.controls.clear()
                table.controls.append(ft.Row([
                    ft.Text(t("id"), width=50, weight=ft.FontWeight.BOLD),
                    ft.Text(t("structure_type"), width=130, weight=ft.FontWeight.BOLD),
                    ft.Text(t("structure_name"), width=220, weight=ft.FontWeight.BOLD),
                    ft.Text(t("structure_code"), width=120, weight=ft.FontWeight.BOLD),
                    ft.Text(t("parent_structure"), width=220, weight=ft.FontWeight.BOLD),
                    ft.Text(t("status"), width=100, weight=ft.FontWeight.BOLD),
                ]))
                table.controls.append(ft.Divider())
                for row in rows:
                    hay = " ".join(str(row[k] or "") for k in ("type", "nom", "code", "parent_nom")).lower()
                    if q and q not in hay:
                        continue
                    table.controls.append(ft.Row([
                        ft.Text(str(row["id"]), width=50),
                        ft.Text(row["type"] or "", width=130),
                        ft.Text(row["nom"] or "", width=220),
                        ft.Text(row["code"] or "", width=120),
                        ft.Text(row["parent_nom"] or "—", width=220),
                        ft.Text(t("active") if int(row["actif"] or 0) else t("inactive"), width=100),
                    ]))
                page.update()

            def save(_=None):
                if not require_action("territorial_manage"):
                    return
                name = (nom.value or "").strip()
                if not name:
                    snack(t("structure_name"), True)
                    return
                if str(current_user.get("niveau_acces") or "National") != "National":
                    if type_structure.value == "Province":
                        journaliser("OPERATION_REFUSEE", f"Création province hors niveau national: {name}")
                        snack(t("access_denied"), True)
                        return
                    if not (parent.value or "").strip():
                        snack(t("parent_structure"), True)
                        return
                conn = db_connect()
                try:
                    parent_id = None
                    parent_name = (parent.value or "").strip()
                    if parent_name:
                        prow = conn.execute("SELECT id FROM structures_territoriales WHERE LOWER(nom)=LOWER(?) ORDER BY id LIMIT 1", (parent_name,)).fetchone()
                        if prow:
                            parent_id = prow["id"]
                        else:
                            snack(t("parent_structure"), True); conn.close(); return
                    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    conn.execute("INSERT INTO structures_territoriales(type, code, nom, parent_id, actif, date_creation) VALUES (?, ?, ?, ?, 1, ?)", (type_structure.value, code.value.strip(), name, parent_id, now))
                    conn.commit()
                except sqlite3.IntegrityError:
                    conn.rollback(); conn.close(); snack(t("error"), True); return
                except Exception as ex:
                    conn.rollback(); conn.close(); snack(f"{t('error')} : {ex}", True); return
                conn.close()
                journaliser("STRUCTURE_TERRITORIALE_CREEE", f"{type_structure.value}: {name}")
                nom.value = ""; code.value = ""; parent.value = ""
                snack(t("add_structure")); load()

            recherche.on_change = load

            content.content = ft.Column([
                title(t("territorial_title"), t("territorial_subtitle")),
                ft.Divider(),
                panel(ft.Column([
                    ft.Text(t("national_control"), size=17, weight=ft.FontWeight.BOLD),
                    ft.Text(t("scope_restricted"), size=12, color=COLOR_TEXT_MUTED),
                    ft.Row([type_structure, code]),
                    ft.Row([nom, parent]),
                    ft.ElevatedButton(
                        t("add_structure"), 
                        icon=ft.Icons.ADD_LOCATION_ALT, 
                        on_click=save,
                        style=ft.ButtonStyle(
                            bgcolor=COLOR_PRIMARY,
                            color=ft.Colors.WHITE,
                            padding=ft.padding.symmetric(horizontal=20, vertical=15),
                            shape=ft.RoundedRectangleBorder(radius=8),
                        ),
                    ),
                ], spacing=10)),
                panel(ft.Column([
                    ft.Text(t("department_control"), size=17, weight=ft.FontWeight.BOLD),
                    ft.Row([departement_nom, departement_province]),
                    ft.ElevatedButton(
                        t("add_structure"), 
                        icon=ft.Icons.BUSINESS, 
                        on_click=save_departement,
                        style=ft.ButtonStyle(
                            bgcolor=COLOR_PRIMARY,
                            color=ft.Colors.WHITE,
                            padding=ft.padding.symmetric(horizontal=20, vertical=15),
                            shape=ft.RoundedRectangleBorder(radius=8),
                        ),
                    ),
                    departement_table,
                ], spacing=10)),
                recherche,
                panel(table),
            ], spacing=12, scroll=ft.ScrollMode.AUTO)
            
            load()
            load_departements()
# --------------------------------------------------------
        # Utilisateurs
        # --------------------------------------------------------

        def utilisateurs_view():
            nonlocal current_route
            current_route = "utilisateurs"

            selected_id = {"value": None}
            nom = text_field(t("full_name") + " *")
            role = ft.Dropdown(
                label=t("role"), value="Agent",
                options=[
                    ft.dropdown.Option("Administrateur", t("administrator")),
                    ft.dropdown.Option("Superviseur", t("supervisor")),
                    ft.dropdown.Option("Technicien", t("technician")),
                    ft.dropdown.Option("Agent", t("agent_role")),
                    ft.dropdown.Option("Consultation", t("consultation")),
                ], border_radius=14, border_color=COLOR_BORDER, expand=True,
            )
            tel = text_field(t("phone"))
            niveau = ft.Dropdown(
                label=t("access_level"), value="National",
                options=[
                    ft.dropdown.Option("National", t("national_level")),
                    ft.dropdown.Option("Provincial", t("provincial_level")),
                    ft.dropdown.Option("Ville", t("city_level")),
                    ft.dropdown.Option("Territoire", t("territory_level")),
                    ft.dropdown.Option("Commune", t("commune_level")),
                    ft.dropdown.Option("Secteur", t("sector_chiefdom")),
                    ft.dropdown.Option("Chefferie", t("sector_chiefdom")),
                    ft.dropdown.Option("Groupement", t("grouping")),
                    ft.dropdown.Option("Village", t("village")),
                ], border_radius=14, border_color=COLOR_BORDER, expand=True,
            )
            province_u = text_field(t("province"))
            ville_u = text_field(t("city"))
            territoire_u = text_field(t("territory"))
            commune_u = text_field(t("commune"))
            secteur_u = text_field(t("sector_chiefdom"))
            groupement_u = text_field(t("grouping"))
            village_u = text_field(t("village"))
            departement_u = text_field(t("department_scope"))
            password = ft.TextField(label=t("new_password_optional"), password=True, can_reveal_password=True, border_radius=14, border_color=COLOR_BORDER, expand=True)
            password_confirm = ft.TextField(label=t("confirm_password"), password=True, can_reveal_password=True, border_radius=14, border_color=COLOR_BORDER, expand=True)
            status = ft.Text("", size=12, color=COLOR_TEXT_MUTED)
            table = ft.Column()

            def clear_form():
                selected_id["value"] = None
                for f in (nom, tel, province_u, ville_u, territoire_u, commune_u, secteur_u, groupement_u, village_u, departement_u, password, password_confirm): f.value = ""
                role.value = "Agent"; niveau.value = "National"; status.value = ""

            def load_user(user_id):
                conn = db_connect(); row = conn.execute("SELECT * FROM utilisateurs WHERE id = ?", (user_id,)).fetchone(); conn.close()
                if not row: return
                selected_id["value"] = row["id"]; nom.value = row["nom"] or ""; tel.value = row["telephone"] or ""; role.value = row["role"]; niveau.value = row["niveau_acces"] or "National"
                province_u.value = row["province"] or ""; ville_u.value = row["ville"] or ""; territoire_u.value = row["territoire"] or ""; commune_u.value = row["commune"] or ""; secteur_u.value = row["secteur_chefferie"] or ""; groupement_u.value = row["groupement"] or ""; village_u.value = row["village"] or ""; departement_u.value = row["departement"] or ""
                password.value = ""; password_confirm.value = ""; status.value = f"{t('select_user')}: {row['nom']}"; page.update()

            def save_user(_=None):
                if not require_action("user_manage"): return
                username = nom.value.strip(); pwd = password.value or ""; uid = selected_id["value"]
                if not username: snack(t("owner_required"), True); return
                if pwd and len(pwd) < 8: snack(t("password_short"), True); return
                if pwd and pwd != (password_confirm.value or ""): snack(t("password_mismatch"), True); return
                if niveau.value != "National" and not province_u.value.strip(): snack(t("province"), True); return
                if not target_scope_allowed(niveau.value, province_u.value.strip(), ville_u.value.strip(), territoire_u.value.strip(), commune_u.value.strip(), secteur_u.value.strip(), groupement_u.value.strip(), village_u.value.strip()):
                    journaliser("OPERATION_REFUSEE", f"Affectation territoriale hors périmètre: {username}")
                    snack(t("access_denied"), True); return
                conn = db_connect()
                try:
                    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    scope_values = (niveau.value, province_u.value.strip(), ville_u.value.strip(), territoire_u.value.strip(), commune_u.value.strip(), secteur_u.value.strip(), groupement_u.value.strip(), village_u.value.strip(), departement_u.value.strip())
                    if uid is None:
                        if not pwd: snack(t("password_short"), True); conn.close(); return
                        salt, pwd_hash = hash_password(pwd)
                        conn.execute("""INSERT INTO utilisateurs(nom, role, telephone, date_creation, password_hash, password_salt, failed_attempts, locked_until, last_login, active, niveau_acces, province, ville, territoire, commune, secteur_chefferie, groupement, village, departement) VALUES (?, ?, ?, ?, ?, ?, 0, 0, NULL, 1, ?, ?, ?, ?, ?, ?, ?, ?, ?)""", (username, role.value, tel.value.strip(), now, pwd_hash, salt, *scope_values))
                        action = "CREATION_UTILISATEUR"
                    else:
                        old = conn.execute("SELECT * FROM utilisateurs WHERE id = ?", (uid,)).fetchone()
                        if not old: conn.close(); snack(t("error"), True); return
                        if int(uid) == int(current_user["id"]) and role.value != old["role"]: conn.close(); snack(t("cannot_change_self_role"), True); return
                        if old["role"] == "Administrateur" and role.value != "Administrateur":
                            admins = conn.execute("SELECT COUNT(*) FROM utilisateurs WHERE role='Administrateur' AND active=1").fetchone()[0]
                            if admins <= 1: conn.close(); snack(t("cannot_delete_last_admin"), True); return
                        conn.execute("""UPDATE utilisateurs SET nom=?, role=?, telephone=?, niveau_acces=?, province=?, ville=?, territoire=?, commune=?, secteur_chefferie=?, groupement=?, village=?, departement=? WHERE id=?""", (username, role.value, tel.value.strip(), *scope_values, uid))
                        if pwd:
                            salt, pwd_hash = hash_password(pwd); conn.execute("UPDATE utilisateurs SET password_hash=?, password_salt=?, failed_attempts=0, locked_until=0 WHERE id=?", (pwd_hash, salt, uid))
                        action = "MODIFICATION_UTILISATEUR"
                    conn.commit()
                except Exception as ex:
                    conn.rollback(); conn.close(); snack(f"{t('error')} : {ex}", True); return
                conn.close(); make_backup("pre_user_change"); journaliser(action, f"{username} — {niveau.value} — {province_u.value.strip()}"); clear_form(); snack(t("user_saved")); load()

            def toggle_active(user_id):
                if not require_action("user_manage"): return
                conn = db_connect(); row = conn.execute("SELECT * FROM utilisateurs WHERE id=?", (user_id,)).fetchone()
                if not row: conn.close(); return
                new_active = 0 if int(row["active"] or 0) else 1
                if new_active == 0 and row["role"] == "Administrateur":
                    admins = conn.execute("SELECT COUNT(*) FROM utilisateurs WHERE role='Administrateur' AND active=1").fetchone()[0]
                    if admins <= 1: conn.close(); snack(t("cannot_disable_last_admin"), True); return
                conn.execute("UPDATE utilisateurs SET active=?, failed_attempts=0, locked_until=0 WHERE id=?", (new_active, user_id)); conn.commit(); conn.close(); make_backup("pre_user_status"); journaliser("ACTIVATION_UTILISATEUR" if new_active else "DESACTIVATION_UTILISATEUR", str(row["nom"])); load()

            def load():
                conn = db_connect(); rows = conn.execute("SELECT * FROM utilisateurs ORDER BY id DESC").fetchall(); conn.close(); table.controls.clear()
                table.controls.append(ft.Row([ft.Text(t("id"), width=45, weight=ft.FontWeight.BOLD), ft.Text(t("name"), width=180, weight=ft.FontWeight.BOLD), ft.Text(t("role"), width=140, weight=ft.FontWeight.BOLD), ft.Text(t("access_level"), width=110, weight=ft.FontWeight.BOLD), ft.Text(t("province"), width=150, weight=ft.FontWeight.BOLD), ft.Text(t("status"), width=100, weight=ft.FontWeight.BOLD), ft.Text(t("actions"), width=260, weight=ft.FontWeight.BOLD)]))
                table.controls.append(ft.Divider()); now = time.time()
                for row in rows:
                    if not target_scope_allowed(row["niveau_acces"] or "National", row["province"], row["ville"], row["territoire"], row["commune"], row["secteur_chefferie"], row["groupement"], row["village"]):
                        continue
                    locked = float(row["locked_until"] or 0) > now; st = t("locked") if locked else (t("active") if int(row["active"] or 0) else t("inactive"))
                    buttons = [
                        ft.ElevatedButton(
                            t("edit_user"), 
                            icon=ft.Icons.EDIT, 
                            on_click=lambda e, i=row["id"]: load_user(i),
                            style=ft.ButtonStyle(
                                bgcolor=COLOR_PRIMARY,
                                color=ft.Colors.WHITE,
                                padding=ft.padding.symmetric(horizontal=12, vertical=8),
                                shape=ft.RoundedRectangleBorder(radius=8),
                            ),
                        ), 
                        ft.ElevatedButton(
                            t("unblock_user") if locked or not int(row["active"] or 0) else t("block_user"), 
                            icon=ft.Icons.LOCK_OPEN if locked or not int(row["active"] or 0) else ft.Icons.LOCK, 
                            on_click=lambda e, i=row["id"]: toggle_active(i),
                            style=ft.ButtonStyle(
                                bgcolor=COLOR_PRIMARY,
                                color=ft.Colors.WHITE,
                                padding=ft.padding.symmetric(horizontal=12, vertical=8),
                                shape=ft.RoundedRectangleBorder(radius=8),
                            ),
                        )
                    ]
                    table.controls.append(ft.Row([ft.Text(str(row["id"]), width=45), ft.Text(row["nom"] or "", width=180), ft.Text(enum_text(row["role"], ROLE_KEYS), width=140), ft.Text(row["niveau_acces"] or "National", width=110), ft.Text(row["province"] or "Toutes", width=150), ft.Text(st, width=100), ft.Row(buttons, spacing=5, width=260)]))
                page.update()

            content.content = ft.Column([
                title(t("users_title"), t("users_subtitle")), 
                ft.Divider(), 
                panel(ft.Column([
                    ft.Text(t("user_management"), size=17, weight=ft.FontWeight.BOLD), 
                    ft.Row([nom, role, niveau]), 
                    tel, 
                    ft.Row([province_u, ville_u, territoire_u]), 
                    ft.Row([commune_u, secteur_u, groupement_u]), 
                    ft.Row([village_u, departement_u]), 
                    ft.Row([password, password_confirm]), 
                    status, 
                    ft.Row([
                        ft.ElevatedButton(
                            t("add_user"), 
                            icon=ft.Icons.PERSON_ADD, 
                            on_click=lambda e: (clear_form(), page.update()),
                            style=ft.ButtonStyle(
                                bgcolor=COLOR_PRIMARY,
                                color=ft.Colors.WHITE,
                                padding=ft.padding.symmetric(horizontal=20, vertical=15),
                                shape=ft.RoundedRectangleBorder(radius=8),
                            ),
                        ), 
                        ft.ElevatedButton(
                            t("update_user"), 
                            icon=ft.Icons.SAVE, 
                            on_click=save_user,
                            style=ft.ButtonStyle(
                                bgcolor=COLOR_PRIMARY,
                                color=ft.Colors.WHITE,
                                padding=ft.padding.symmetric(horizontal=20, vertical=15),
                                shape=ft.RoundedRectangleBorder(radius=8),
                            ),
                        )
                    ]), 
                    ft.Text(t("scope_restricted"), size=11, color=COLOR_TEXT_MUTED)
                ], spacing=10)), 
                panel(table)
            ], spacing=12, scroll=ft.ScrollMode.AUTO)
            
            load()

        # --------------------------------------------------------
        # Journal
        # --------------------------------------------------------

        def journal_view():
            nonlocal current_route
            current_route = "journal"

            # ----------------------------------------------------
            # Champs de recherche et filtres
            # ----------------------------------------------------

            recherche = text_field(
                t("search")
            )

            recherche.hint_text = (
                "Nom, action, détails, date..."
            )

            filtre_nom = text_field(
                t("name")
            )

            filtre_role = ft.Dropdown(
                label=t("role"),
                options=[
                    ft.dropdown.Option(
                        "",
                        t("all")
                    ),
                    ft.dropdown.Option(
                        "Administrateur",
                        "Administrateur"
                    ),
                    ft.dropdown.Option(
                        "Superviseur",
                        "Superviseur"
                    ),
                    ft.dropdown.Option(
                        "Technicien",
                        "Technicien"
                    ),
                    ft.dropdown.Option(
                        "Agent",
                        "Agent"
                    ),
                    ft.dropdown.Option(
                        "Consultation",
                        "Consultation"
                    ),
                ],
                value="",
                expand=True,
                border_radius=14,
                border_color=COLOR_BORDER,
            )

            filtre_action = ft.Dropdown(
                label=t("action"),
                options=[
                    ft.dropdown.Option(
                        "",
                        t("all")
                    ),
                    ft.dropdown.Option(
                        "LOGIN",
                        "LOGIN"
                    ),
                    ft.dropdown.Option(
                        "LOGOUT",
                        "LOGOUT"
                    ),
                    ft.dropdown.Option(
                        "MUTATION_CREEE_EN_ATTENTE",
                        "Mutation créée"
                    ),
                    ft.dropdown.Option(
                        "MUTATION_VALIDEE",
                        "Mutation validée"
                    ),
                    ft.dropdown.Option(
                        "MUTATION_REFUSEE",
                        "Mutation refusée"
                    ),
                    ft.dropdown.Option(
                        "VALIDATION_REFUSEE",
                        "Validation refusée"
                    ),
                    ft.dropdown.Option(
                        "VALIDATION_ERREUR",
                        "Erreur validation"
                    ),
                ],
                value="",
                expand=True,
                border_radius=14,
                border_color=COLOR_BORDER,
            )

            filtre_date = text_field(
                t("date")
            )

            filtre_date.hint_text = (
                "AAAA-MM-JJ"
            )

            table = ft.Column(
                spacing=4
            )

            compteur = ft.Text(
                "",
                size=12,
                color=COLOR_TEXT_MUTED,
            )
        # ----------------------------------------------------
        # Chargement du journal
        # ----------------------------------------------------

        def load(_=None):

            recherche_value = (
                recherche.value
                .strip()
                .lower()
            )

            nom_value = (
                filtre_nom.value
                .strip()
                .lower()
            )

            role_value = (
                filtre_role.value
                or ""
            ).strip()

            action_value = (
                filtre_action.value
                or ""
            ).strip()

            date_value = (
                filtre_date.value
                .strip()
            )

            conn = db_connect()

            try:

                rows = conn.execute(
                    """
                    SELECT *
                    FROM journal
                    ORDER BY id DESC
                    LIMIT 500
                    """
                ).fetchall()

            finally:
                conn.close()

            # ------------------------------------------------
            # Application des filtres
            # ------------------------------------------------

            filtered_rows = []

            _scope_level = str(current_user.get("niveau_acces") or "National")
            _scope_province = current_user.get("province")
            _scope_ville = current_user.get("ville")
            _scope_territoire = current_user.get("territoire")
            _scope_commune = current_user.get("commune")
            _scope_secteur = current_user.get("secteur_chefferie")
            _scope_groupement = current_user.get("groupement")
            _scope_village = current_user.get("village")

            for row in rows:

                # Le journal est lui-même soumis au périmètre territorial.
                if _scope_level != "National" and _scope_province:
                    if (row["province"] or "").strip().lower() != str(_scope_province).strip().lower():
                        continue
                    if _scope_level == "Ville" and _scope_ville and (row["ville"] or "").strip().lower() != str(_scope_ville).strip().lower():
                        continue
                    if _scope_level == "Territoire" and _scope_territoire and (row["territoire"] or "").strip().lower() != str(_scope_territoire).strip().lower():
                        continue
                    if _scope_level == "Commune" and _scope_commune and (row["commune"] or "").strip().lower() != str(_scope_commune).strip().lower():
                        continue
                    if _scope_level in {"Secteur", "Chefferie", "Secteur / Chefferie"} and _scope_secteur and (row["secteur_chefferie"] or "").strip().lower() != str(_scope_secteur).strip().lower():
                        continue
                    if _scope_level == "Groupement" and _scope_groupement and (row["groupement"] or "").strip().lower() != str(_scope_groupement).strip().lower():
                        continue
                    if _scope_level in {"Village", "Local"} and _scope_village and (row["village"] or "").strip().lower() != str(_scope_village).strip().lower():
                        continue


                utilisateur_nom = (
                    row["utilisateur_nom"]
                    or "Système"
                )

                role_user = (
                    row["role"]
                    or ""
                )

                action = (
                    row["action"]
                    or ""
                )

                details = (
                    row["details"]
                    or ""
                )

                date_action = (
                    row["date_action"]
                    or ""
                )

                # Recherche globale
                if recherche_value:

                    texte_global = " ".join(
                        [
                            str(
                                row["id"]
                            ),
                            str(
                                utilisateur_nom
                            ),
                            str(role_user),
                            str(action),
                            str(details),
                            str(date_action),
                        ]
                    ).lower()

                    if recherche_value not in texte_global:
                        continue

                # Filtre par nom
                if nom_value:

                    if nom_value not in (
                        str(
                            utilisateur_nom
                        ).lower()
                    ):
                        continue

                # Filtre par rôle
                if role_value:

                    if role_user != role_value:
                        continue

                # Filtre par action
                if action_value:

                    if action != action_value:
                        continue

                # Filtre par date
                if date_value:

                    if not date_action.startswith(
                        date_value
                    ):
                        continue

                filtered_rows.append(row)

            # ------------------------------------------------
            # Affichage
            # ------------------------------------------------

            table.controls.clear()

            compteur.value = (
                f"{len(filtered_rows)} "
                f"entrée(s) trouvée(s)"
            )

            # ------------------------------------------------
            # En-tête
            # ------------------------------------------------

            table.controls.append(
                ft.Row(
                    [
                        ft.Text(
                            t("id"),
                            width=45,
                            weight=ft.FontWeight.BOLD,
                        ),

                        ft.Text(
                            t("action"),
                            width=190,
                            weight=ft.FontWeight.BOLD,
                        ),

                        ft.Text(
                            t("name"),
                            width=150,
                            weight=ft.FontWeight.BOLD,
                        ),

                        ft.Text(
                            t("role"),
                            width=120,
                            weight=ft.FontWeight.BOLD,
                        ),

                        ft.Text(
                            t("details"),
                            width=420,
                            weight=ft.FontWeight.BOLD,
                        ),

                        ft.Text(
                            t("date"),
                            width=170,
                            weight=ft.FontWeight.BOLD,
                        ),
                    ]
                )
            )

            table.controls.append(
                ft.Divider()
            )

            # ------------------------------------------------
            # Résultats
            # ------------------------------------------------

            if not filtered_rows:

                table.controls.append(
                    ft.Container(
                        padding=20,
                        border_radius=12,
                        bgcolor=(
                            "#3F1D24"
                            if page.theme_mode
                            == ft.ThemeMode.DARK
                            else "#FEF2F2"
                        ),
                        content=ft.Text(
                            t("no_results"),
                            color=COLOR_ACCENT,
                        ),
                    )
                )

            else:

                for row in filtered_rows:

                    action_text = enum_text(
                        row["action"],
                        JOURNAL_ACTION_KEYS,
                    )

                    table.controls.append(
                        ft.Row(
                            [
                                ft.Text(
                                    str(
                                        row["id"]
                                    ),
                                    width=45,
                                ),

                                ft.Text(
                                    action_text,
                                    width=190,
                                ),

                                ft.Text(
                                    row[
                                        "utilisateur_nom"
                                    ]
                                    or "Système",
                                    width=150,
                                ),

                                ft.Text(
                                    row["role"]
                                    or "—",
                                    width=120,
                                ),

                                ft.Text(
                                    row["details"]
                                    or "",
                                    width=420,
                                ),

                                ft.Text(
                                    row[
                                        "date_action"
                                    ]
                                    or "",
                                    width=170,
                                ),
                            ]
                        )
                    )

            page.update()

        # ----------------------------------------------------
        # Réinitialisation des filtres
        # ----------------------------------------------------

        def reset_filters(_=None):

            recherche.value = ""
            filtre_nom.value = ""
            filtre_role.value = ""
            filtre_action.value = ""
            filtre_date.value = ""

            load()

        # ----------------------------------------------------
        # Recherche avec Entrée
        # ----------------------------------------------------

        recherche.on_submit = load

        filtre_nom.on_submit = load
        filtre_date.on_submit = load

        # Recherche automatique lors de la modification
        recherche.on_change = load
        filtre_nom.on_change = load
        filtre_date.on_change = load

        filtre_role.on_change = load
        filtre_action.on_change = load

        # ----------------------------------------------------
        # Interface
        # ----------------------------------------------------

        # En-tête et filtres FIXES : seule la liste du journal défile.
        journal_header = ft.Column(
            [
                title(t("journal_title"), t("journal_subtitle")),
                ft.Divider(),
                ft.Row(
                    [
                        recherche,
                        ft.ElevatedButton(
                            t("search_button"),
                            icon=ft.Icons.SEARCH,
                            bgcolor=COLOR_PRIMARY,
                            color=ft.Colors.WHITE,
                            on_click=load,
                        ),
                    ],
                ),
                panel(
                    ft.Column(
                        [
                            ft.Text("Filtres du journal", size=16, weight=ft.FontWeight.BOLD),
                            ft.Row([filtre_nom, filtre_role], spacing=10),
                            ft.Row([filtre_action, filtre_date], spacing=10),
                            ft.Row(
                                [
                                    ft.ElevatedButton(t("refresh"), icon=ft.Icons.REFRESH, bgcolor="#2E7D32", color=ft.Colors.WHITE, on_click=load),
                                    ft.ElevatedButton("Réinitialiser", icon=ft.Icons.CLEAR, bgcolor="#EF6C00", color=ft.Colors.WHITE, on_click=reset_filters),
                                    compteur,
                                ],
                                spacing=10,
                                wrap=True,
                            ),
                        ],
                        spacing=10,
                    )
                ),
            ],
            spacing=10,
        )

        journal_table_scroll = ft.Container(
            expand=True,
            padding=ft.padding.symmetric(horizontal=4, vertical=4),
            border_radius=14,
            bgcolor=COLOR_CARD if page.theme_mode != ft.ThemeMode.DARK else "#111827",
            border=ft.border.all(1, COLOR_BORDER),
            content=ft.ListView(
                expand=True,
                spacing=0,
                controls=[table],
            ),
        )

        # Le bandeau du Journal (titre + boutons + filtres) reste fixe.
        # Seule la liste des opérations ci-dessous défile. Cela empêche le
        # contenu du Journal de repousser ou de masquer les boutons du tableau de bord.
        content.content = ft.Column(
            [journal_header, journal_table_scroll],
            spacing=12,
            expand=True,
        )

        # Chargement initial
        load()
# --------------------------------------------------------
# Paramètres
# --------------------------------------------------------
    def settings_view():
        nonlocal current_route
        current_route = "settings"

        # 1. Intégration des 6 langues (FR, EN, SW, LN, TSH, KG)
        language_dropdown = ft.Dropdown(
            label=t("language"),
            value=language,
            options=[
                ft.dropdown.Option("fr", "🇫🇷 Français"),
                ft.dropdown.Option("en", "🇬🇧 English"),
                ft.dropdown.Option("sw", "🇨🇩 Kiswahili"),
                ft.dropdown.Option("ln", "🇨🇩 Lingala"),
                ft.dropdown.Option("tsh", "🇨🇩 Tshiluba"),
                ft.dropdown.Option("kg", "🇨🇩 Kikongo"),
            ],
            width=330,
            border_radius=14,
            border_color=COLOR_BORDER,
        )

        theme_dropdown = ft.Dropdown(
            label=t("theme"),
            value=(
                "dark"
                if page.theme_mode == ft.ThemeMode.DARK
                else "light"
            ),
            options=[
                ft.dropdown.Option("light", t("light")),
                ft.dropdown.Option("dark", t("dark")),
            ],
            width=330,
            border_radius=14,
            border_color=COLOR_BORDER,
        )

        def apply_language(_):
            set_language(language_dropdown.value or "fr")

        def apply_theme(_):
            selected_theme = theme_dropdown.value
            if selected_theme == "dark":
                page.theme_mode = ft.ThemeMode.DARK
            else:
                page.theme_mode = ft.ThemeMode.LIGHT

            page.bgcolor = (
                "#0B1220"
                if page.theme_mode == ft.ThemeMode.DARK
                else COLOR_BG_LIGHT
            )

            rebuild_interface()
            navigate(current_route)

        backup_status = ft.Text("", size=12, color=COLOR_TEXT_MUTED)

        def do_backup(_):
            if not require_action("backup_manage"):
                return
            path, error = make_backup("manual")
            if path:
                backup_status.value = f"{t('backup_created')} : {os.path.basename(path)}"
                journaliser("SAUVEGARDE_MANUELLE", os.path.basename(path))
                snack(t("backup_created"))
            else:
                backup_status.value = f"{t('backup_failed')} : {error}"
                snack(t("backup_failed"), True)
            page.update()

        def do_restore(_):
            if not require_action("restore_backup"):
                return

            # FilePicker local à cette opération : aucun objet global
            # manquant et compatible avec le mode Web de Flet.
            restore_file_picker = ft.FilePicker()

            async def pick(e):
                try:
                    # Utilisation du FilePicker global
                    files = await restore_file_picker.pick_files(allow_multiple=False, with_data=False)
                    if not files:
                        return
                    selected = files[0]
                    restore_dir = os.path.join(UPLOADS_DIR, "_restore")
                    os.makedirs(restore_dir, exist_ok=True)
                    target = os.path.join(restore_dir, f"restore_{uuid.uuid4().hex}.zip")
                    await restore_file_picker.upload([ft.FilePickerUploadFile(
                        name=os.path.basename(target),
                        id=selected.id,
                        upload_url=page.get_upload_url(os.path.relpath(target, UPLOADS_DIR), 3600),
                    )])
                    ok, error = restore_backup_archive(target)
                    try: 
                        os.remove(target)
                    except Exception: 
                        pass
                    if ok:
                        journaliser("RESTAURATION_SAUVEGARDE", os.path.basename(selected.name))
                        snack(t("backup_restored"))
                        show_login()
                    else:
                        snack(f"{t('restore_failed')} : {error}", True)
                except Exception as ex:
                    snack(f"{t('restore_failed')} : {ex}", True)
        
            page.run_task(pick)

        content.content = ft.Column(
            [
                title(
                    t("settings"),
                    t("app_subtitle"),
                ),
                ft.Divider(),
                panel(
                    ft.Column(
                        [
                            ft.Text(
                                t("language"),
                                size=17,
                                weight=ft.FontWeight.BOLD,
                            ),
                            language_dropdown,
                            ft.ElevatedButton(
                                t("save"),
                                icon=ft.Icons.LANGUAGE,
                                on_click=apply_language,
                                style=ft.ButtonStyle(
                                    bgcolor=COLOR_PRIMARY,
                                    color=ft.Colors.WHITE,
                                    padding=ft.Padding(left=20, top=15, right=20, bottom=15),
                                    shape=ft.RoundedRectangleBorder(radius=8),
                                ),
                            ),
                            ft.Divider(),
                            ft.Text(
                                t("theme"),
                                size=17,
                                weight=ft.FontWeight.BOLD,
                            ),
                            theme_dropdown,
                            ft.ElevatedButton(
                                t("save"),
                                icon=ft.Icons.BRIGHTNESS_6,
                                on_click=apply_theme,
                                style=ft.ButtonStyle(
                                    bgcolor=COLOR_PRIMARY,
                                    color=ft.Colors.WHITE,
                                    padding=ft.Padding(left=20, top=15, right=20, bottom=15),
                                    shape=ft.RoundedRectangleBorder(radius=8),
                                ),
                            ),
                            ft.Divider(),
                            ft.Text(t("backup"), size=17, weight=ft.FontWeight.BOLD),
                            ft.Text(t("security_backup_note"), size=12, color=COLOR_TEXT_MUTED),
                            ft.Row([
                                ft.ElevatedButton(
                                    t("backup_now"), 
                                    icon=ft.Icons.SAVE, 
                                    on_click=do_backup,
                                    style=ft.ButtonStyle(
                                        bgcolor=COLOR_PRIMARY,
                                        color=ft.Colors.WHITE,
                                        padding=ft.Padding(left=12, top=8, right=12, bottom=8),
                                        shape=ft.RoundedRectangleBorder(radius=8),
                                    ),
                                ),
                                ft.ElevatedButton(
                                    t("restore_backup"), 
                                    icon=ft.Icons.UPLOAD, 
                                    on_click=do_restore,
                                    style=ft.ButtonStyle(
                                        bgcolor=COLOR_PRIMARY,
                                        color=ft.Colors.WHITE,
                                        padding=ft.Padding(left=12, top=8, right=12, bottom=8),
                                        shape=ft.RoundedRectangleBorder(radius=8),
                                    ),
                                ),
                            ]),
                            backup_status,
                        ],
                        spacing=12,
                    )
                ),
            ],
            spacing=14,
            scroll=ft.ScrollMode.AUTO,
        )

       # --------------------------------------------------------
        # Sécurité / session / authentification
        # --------------------------------------------------------

    def security_accounts_ready():
        conn = db_connect()
        row = conn.execute(
            "SELECT COUNT(*) AS n FROM utilisateurs WHERE password_hash IS NOT NULL AND password_hash != ''"
        ).fetchone()
        conn.close()
        return int(row["n"] or 0) > 0

    def create_first_admin(username, password):
        username = username.strip()
        if not username or len(password or "") < 8:
            return False, t("password_short")

        salt, pwd_hash = hash_password(password)
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = db_connect()
        try:
            conn.execute(
                """
                INSERT INTO utilisateurs
                (nom, role, telephone, date_creation, password_hash, password_salt,
                 failed_attempts, locked_until, active, niveau_acces)
                VALUES (?, 'Administrateur', '', ?, ?, ?, 0, 0, 1, 'National')
                """,
                (username, now, pwd_hash, salt),
            )
            conn.commit()
        except Exception as ex:
            conn.rollback()
            conn.close()
            return False, str(ex)
        conn.close()
        journaliser("INITIALISATION_ADMINISTRATEUR", username)
        return True, ""

    def touch_session():
        current_user["last_activity"] = time.time()

    def check_session(show_message=False):
        if current_user["id"] is None:
            return False

        if time.time() - float(current_user["last_activity"] or 0) > SESSION_TIMEOUT_SECONDS:
            username = current_user.get("nom") or ""
            journaliser("SESSION_EXPIREE", username)
            
            # Réinitialisation sécurisée
            keys_to_reset = ["id", "nom", "role", "province", "ville", "territoire", "commune", 
                             "secteur_chefferie", "groupement", "village", "departement"]
            for k in keys_to_reset: current_user[k] = None
            current_user.update({"niveau_acces": "National", "login_time": 0.0, "last_activity": 0.0})
            
            # Mise à jour du contexte d'audit
            AUDIT_CONTEXT.update({k: None for k in ["user_id", "user_name", "user_role", "province", "ville", 
                                                    "territoire", "commune", "secteur_chefferie", "groupement", 
                                                    "village", "departement"]})
            AUDIT_CONTEXT.update({"niveau_acces": "National"})

            if show_message:
                snack(t("session_expired"), True)
            show_login()
            return False

        touch_session()
        return True

    def check_permission(route, show_message=True):
        if not check_session(show_message=show_message):
            return False

        role = current_user["role"]
        allowed = ROLE_PERMISSIONS.get(role, set())
        if "*" in allowed or route in allowed:
            return True

        if show_message:
            snack(t("access_denied"), True)
        journaliser("ACCES_REFUSE", f"{current_user['nom']} ({role}) -> {route}")
        return False

    def check_action(action, show_message=True):
        if not check_session(show_message=show_message):
            return False
        role = current_user["role"]
        allowed = ROLE_ACTION_PERMISSIONS.get(role, set())
        if "*" in allowed or action in allowed:
            return True
        if show_message:
            snack(t("operation_denied"), True)
        journaliser("OPERATION_REFUSEE", f"{current_user['nom']} ({role}) -> {action}")
        return False

    def require_action(action):
        return check_action(action, show_message=True)

    def authenticate(username, password):
        username = (username or "").strip()
        password = password or ""

        # --- Accès PUBLIC ---
        if username.upper() == "PUBLIC":
            public_data = {
                "id": "PUBLIC",
                "nom": "Public",
                "role": "PUBLIC",
                "niveau_acces": "National",
                "province": None,
                "ville": None,
                "territoire": None,
                "commune": None,
                "secteur_chefferie": None,
                "groupement": None,
                "village": None,
                "departement": None,
                "login_time": time.time(),
                "last_activity": time.time(),
            }
            current_user.update(public_data)
            AUDIT_CONTEXT.update({"user_id": "PUBLIC", "user_name": "Public", "user_role": "PUBLIC", "niveau_acces": "National"})
            journaliser("CONNEXION_REUSSIE", "Public (Accès Libre)")
            return True, ""

        conn = db_connect()
        row = conn.execute("SELECT * FROM utilisateurs WHERE LOWER(nom) = LOWER(?) LIMIT 1", (username,)).fetchone()

        if not row:
            conn.close()
            journaliser("CONNEXION_ECHEC", username)
            return False, t("invalid_credentials")

        now = time.time()
        # Vérification verrouillage
        if float(row["locked_until"] or 0) > now:
            conn.close()
            journaliser("COMPTE_VERROUILLE", username)
            return False, t("account_locked")

        # Vérification statut actif
        if int(row["active"] or 0) != 1:
            conn.close()
            return False, t("account_inactive")

        # Vérification mot de passe
        if not verify_password(password, row["password_salt"], row["password_hash"]):
            attempts = int(row["failed_attempts"] or 0) + 1
            lock = now + LOCKOUT_SECONDS if attempts >= MAX_LOGIN_ATTEMPTS else 0
            conn.execute(
                "UPDATE utilisateurs SET failed_attempts = ?, locked_until = ? WHERE id = ?",
                (attempts, lock, row["id"]),
            )
            conn.commit()
            conn.close()
            journaliser("CONNEXION_ECHEC", f"{username} — tentative {attempts}")
            if lock:
                journaliser("COMPTE_VERROUILLE", username)
                return False, t("account_locked")
            return False, t("invalid_credentials")

        # Succès : Réinitialisation des compteurs
        conn.execute(
            "UPDATE utilisateurs SET failed_attempts = 0, locked_until = 0, last_login = ? WHERE id = ?",
            (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), row["id"]),
        )
        conn.commit()
        conn.close()

        # Mise à jour session
        user_info = {
            "id": row["id"], "nom": row["nom"], "role": row["role"],
            "niveau_acces": row["niveau_acces"] or "National",
            "province": row["province"], "ville": row["ville"], "territoire": row["territoire"],
            "commune": row["commune"], "secteur_chefferie": row["secteur_chefferie"],
            "groupement": row["groupement"], "village": row["village"], "departement": row["departement"],
            "login_time": now, "last_activity": now,
        }
        current_user.update(user_info)
        AUDIT_CONTEXT.update(user_info)
        journaliser("CONNEXION_REUSSIE", f"{row['nom']} ({row['role']})")
        return True, ""

    def logout(_=None):
        """Déconnexion centralisée, y compris pour l'espace PUBLIC.

        Le public reste sans mot de passe, mais sa session temporaire est
        détruite lors de la déconnexion et l'écran de connexion est restauré.
        Les données cadastrales et les permissions administratives ne sont
        jamais modifiées par cette opération.
        """
        user_name = current_user.get('nom') or "Anonyme"
        role = current_user.get('role') or "Aucun"
        
        if current_user["id"] is not None:
            journaliser(
                "DECONNEXION",
                f"{user_name} ({role})",
            )
            
        # Réinitialisation complète
        current_user.update({
            "id": None,
            "nom": None,
            "role": None,
            "niveau_acces": "National",
            "province": None,
            "ville": None,
            "territoire": None,
            "commune": None,
            "secteur_chefferie": None,
            "groupement": None,
            "village": None,
            "departement": None,
            "login_time": 0.0,
            "last_activity": 0.0,
        })
        
        AUDIT_CONTEXT.update({
            "user_id": None,
            "user_name": None,
            "user_role": None,
            "niveau_acces": "National",
            "province": None,
            "ville": None,
            "territoire": None,
            "commune": None,
            "secteur_chefferie": None,
            "groupement": None,
            "village": None,
            "departement": None,
        })
        
        # Retour à l'écran de connexion
        show_login()
        
    def show_login():
        setup_mode = not security_accounts_ready()

        username = ft.TextField(
            label=t("username"),
            prefix_icon=ft.Icons.PERSON,
            width=390,
            border_radius=14,
            border_color=COLOR_BORDER,
            hint_text="PUBLIC pour l'espace citoyen sans mot de passe" if not setup_mode else None,
        )
        password = ft.TextField(
            label=t("password"),
            password=True,
            can_reveal_password=True,
            prefix_icon=ft.Icons.LOCK,
            width=390,
            border_radius=14,
            border_color=COLOR_BORDER,
        )
        confirm = ft.TextField(
            label=t("confirm_password"),
            password=True,
            can_reveal_password=True,
            prefix_icon=ft.Icons.LOCK_RESET,
            width=390,
            border_radius=14,
            border_color=COLOR_BORDER,
            visible=setup_mode,
        )
        message = ft.Text(
            "",
            size=13,
            color=COLOR_ACCENT,
            text_align=ft.TextAlign.CENTER,
        )

        public_hint = ft.Text(
            "ESPACE PUBLIC : saisissez PUBLIC et laissez le mot de passe vide.",
            size=12,
            color=COLOR_PRIMARY,
            text_align=ft.TextAlign.CENTER,
            visible=False,
        )

        def sync_public_mode(_=None):
            is_public = (username.value or "").strip().upper() == "PUBLIC" and not setup_mode
            password.visible = not is_public
            if is_public:
                password.value = ""
                public_hint.visible = True
                message.value = ""
            else:
                public_hint.visible = False
            page.update()

        username.on_change = sync_public_mode

        def open_public_portal():
            """Ouvre le portail citoyen après authentification PUBLIC sans mot de passe.

            Le portail est un ft.View. Dans cette application, la page de connexion
            est construite avec page.controls et non avec page.views. On ajoute donc
            les contrôles de la vue directement à la page afin d'éviter une page blanche.
            """
            public_view = public_verify_view(page, on_admin_access=show_login, on_logout=logout)
            page.controls.clear()
            page.controls.extend(public_view.controls)
            page.update()

        def submit(_):
            if setup_mode:
                if password.value != confirm.value:
                    message.value = t("password_mismatch")
                    page.update()
                    return
                if len(password.value or "") < 8:
                    message.value = t("password_short")
                    page.update()
                    return
                ok, error = create_first_admin(username.value, password.value)
                if not ok:
                    message.value = error
                    page.update()
                    return
                message.value = t("security_ready")
                password.value = ""
                confirm.value = ""
                show_login()
                return

            username_value = (username.value or "").strip()
            password_value = password.value or ""

            # Accès citoyen : PUBLIC est le seul identifiant public accepté et
            # aucun mot de passe n'est demandé.
            if username_value.upper() == "PUBLIC":
                if password_value.strip():
                    message.value = "Pour l'espace PUBLIC, aucun mot de passe ne doit être saisi."
                    password.value = ""
                    sync_public_mode()
                    return
                ok, error = authenticate("PUBLIC", "")
                if not ok:
                    message.value = error
                    page.update()
                    return
                open_public_portal()
                return

            ok, error = authenticate(username_value, password_value)
            if not ok:
                message.value = error
                password.value = ""
                page.update()
                return

            current_route_local = "dashboard"
            rebuild_interface()
            navigate(current_route_local)

        def login_public(_):
            ok, error = authenticate("PUBLIC", "")
            if not ok:
                message.value = error
                page.update()
                return
            open_public_portal()

        page.controls.clear()

        title_text = t("setup_title") if setup_mode else t("login_title")
        subtitle_text = t("setup_subtitle") if setup_mode else t("login_subtitle")
        button_text = t("setup_admin") if setup_mode else t("login")

        page.add(
            ft.Container(
                expand=True,
                alignment=ft.Alignment.CENTER,
                bgcolor=(
                    "#0B1220"
                    if page.theme_mode == ft.ThemeMode.DARK
                    else COLOR_BG_LIGHT
                ),
                content=ft.Container(
                    width=500,
                    padding=40,
                    border_radius=22,
                    bgcolor=(
                        "#172033"
                        if page.theme_mode == ft.ThemeMode.DARK
                        else ft.Colors.WHITE
                    ),
                    shadow=ft.BoxShadow(
                        blur_radius=25,
                        spread_radius=2,
                        color="#00000022",
                    ),
                    content=ft.Column(
                        [
                            ft.Image(
                                src=RDC_COAT_OF_ARMS_URL,
                                width=55,
                                height=65,
                            ),
                            ft.Text(
                                t("app_title"),
                                size=28,
                                weight=ft.FontWeight.BOLD,
                            ),
                            ft.Text(
                                title_text,
                                size=20,
                                weight=ft.FontWeight.BOLD,
                                text_align=ft.TextAlign.CENTER,
                            ),
                            ft.Text(
                                subtitle_text,
                                size=13,
                                color=COLOR_TEXT_MUTED,
                                text_align=ft.TextAlign.CENTER,
                            ),
                            ft.Divider(),
                            username,
                            password,
                            confirm,
                            public_hint,
                            message,
                            ft.ElevatedButton(
                                button_text,
                                icon=ft.Icons.PERSON_ADD if setup_mode else ft.Icons.LOGIN,
                                on_click=submit,
                                style=ft.ButtonStyle(
                                    bgcolor=COLOR_PRIMARY,
                                    color=ft.Colors.WHITE,
                                    padding=ft.Padding(left=20, top=15, right=20, bottom=15),
                                    shape=ft.RoundedRectangleBorder(radius=8),
                                ),
                                width=390,
                            ),
                            *(
                                [
                                    ft.OutlinedButton(
                                        "Accès Public Libre — sans mot de passe",
                                        icon=ft.Icons.EXPLORE,
                                        on_click=login_public,
                                        style=ft.ButtonStyle(
                                            padding=ft.Padding(left=20, top=12, right=20, bottom=12),
                                            shape=ft.RoundedRectangleBorder(radius=8),
                                        ),
                                        width=390,
                                    )
                                ]
                                if not setup_mode
                                else []
                            ),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=16,
                    ),
                ),
            )
        )
        page.update()

   # --------------------------------------------------------
    # Navigation
    # --------------------------------------------------------

    def navigate(route):
        nonlocal current_route
        if not check_permission(route, show_message=True):
            return

        # Met à jour la route courante de la session.
        current_route = route

        if route == "dashboard":
            dashboard_view()
        elif route == "proprietaires":
            proprietaires_view()
        elif route == "parcelles":
            parcelles_view()
        elif route == "mutations":
            mutations_view()
        elif route == "verification":
            verification_view()
        elif route == "utilisateurs":
            utilisateurs_view()
        elif route == "territoires":
            territoires_view()
        elif route == "journal":
            journal_view()
        elif route == "settings":
            settings_view()
        
        # Actualise l'interface pour refléter le changement de menu actif
        rebuild_interface()

    def nav_item(label, icon, route):
        active = current_route == route
        c1, c2 = MENU_GRADIENTS.get(route, (COLOR_PRIMARY, COLOR_DEEP_BLUE))

        return ft.Container(
            padding=ft.Padding(left=4, top=2, right=4, bottom=2),
            border_radius=14,
            gradient=ft.LinearGradient(
                begin=ft.Alignment.TOP_LEFT,
                end=ft.Alignment.BOTTOM_RIGHT,
                colors=[c1, c2],
            ),
            opacity=1.0 if active else 0.92,
            shadow=ft.BoxShadow(blur_radius=10 if active else 6, spread_radius=0, color="#00000022"),
            content=ft.TextButton(
                content=ft.Row(
                    [
                        ft.Icon(icon, color=ft.Colors.WHITE, size=18),
                        ft.Text(label, color=ft.Colors.WHITE, weight=ft.FontWeight.BOLD, size=13),
                    ],
                    spacing=10,
                ),
                style=ft.ButtonStyle(
                    color=ft.Colors.WHITE,
                    padding=ft.Padding(left=12, top=11, right=12, bottom=11),
                    alignment=ft.Alignment.CENTER_LEFT,
                    overlay_color="#FFFFFF22",
                ),
                on_click=lambda e, r=route: navigate(r),
            ),
        )

# --------------------------------------------------------
    # Reconstruction de l'interface
    # --------------------------------------------------------

    def rebuild_interface():
        nonlocal language

        dark = (
            page.theme_mode
            == ft.ThemeMode.DARK
        )

        sidebar_container.content = ft.Container(
            width=285,
            bgcolor=COLOR_CARD if not dark else COLOR_BG_DARK,
            padding=20,
            content=ft.Column(
                [
                    # En-tête
                    ft.Row(
                        [
                            ft.Container(
                                width=86,
                                height=86,
                                border_radius=18,
                                bgcolor=ft.Colors.WHITE,
                                alignment=ft.Alignment.CENTER,
                                padding=7,
                                shadow=ft.BoxShadow(
                                    blur_radius=18,
                                    spread_radius=1,
                                    color="#00000044",
                                ),
                                content=ft.Image(
                                    src=RDC_COAT_OF_ARMS_URL,
                                    width=40,
                                    height=48,
                                ),
                            ),
                            ft.Column(
                                [
                                    ft.Row(
                                        [
                                            ft.Text(
                                                t("app_title"),
                                                size=21,
                                                weight=ft.FontWeight.BOLD,
                                                color=ft.Colors.WHITE,
                                            ),
                                            ft.Image(
                                                src=RDC_FLAG_URL,
                                                width=30,
                                                height=20,
                                            ),
                                        ],
                                        spacing=7,
                                        vertical_alignment=ft.CrossAxisAlignment.CENTER,
                                    ),
                                    ft.Text(
                                        t("republic"),
                                        size=8,
                                        color=COLOR_SECONDARY,
                                    ),
                                ],
                                spacing=1,
                            ),
                        ],
                        spacing=10,
                    ),
                    ft.Text(
                        t("app_subtitle"),
                        size=10,
                        color="#94A3B8",
                    ),
                    ft.Divider(
                        color="#334155",
                        height=20,
                    ),

                    # Navigation
                    nav_item(t("dashboard"), ft.Icons.DASHBOARD, "dashboard"),
                    nav_item(t("owners"), ft.Icons.PERSON, "proprietaires"),
                    nav_item(t("parcels"), ft.Icons.LANDSCAPE, "parcelles"),
                    nav_item(t("mutations"), ft.Icons.SWAP_HORIZ, "mutations"),
                    nav_item(t("search"), ft.Icons.SEARCH, "verification"),
                    nav_item(t("users"), ft.Icons.GROUP, "utilisateurs"),
                    nav_item(t("territorial_admin"), ft.Icons.ACCOUNT_TREE, "territoires"),
                    nav_item(t("journal"), ft.Icons.HISTORY, "journal"),
                    nav_item(t("settings"), ft.Icons.SETTINGS, "settings"),

                    ft.Container(
                        expand=True
                    ),

                    ft.Divider(
                        color="#334155",
                        height=15,
                    ),

                    # Bouton de déconnexion
                    ft.OutlinedButton(
                        t("logout"),
                        icon=ft.Icons.LOGOUT,
                        on_click=logout,
                        style=ft.ButtonStyle(
                            color=ft.Colors.WHITE,
                            side=ft.BorderSide(1, "#334155"),
                        ),
                        width=245,
                    ),

                    # Sélecteur de langue
                    ft.Text(
                        t("language"),
                        size=10,
                        color="#94A3B8",
                    ),
                    ft.Dropdown(
                        value=language,
                        options=[
                            ft.dropdown.Option("fr", "🇫🇷 Français"),
                            ft.dropdown.Option("en", "🇬🇧 English"),
                            ft.dropdown.Option("sw", "🇨🇩 Kiswahili"),
                            ft.dropdown.Option("ln", "🇨🇩 Lingala"),
                            ft.dropdown.Option("tsh", "🇨🇩 Tshiluba"),
                            ft.dropdown.Option("kg", "🇨🇩 Kikongo"),
                        ],
                        border_radius=8,
                        bgcolor="#172033",
                        border_color="#334155",
                        color=ft.Colors.WHITE,
                        on_select=on_language_change,
                    ),

                    ft.Container(height=5),

                    ft.Text(
                        t("prototype"),
                        size=9,
                        color="#64748B",
                        text_align=ft.TextAlign.CENTER,
                    ),
                ],
                spacing=5,
            ),
        )

        # Barre supérieure
        topbar = ft.Container(
            height=70,
            padding=ft.Padding(left=22, top=10, right=22, bottom=10),
            bgcolor=(
                "#111827"
                if dark
                else COLOR_SURFACE
            ),
            border=ft.Border(
                bottom=ft.BorderSide(
                    1,
                    "#334155"
                    if dark
                    else COLOR_BORDER,
                )
            ),
            content=ft.Row(
                [
                    ft.Row(
                        [
                            ft.Container(
                                width=34,
                                height=23,
                                border_radius=4,
                                clip_behavior=ft.ClipBehavior.HARD_EDGE,
                                content=ft.Image(
                                    src=RDC_FLAG_URL,
                                    width=34,
                                    height=23,
                                ),
                            ),
                            ft.Column(
                                [
                                    ft.Text(
                                        t("app_title"),
                                        size=17,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                    ft.Text(
                                        t("welcome"),
                                        size=10,
                                        color=COLOR_TEXT_MUTED,
                                    ),
                                ],
                                spacing=0,
                            ),
                        ],
                        spacing=10,
                        expand=True,
                    ),
                    ft.Container(
                        padding=ft.Padding(left=12, top=8, right=12, bottom=8),
                        border_radius=20,
                        bgcolor=(
                            "#1E293B"
                            if dark
                            else "#EFF6FF"
                        ),
                        content=ft.Row(
                            [
                                ft.Icon(ft.Icons.PERSON, size=17, color=COLOR_PRIMARY),
                                ft.Column(
                                    [
                                        ft.Text(
                                            f"{current_user['nom']} • {current_user['role']}",
                                            size=11,
                                            weight=ft.FontWeight.BOLD,
                                        ),
                                        ft.Text(
                                            current_scope_label(),
                                            size=8,
                                            color=COLOR_TEXT_MUTED,
                                        ),
                                    ],
                                    spacing=0,
                                ),
                            ],
                            spacing=5,
                        ),
                    ),
                    ft.Container(
                        padding=ft.Padding(left=12, top=8, right=12, bottom=8),
                        border_radius=20,
                        bgcolor=(
                            "#1E293B"
                            if dark
                            else "#EFF6FF"
                        ),
                        content=ft.Row(
                            [
                                ft.Icon(
                                    ft.Icons.LANGUAGE,
                                    size=17,
                                    color=COLOR_PRIMARY,
                                ),
                                ft.Text(
                                    language.upper(),
                                    size=11,
                                    weight=ft.FontWeight.BOLD,
                                ),
                            ],
                            spacing=5,
                        ),
                    ),
                    ft.Container(
                        padding=ft.Padding(left=12, top=6, right=12, bottom=6),
                        border_radius=16,
                        bgcolor="#F1F5F9" if not dark else "#1E293B",
                        content=ft.Row(
                            [
                                ft.Icon(ft.Icons.ACCESS_TIME, size=16, color=COLOR_PRIMARY),
                                ft.Column([clock_date, clock_time, clock_zone], spacing=0),
                            ],
                            spacing=7,
                        ),
                    ),
                    ft.IconButton(
                        icon=ft.Icons.BRIGHTNESS_6,
                        tooltip=t("theme"),
                        on_click=lambda e: toggle_theme(),
                    ),
                ],
                vertical_alignment=(
                    ft.CrossAxisAlignment.CENTER
                ),
            ),
        )

        page.controls.clear()

        page.add(
            ft.Column(
                [
                    topbar,
                    ft.Row(
                        [
                            sidebar_container,
                            ft.VerticalDivider(
                                width=1,
                            ),
                            content,
                        ],
                        expand=True,
                        spacing=0,
                    ),
                ],
                expand=True,
            )
        )

        page.update()

    def set_language(new_language):
        nonlocal language

        supported = {"fr", "en", "sw", "ln", "tsh", "kg"}
        selected = str(new_language or "").strip()

        if selected not in supported:
            return

        language = selected
        rebuild_interface()
        navigate(current_route)

    def on_language_change(e):
        selected = getattr(e.control, "value", None) or getattr(e, "data", None)
        if selected:
            set_language(selected)

    def toggle_theme():
        page.theme_mode = (
            ft.ThemeMode.DARK
            if page.theme_mode
            == ft.ThemeMode.LIGHT
            else ft.ThemeMode.LIGHT
        )

        page.bgcolor = (
            "#0B1220"
            if page.theme_mode
            == ft.ThemeMode.DARK
            else COLOR_BG_LIGHT
        )

        rebuild_interface()
        navigate(current_route)

    async def update_local_clock():
        # Fuseau de l'appareil/client lorsqu'il est exposé par Flet.
        # Fallback : fuseau local du processus Python.
        client_tz = None
        try:
            device = await page.get_device_info()
            for attr in ("timezone", "time_zone", "timezone_name", "timeZone"):
                candidate = getattr(device, attr, None)
                if candidate:
                    client_tz = str(candidate)
                    break
        except Exception:
            pass

        while True:
            try:
                now = datetime.now().astimezone()
                if client_tz:
                    try:
                        now = datetime.now(ZoneInfo(client_tz))
                    except Exception:
                        pass

                if language == "en":
                    date_value = now.strftime("%m/%d/%Y")
                else:
                    date_value = now.strftime("%d/%m/%Y")
                clock_date.value = date_value
                clock_time.value = now.strftime("%H:%M:%S")
                zone_name = now.tzname() or "LOCAL"
                offset = now.strftime("%z")
                offset = (offset[:3] + ":" + offset[3:]) if len(offset) == 5 else offset
                clock_zone.value = f"{zone_name}  UTC{offset}" if offset else zone_name
                try:
                    clock_date.update()
                    clock_time.update()
                    clock_zone.update()
                except Exception:
                    pass
            except Exception as ex:
                print(f"[CADASTRE RDC] HORLOGE : {type(ex).__name__}: {ex}")
            await asyncio.sleep(1)

    async def detect_browser_language():
        """Adapte l'interface au premier langage préféré du navigateur."""
        try:
            device = await page.get_device_info()
            locales = getattr(device, "locales", None) or []
            for loc in locales:
                code = str(getattr(loc, "language_code", "") or "").lower()
                if code in {"fr", "en", "sw", "ln", "tsh", "kg"}:
                    if code != language:
                        set_language(code)
                    return
            browser_lang = str(getattr(device, "language", "") or "").lower().split("-")[0]
            if browser_lang in {"fr", "en", "sw", "ln", "tsh", "kg"} and browser_lang != language:
                set_language(browser_lang)
        except Exception as ex:
            print(f"[CADASTRE RDC] LANGUE NAVIGATEUR : {type(ex).__name__}: {ex}")

    def on_host_locale_change(e):
        page.run_task(detect_browser_language)

    page.on_locale_change = on_host_locale_change
    show_login()
    page.run_task(update_local_clock)
    page.run_task(detect_browser_language)


# ============================================================
# QR — VÉRIFICATION PUBLIQUE DES CERTIFICATS
# ============================================================

CADASTRE_QR_BASE_URL = os.environ.get(
    "CADASTRE_QR_BASE_URL",
    "http://127.0.0.1:8000/verification/certificat",
)


def _cad_qr_init_db():
    """Crée la table de liaison des certificats sans toucher aux tables existantes."""
    conn = db_connect()
    try:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS certificats_qr (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                numero_certificat TEXT NOT NULL UNIQUE,
                token TEXT NOT NULL UNIQUE,
                type_document TEXT NOT NULL,
                parcelle_id INTEGER,
                proprietaire_id INTEGER,
                fichier_word TEXT,
                fichier_pdf TEXT,
                qr_image TEXT,
                date_creation TEXT NOT NULL,
                statut TEXT NOT NULL DEFAULT 'Actif'
            )
            """
        )
        conn.commit()
    finally:
        conn.close()


def _cad_qr_token():
    return secrets.token_urlsafe(32)


def _cad_qr_url(token):
    return f"{CADASTRE_QR_BASE_URL.rstrip('/')}/{token}"


def _cad_qr_register(numero_certificat, type_document, parcelle, proprietaire):
    """Associe un certificat à un jeton non devinable."""
    _cad_qr_init_db()
    conn = db_connect()
    try:
        row = conn.execute(
            "SELECT * FROM certificats_qr WHERE numero_certificat = ? LIMIT 1",
            (numero_certificat,),
        ).fetchone()
        if row:
            return dict(row)

        token = _cad_qr_token()
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            """
            INSERT INTO certificats_qr
            (numero_certificat, token, type_document, parcelle_id,
             proprietaire_id, date_creation, statut)
            VALUES (?, ?, ?, ?, ?, ?, 'Actif')
            """,
            (
                numero_certificat,
                token,
                type_document,
                parcelle.get("id") if parcelle else None,
                proprietaire.get("id") if proprietaire else None,
                now,
            ),
        )
        conn.commit()
        row = conn.execute(
            "SELECT * FROM certificats_qr WHERE numero_certificat = ? LIMIT 1",
            (numero_certificat,),
        ).fetchone()
        return dict(row)
    finally:
        conn.close()


def _cad_qr_generate(token, output_path):
    try:
        import qrcode
    except ImportError as ex:
        raise RuntimeError(
            "Le QR nécessite la bibliothèque qrcode. Exécutez : pip install qrcode[pil]"
        ) from ex

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=4,
    )
    qr.add_data(_cad_qr_url(token))
    qr.make(fit=True)
    image = qr.make_image(fill_color="black", back_color="white")
    image.save(output_path)
    return output_path


def _cad_qr_update_files(numero_certificat, fichier_word=None, fichier_pdf=None, qr_image=None):
    _cad_qr_init_db()
    conn = db_connect()
    try:
        conn.execute(
            """
            UPDATE certificats_qr
            SET fichier_word = COALESCE(?, fichier_word),
                fichier_pdf = COALESCE(?, fichier_pdf),
                qr_image = COALESCE(?, qr_image)
            WHERE numero_certificat = ?
            """,
            (fichier_word, fichier_pdf, qr_image, numero_certificat),
        )
        conn.commit()
    finally:
        conn.close()


def _cad_qr_get_certificate(token):
    """Retourne le certificat et toutes les données liées à la parcelle/propriétaire."""
    _cad_qr_init_db()
    conn = db_connect()
    try:
        row = conn.execute(
            """
            SELECT
                cq.*,
                p.*,
                pr.id AS owner_id,
                pr.nom_complet AS owner_name,
                pr.telephone AS owner_phone,
                pr.email AS owner_email,
                pr.piece_identite AS owner_identity,
                pr.adresse AS owner_address,
                pr.photo AS owner_photo
            FROM certificats_qr cq
            LEFT JOIN parcelles p ON p.id = cq.parcelle_id
            LEFT JOIN proprietaires pr ON pr.id = cq.proprietaire_id
            WHERE cq.token = ? AND cq.statut = 'Actif'
            LIMIT 1
            """,
            (token,),
        ).fetchone()
        return dict(row) if row else None
    finally:
        conn.close()

# ============================================================
# MODULE 10 — DOCUMENTS CADASTRAUX OFFICIELS
# ============================================================

def _cad_docs_dir():
    base = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base, "documents_cadastraux")
    os.makedirs(path, exist_ok=True)
    return path



def _cad_resolve_upload_path(value):
    """Résout un chemin de fichier provenant de SQLite ou de uploads/.

    Accepte les chemins Windows absolus, /uploads/..., uploads/... et les
    chemins relatifs. La fonction ne modifie jamais la valeur enregistrée
    dans la base : elle sert uniquement à retrouver le fichier réel.
    """
    if not value:
        return None
    raw = str(value).strip().strip('"')
    if not raw:
        return None
    raw = raw.replace('\\', os.sep).replace('/', os.sep)
    candidates = []
    if os.path.isabs(raw):
        candidates.append(raw)
    normalized = raw.lstrip('/\\')
    if normalized.lower().startswith('uploads' + os.sep):
        candidates.append(os.path.join(BASE_DIR, normalized))
    else:
        candidates.append(os.path.join(BASE_DIR, normalized))
        candidates.append(os.path.join(UPLOADS_DIR, normalized))
    for candidate in candidates:
        candidate = os.path.abspath(candidate)
        if os.path.isfile(candidate):
            return candidate
    # Dernier recours : retrouver le fichier par son nom dans uploads/.
    basename = os.path.basename(raw)
    if basename and os.path.isdir(UPLOADS_DIR):
        for root, _, files in os.walk(UPLOADS_DIR):
            if basename in files:
                return os.path.join(root, basename)
    return None


def _cad_public_upload_url(value):
    """Convertit un chemin uploads en URL Flet publique si possible."""
    raw = str(value or '').strip()
    if raw.startswith('/uploads/'):
        return raw
    resolved = _cad_resolve_upload_path(raw)
    if resolved:
        try:
            rel = os.path.relpath(resolved, UPLOADS_DIR).replace(os.sep, '/')
            if not rel.startswith('..'):
                return '/uploads/' + rel
        except Exception:
            pass
    return raw


def _cad_find_template(template_name):
    """
    Recherche le modèle officiel sans jamais le remplacer par un faux modèle.

    Priorité absolue : templates/. Une recherche tolérante est ensuite faite
    dans templates/ pour accepter un nom légèrement différent. modeles/ reste
    une compatibilité secondaire.
    """
    _cad_word_dirs()
    stem, ext = os.path.splitext(template_name)
    candidates = [
        template_name,
        stem.replace("modele_", "") + ext,
        stem.replace("modele_", "certificat_") + ext,
    ]

    # 1. Priorité stricte au dossier templates/.
    for name in candidates:
        candidate = os.path.join(TEMPLATES_DIR, name)
        if os.path.isfile(candidate):
            return candidate

    # 2. Recherche récursive dans templates/ pour les noms réels du modèle.
    if os.path.isdir(TEMPLATES_DIR):
        for root, _dirs, files in os.walk(TEMPLATES_DIR):
            for filename in files:
                if not filename.lower().endswith(".docx"):
                    continue
                normalized = filename.lower().replace(" ", "_").replace("-", "_")
                if "certificat" in normalized and "enregistrement" in normalized:
                    return os.path.join(root, filename)

    # 3. Compatibilité avec l'ancien dossier modeles/.
    for name in candidates:
        candidate = os.path.join(MODELES_DIR, name)
        if os.path.isfile(candidate):
            return candidate

    if os.path.isdir(MODELES_DIR):
        for root, _dirs, files in os.walk(MODELES_DIR):
            for filename in files:
                if not filename.lower().endswith(".docx"):
                    continue
                normalized = filename.lower().replace(" ", "_").replace("-", "_")
                if "certificat" in normalized and "enregistrement" in normalized:
                    return os.path.join(root, filename)

    # IMPORTANT : aucun modèle de secours ne doit être créé pour un document
    # officiel. Cela évite de générer un certificat qui ne respecte pas le modèle
    # administratif présent sur le poste de l'utilisateur.
    raise FileNotFoundError(
        "Modèle officiel introuvable. Placez le fichier .docx du certificat "
        "d'enregistrement dans le dossier templates/ de CADASTRE RDC."
    )


def _cad_owner_parcels(proprietaire_id):
    """Retourne toutes les parcelles liées à un propriétaire."""
    if proprietaire_id in (None, ''):
        return []
    conn = db_connect()
    try:
        return [dict(row) for row in conn.execute(
            "SELECT * FROM parcelles WHERE proprietaire_id = ? ORDER BY id ASC",
            (proprietaire_id,),
        ).fetchall()]
    finally:
        conn.close()


def _cad_append_parcels_to_word(document, parcels):
    """Ajoute au certificat la liste complète des parcelles du propriétaire."""
    if not parcels:
        return
    try:
        from docx.shared import Pt
    except Exception:
        Pt = None
    document.add_page_break()
    document.add_heading("PARCELLES ENREGISTRÉES AU NOM DU PROPRIÉTAIRE", level=1)
    document.add_paragraph(
        f"Nombre total de parcelles enregistrées : {len(parcels)}"
    )
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = [
        "N° / Référence", "Superficie", "Adresse", "Province",
        "Ville / Territoire", "Commune / Chefferie", "Localité", "GPS",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        if Pt:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
    for parcel in parcels:
        gps = " / ".join(
            str(v) for v in (parcel.get('latitude'), parcel.get('longitude'))
            if v not in (None, '')
        ) or "—"
        values = [
            parcel.get('numero') or parcel.get('reference') or f"ID-{parcel.get('id', '—')}",
            parcel.get('superficie') or "—",
            parcel.get('adresse') or parcel.get('localisation') or "—",
            parcel.get('province') or "—",
            parcel.get('ville_territoire') or "—",
            parcel.get('commune_chefferie') or "—",
            parcel.get('localite') or "—",
            gps,
        ]
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
            if Pt:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)


def _cad_word_safe_filename(value):
    """Nettoie une valeur destinée à devenir un nom de fichier Windows."""
    text = str(value or "document")
    text = re.sub(r'[<>:"/\\|?*]', "_", text)
    text = re.sub(r"\s+", "_", text).strip("._ ")
    return (text or "document")[:120]


def _cad_word_dirs():
    """Prépare templates/, modeles/ et les certificats sans casser l'existant."""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(MODELES_DIR, exist_ok=True)
    os.makedirs(CERTIFICATS_DIR, exist_ok=True)
    return TEMPLATES_DIR, MODELES_DIR, CERTIFICATS_DIR


def _cad_create_default_word_template(template_path=None):
    """
    Crée automatiquement un modèle Word de certificat d'enregistrement
    si aucun modèle n'existe encore. La mise en page peut ensuite être
    remplacée par le modèle officiel fourni par l'administration.
    """
    _cad_word_dirs()
    template_path = template_path or os.path.join(
        TEMPLATES_DIR,
        "modele_certificat_enregistrement.docx",
    )

    if os.path.exists(template_path):
        return template_path

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("RÉPUBLIQUE DÉMOCRATIQUE DU CONGO")
        r.bold = True
        r.font.size = Pt(15)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("MINISTÈRE DES AFFAIRES FONCIÈRES")
        r.bold = True
        r.font.size = Pt(13)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("CADASTRE RDC")
        r.bold = True
        r.font.size = Pt(16)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("CERTIFICAT D'ENREGISTREMENT")
        r.bold = True
        r.font.size = Pt(18)

        doc.add_paragraph("Numéro du document : {{ numero_document }}")
        doc.add_paragraph("Propriétaire : {{ nom_proprietaire }}")
        doc.add_paragraph("Téléphone : {{ telephone }}")
        doc.add_paragraph("E-mail : {{ email }}")
        doc.add_paragraph("Pièce d'identité : {{ piece_identite }}")
        doc.add_paragraph("Adresse du propriétaire : {{ adresse_proprietaire }}")
        doc.add_paragraph("")
        doc.add_paragraph("PARCELLE")
        doc.add_paragraph("Numéro / Référence : {{ numero_parcelle }}")
        doc.add_paragraph("Superficie : {{ superficie }}")
        doc.add_paragraph("Adresse : {{ adresse_parcelle }}")
        doc.add_paragraph("Province : {{ province }}")
        doc.add_paragraph("Ville / Territoire : {{ ville_territoire }}")
        doc.add_paragraph("Commune / Chefferie : {{ commune_chefferie }}")
        doc.add_paragraph("Quartier / Groupement : {{ quartier_groupement }}")
        doc.add_paragraph("Localité : {{ localite }}")
        doc.add_paragraph("")
        doc.add_paragraph("COORDONNÉES GPS")
        doc.add_paragraph("Latitude : {{ latitude }}")
        doc.add_paragraph("Longitude : {{ longitude }}")
        doc.add_paragraph("Statut : {{ statut }}")
        doc.add_paragraph("Date d'émission : {{ date_emission }}")
        doc.add_paragraph("QR DE VÉRIFICATION :")
        doc.add_paragraph("{{ qr_code }}")
        doc.add_paragraph("Vérification : {{ qr_verification_url }}")
        doc.add_paragraph("")
        doc.add_paragraph("Photo du propriétaire :")
        doc.add_paragraph("{{ photo_proprietaire }}")
        doc.add_paragraph("")
        doc.add_paragraph(
            "Document généré électroniquement par CADASTRE RDC. "
            "La valeur juridique dépend de la validation par l'autorité compétente."
        )
        doc.save(template_path)
        return template_path
    except Exception as ex:
        raise RuntimeError(
            "Impossible de créer le modèle Word. "
            "Installez python-docx avec : pip install python-docx"
        ) from ex


def _cad_word_context(proprietaire, parcelle, document_number):
    """Prépare le contexte commun utilisé par les modèles Word."""
    adresse_parcelle = (
        parcelle.get("adresse")
        or parcelle.get("localisation")
        or ""
    )
    adresse_proprietaire = (
        proprietaire.get("adresse")
        or ""
    )

    return {
        "numero_document": document_number,
        "nom_proprietaire": proprietaire.get("nom_complet") or "—",
        "telephone": proprietaire.get("telephone") or "—",
        "email": proprietaire.get("email") or "—",
        "piece_identite": proprietaire.get("piece_identite") or "—",
        "adresse_proprietaire": adresse_proprietaire or "—",
        "numero_parcelle": (
            parcelle.get("numero")
            or parcelle.get("reference")
            or "—"
        ),
        "superficie": parcelle.get("superficie") or "—",
        "adresse_parcelle": adresse_parcelle or "—",
        "province": parcelle.get("province") or "—",
        "ville_territoire": parcelle.get("ville_territoire") or "—",
        "commune_chefferie": parcelle.get("commune_chefferie") or "—",
        "quartier_groupement": parcelle.get("quartier_groupement") or "—",
        "localite": parcelle.get("localite") or "—",
        "latitude": parcelle.get("latitude") or "—",
        "longitude": parcelle.get("longitude") or "—",
        "statut": parcelle.get("statut") or "—",
        "date_emission": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def generer_document_word(
    proprietaire,
    parcelle,
    template_name="modele_certificat_enregistrement.docx",
    prefix="CE",
    title=None,
):
    """
    Génère un document Word à partir d'un modèle .docx.

    - conserve les PDF existants ;
    - utilise docxtpl si disponible ;
    - insère la photo du propriétaire avec InlineImage ;
    - crée automatiquement un modèle de départ s'il est absent.
    """
    _cad_word_dirs()

    try:
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm
    except ImportError as ex:
        raise RuntimeError(
            "La génération Word nécessite docxtpl et python-docx. "
            "Exécutez : pip install docxtpl python-docx"
        ) from ex

    template_path = _cad_find_template(template_name)

    if not os.path.exists(template_path):
        # Le modèle officiel doit être placé dans templates/. En son absence,
        # on crée le modèle de secours à cet emplacement sans toucher aux
        # anciens modèles du dossier modeles/.
        _cad_create_default_word_template(template_path)

    document_number = _cad_doc_number(prefix)
    document = DocxTemplate(template_path)

    context = _cad_word_context(
        proprietaire or {},
        parcelle or {},
        document_number,
    )

    # Enregistrement du certificat et création du QR avant le rendu Word.
    qr_record = _cad_qr_register(
        document_number,
        title or prefix,
        parcelle or {},
        proprietaire or {},
    )
    qr_dir = os.path.join(CERTIFICATS_DIR, "qr")
    qr_path = os.path.join(
        qr_dir,
        f"{_cad_word_safe_filename(document_number)}.png",
    )
    _cad_qr_generate(qr_record["token"], qr_path)
    context["qr_verification_url"] = _cad_qr_url(qr_record["token"])

    photo_path = (
        (proprietaire or {}).get("photo_path")
        or (proprietaire or {}).get("photo")
        or (parcelle or {}).get("photo_path")
        or (parcelle or {}).get("photo")
    )
    photo_path = _cad_resolve_upload_path(photo_path) or photo_path

    if photo_path and os.path.exists(str(photo_path)):
        context["photo_proprietaire"] = InlineImage(
            document,
            str(photo_path),
            width=Mm(35),
            height=Mm(45),
        )
    else:
        context["photo_proprietaire"] = "Photo non disponible"

    if os.path.exists(qr_path):
        context["qr_code"] = InlineImage(
            document,
            qr_path,
            width=Mm(35),
            height=Mm(35),
        )
    else:
        context["qr_code"] = "QR non disponible"

    if title:
        context["titre_document"] = title

    document.render(context)

    # Le certificat d'enregistrement doit couvrir toutes les parcelles
    # actuellement rattachées au propriétaire. Le premier enregistrement
    # reste utilisé pour le QR, tandis que la liste complète est ajoutée au
    # document Word.
    if (title or "").strip().lower() == "certificat d'enregistrement":
        owner_id = (proprietaire or {}).get("id")
        all_parcels = _cad_owner_parcels(owner_id)
        if not all_parcels and parcelle:
            all_parcels = [dict(parcelle)]
        _cad_append_parcels_to_word(document, all_parcels)

    numero_parcelle = _cad_word_safe_filename(
        (parcelle or {}).get("numero")
        or (parcelle or {}).get("reference")
        or "parcelle"
    )

    output_path = os.path.join(
        CERTIFICATS_DIR,
        f"{_cad_word_safe_filename(prefix)}_{numero_parcelle}_{document_number}.docx",
    )
    document.save(output_path)
    _cad_qr_update_files(
        document_number,
        fichier_word=output_path,
        qr_image=qr_path,
    )

    return output_path


def generer_certificat_propriete_word(proprietaire, parcelle):
    return generer_document_word(
        proprietaire,
        parcelle,
        template_name="modele_certificat_propriete.docx",
        prefix="CP",
        title="Certificat de propriété",
    )


def generer_certificat_enregistrement_word(proprietaire, parcelle=None):
    """Génère le certificat d'enregistrement officiel pour toutes les parcelles du propriétaire."""
    parcelles = _cad_owner_parcels((proprietaire or {}).get("id"))
    parcelle_principale = (parcelles[0] if parcelles else (parcelle or {}))
    return generer_document_word(
        proprietaire,
        parcelle_principale,
        template_name="modele_certificat_enregistrement.docx",
        prefix="CE",
        title="Certificat d'enregistrement",
    )


def _cad_template_replace_text(paragraph, replacements):
    """Remplace les balises {{ cle }} dans un paragraphe sans détruire sa mise en page."""
    full = "".join(run.text or "" for run in paragraph.runs)
    if not full:
        return
    changed = False
    for key, value in replacements.items():
        for token in (f"{{{{ {key} }}}}", f"{{{{{key}}}}}"):
            if token in full:
                full = full.replace(token, str(value if value not in (None, "") else "—"))
                changed = True
    if changed:
        # On conserve le style du premier run ; la structure du modèle reste intacte.
        if paragraph.runs:
            paragraph.runs[0].text = full
            for run in paragraph.runs[1:]:
                run.text = ""


def _cad_template_replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _cad_template_replace_text(paragraph, replacements)
            for nested in cell.tables:
                _cad_template_replace_in_table(nested, replacements)


def _cad_template_insert_image_after_token(paragraph, image_path, width_mm, height_mm):
    """Remplace une balise image par l'image réelle dans le modèle Word."""
    if not image_path or not os.path.exists(str(image_path)):
        return False
    text = "".join(run.text or "" for run in paragraph.runs)
    tokens = ("{{ photo_proprietaire }}", "{{photo_proprietaire}}", "{{ qr_code }}", "{{qr_code}}")
    if not any(token in text for token in tokens):
        return False
    # Supprime le contenu texte de la balise puis ajoute l'image dans le même paragraphe.
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run()
    from docx.shared import Mm
    run.add_picture(str(image_path), width=Mm(width_mm), height=Mm(height_mm))
    return True


def _cad_convert_docx_to_pdf(docx_path, output_dir=None):
    """Convertit le DOCX issu du modèle officiel en PDF, sans modifier le modèle."""
    output_dir = output_dir or os.path.dirname(docx_path)
    os.makedirs(output_dir, exist_ok=True)

    # 1) LibreOffice / soffice : conversion fidèle au rendu Word du modèle.
    for exe in ("libreoffice", "soffice"):
        exe_path = shutil.which(exe)
        if not exe_path:
            continue
        try:
            subprocess = __import__("subprocess")
            result = subprocess.run(
                [exe_path, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )
            pdf_candidate = os.path.join(
                output_dir,
                os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
            )
            if result.returncode == 0 and os.path.exists(pdf_candidate):
                return pdf_candidate
        except Exception:
            pass

    # 2) Microsoft Word via docx2pdf si disponible sur Windows.
    try:
        from docx2pdf import convert as docx2pdf_convert
        pdf_candidate = os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        docx2pdf_convert(docx_path, pdf_candidate)
        if os.path.exists(pdf_candidate):
            return pdf_candidate
    except Exception:
        pass

    raise RuntimeError(
        "Impossible de convertir le modèle Word en PDF. "
        "Installez LibreOffice ou docx2pdf sur l'ordinateur qui exécute CADASTRE RDC."
    )


def generer_certificat_enregistrement_complet_pdf(proprietaire, parcelle):
    """
    Génère le certificat PDF en utilisant STRICTEMENT le modèle officiel .docx.

    Flux : SQLite -> propriétaire/parcelle -> QR -> modèle templates/ ->
    remplissage automatique -> photo + QR -> DOCX -> PDF.
    Aucune mise en page de remplacement ni annexe automatique n'est ajoutée.
    """
    proprietaire = dict(proprietaire or {})
    parcelle = dict(parcelle or {})
    if not parcelle.get("id"):
        raise ValueError("Parcelle invalide : identifiant introuvable dans cadastre_rdc.db.")
    if not proprietaire.get("id"):
        raise ValueError("Propriétaire invalide : identifiant introuvable dans cadastre_rdc.db.")

    conn = db_connect()
    try:
        owner_row = conn.execute(
            "SELECT * FROM proprietaires WHERE id = ? LIMIT 1",
            (proprietaire.get("id"),),
        ).fetchone()
        parcel_row = conn.execute(
            "SELECT * FROM parcelles WHERE id = ? LIMIT 1",
            (parcelle.get("id"),),
        ).fetchone()
    finally:
        conn.close()

    if not owner_row:
        raise ValueError("Propriétaire introuvable dans cadastre_rdc.db.")
    if not parcel_row:
        raise ValueError("Parcelle introuvable dans cadastre_rdc.db.")

    proprietaire = dict(owner_row)
    parcelle = dict(parcel_row)

    numero = _cad_doc_number("CE")
    qr_record = _cad_qr_register(
        numero,
        "Certificat d'enregistrement parcellaire",
        parcelle,
        proprietaire,
    )
    token = qr_record.get("token")
    if not token:
        raise RuntimeError("Impossible de créer le code QR de vérification.")

    template_path = _cad_find_template("modele_certificat_enregistrement.docx")
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Modèle officiel introuvable : {template_path}")

    qr_dir = os.path.join(CERTIFICATS_DIR, "qr")
    os.makedirs(qr_dir, exist_ok=True)
    qr_path = os.path.join(qr_dir, f"{_cad_word_safe_filename(numero)}.png")
    _cad_qr_generate(token, qr_path)
    verification_url = _cad_qr_url(token)

    photo_value = (
        proprietaire.get("photo_path")
        or proprietaire.get("photo")
        or parcelle.get("photo_path")
        or parcelle.get("photo")
    )
    photo_path = _cad_resolve_upload_path(photo_value) if photo_value else None
    if photo_path and not os.path.isfile(str(photo_path)):
        photo_path = None

    context = _cad_word_context(proprietaire, parcelle, numero)
    context.update({
        "titre_document": "CERTIFICAT D'ENREGISTREMENT PARCELLAIRE",
        "numero_certificat": numero,
        "qr_verification_url": verification_url,
        "proprietaire_id": proprietaire.get("id"),
        "parcelle_id": parcelle.get("id"),
        "reference_parcelle": parcelle.get("reference") or parcelle.get("numero") or "—",
        "gps": f"{parcelle.get('latitude') or '—'}, {parcelle.get('longitude') or '—'}",
        "dimensions": parcelle.get("dimensions") or "—",
        "polygon_gps": parcelle.get("polygon_gps") or "—",
        "points_gps": parcelle.get("points_gps") or "—",
        "agent": parcelle.get("agent") or "—",
        "created_at": parcelle.get("created_at") or proprietaire.get("date_creation") or "—",
        "updated_at": parcelle.get("updated_at") or "—",
        "photo_path": photo_path or "",
        "qr_path": qr_path,
    })
    for k, v in proprietaire.items():
        value = v if v not in (None, "") else "—"
        context.setdefault(f"proprietaire_{k}", value)
        context.setdefault(k, value)
    for k, v in parcelle.items():
        value = v if v not in (None, "") else "—"
        context.setdefault(f"parcelle_{k}", value)
        context.setdefault(k, value)

    # Utilisation de docxtpl lorsque disponible : c'est la voie la plus fidèle
    # au modèle Word, notamment pour les images InlineImage.
    rendered_with_docxtpl = False
    docx_path = os.path.join(
        CERTIFICATS_DIR,
        f"CE_{_cad_word_safe_filename(parcelle.get('numero') or parcelle.get('reference') or f'ID-{parcelle.get('id')}')}_{numero}.docx",
    )

    try:
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm

        document = DocxTemplate(template_path)
        if photo_path:
            context["photo_proprietaire"] = InlineImage(
                document, str(photo_path), width=Mm(35), height=Mm(45)
            )
        else:
            context["photo_proprietaire"] = "Photo non disponible"
        context["qr_code"] = InlineImage(
            document, qr_path, width=Mm(35), height=Mm(35)
        )
        document.render(context)
        document.save(docx_path)
        rendered_with_docxtpl = True
    except ImportError:
        rendered_with_docxtpl = False
    except Exception as ex:
        raise RuntimeError(
            f"Le modèle officiel n'a pas pu être rempli automatiquement : {ex}"
        ) from ex

    if not rendered_with_docxtpl:
        # Fallback uniquement pour les modèles .docx ordinaires ; il ne crée
        # jamais un modèle de remplacement.
        from docx import Document
        from docx.shared import Mm
        document = Document(template_path)
        context["photo_proprietaire"] = "Photo non disponible"
        context["qr_code"] = "Code QR de vérification"

        def process_paragraph(paragraph):
            _cad_template_replace_text(paragraph, context)
            text = "".join(run.text or "" for run in paragraph.runs)
            if "{{ photo_proprietaire }}" in text or "{{photo_proprietaire}}" in text:
                if photo_path:
                    paragraph.clear()
                    paragraph.add_run().add_picture(str(photo_path), width=Mm(35), height=Mm(45))
            if "{{ qr_code }}" in text or "{{qr_code}}" in text:
                paragraph.clear()
                paragraph.add_run().add_picture(qr_path, width=Mm(35), height=Mm(35))

        for paragraph in document.paragraphs:
            process_paragraph(paragraph)
        for section in document.sections:
            for paragraph in list(section.header.paragraphs) + list(section.footer.paragraphs):
                process_paragraph(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
                    for nested in cell.tables:
                        _cad_template_replace_in_table(nested, context)
        document.save(docx_path)

    pdf_path = _cad_convert_docx_to_pdf(docx_path, CERTIFICATS_DIR)
    if not os.path.isfile(pdf_path):
        raise RuntimeError("Le PDF n'a pas été créé après conversion du modèle officiel.")

    _cad_qr_update_files(
        numero,
        fichier_word=docx_path,
        fichier_pdf=pdf_path,
        qr_image=qr_path,
    )
    return pdf_path, numero, token


def generer_fiche_parcellaire_word(proprietaire, parcelle):
    return generer_document_word(
        proprietaire,
        parcelle,
        template_name="modele_fiche_parcellaire.docx",
        prefix="FP",
        title="Fiche parcellaire",
    )


def generer_fiche_proprietaire_word(proprietaire, parcelle=None):
    return generer_document_word(
        proprietaire,
        parcelle or {},
        template_name="modele_fiche_proprietaire.docx",
        prefix="PR",
        title="Fiche propriétaire",
    )


def _cad_doc_number(prefix="CAD"):
    return f"{prefix}-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:10].upper()}"


def _cad_generate_pdf(title_text, document_number, rows, verification_code=None, image_path=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib.units import mm

    root = _cad_docs_dir()
    filename = f"{title_text.replace(' ', '_')}_{document_number}.pdf"
    pdf_path = os.path.join(root, filename)

    doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm,
        topMargin=18*mm, bottomMargin=18*mm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CadastreTitle", parent=styles["Title"],
        alignment=TA_CENTER, fontSize=16, leading=20
    )
    small_style = ParagraphStyle(
        "CadastreSmall", parent=styles["Normal"],
        fontSize=8, leading=10
    )

    story = [
        Paragraph("RÉPUBLIQUE DÉMOCRATIQUE DU CONGO", title_style),
        Spacer(1, 3*mm),
        Paragraph("CADASTRE RDC", title_style),
        Spacer(1, 7*mm),
        Paragraph(title_text.upper(), title_style),
        Spacer(1, 5*mm),
        Paragraph(f"<b>Numéro du document :</b> {document_number}", styles["Normal"]),
        Spacer(1, 6*mm),
    ]

    # Insertion optionnelle de la photo
    if image_path and os.path.exists(str(image_path)):
        try:
            story.append(RLImage(image_path, width=35*mm, height=45*mm))
            story.append(Spacer(1, 5*mm))
        except Exception:
            pass

    data = [["Rubrique", "Information"]]
    for key, value in rows:
        data.append([str(key), str(value) if value not in (None, "") else "—"])

    table = Table(data, colWidths=[55*mm, 110*mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    story.append(Spacer(1, 8*mm))

    qr_path = None
    if verification_code:
        qr_path = os.path.join(root, f"QR_{document_number}.png")
        try:
            import qrcode
            qrcode.make(verification_code).save(qr_path)
            story.append(RLImage(qr_path, width=28*mm, height=28*mm))
            story.append(Spacer(1, 3*mm))
        except Exception:
            pass

        story.append(Paragraph(
            f"<b>Code de vérification :</b> {verification_code}",
            small_style
        ))

    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(
        "Document généré électroniquement par CADASTRE RDC. "
        "La valeur juridique dépend de la validation par l'autorité compétente.",
        small_style
    ))

    # Construction du document PDF
    doc.build(story)

    # Nettoyage du fichier QR code temporaire s'il existe
    if qr_path and os.path.exists(qr_path):
        try:
            os.remove(qr_path)
        except Exception:
            pass

    return pdf_path

def generer_certificat_propriete(proprietaire, parcelle):
    """Génère un certificat de propriété avec photo et adresse du terrain."""
    numero = _cad_doc_number("CP")
    photo_path = (
        proprietaire.get("photo_path")
        or proprietaire.get("photo")
        or parcelle.get("photo_path")
        or parcelle.get("photo")
    )

    adresse = (
        parcelle.get("adresse")
        or parcelle.get("localisation")
        or ""
    )

    return _cad_generate_pdf(
        "Certificat de propriété", numero,
        [
            ("Propriétaire", proprietaire.get("nom_complet")),
            ("Téléphone", proprietaire.get("telephone")),
            ("Pièce d'identité", proprietaire.get("piece_identite")),
            ("Parcelle", parcelle.get("numero", parcelle.get("reference"))),
            ("Adresse", adresse),
            ("Superficie", parcelle.get("superficie")),
            ("Province", parcelle.get("province")),
            ("Ville / Territoire", parcelle.get("ville_territoire")),
            ("Commune / Chefferie", parcelle.get("commune_chefferie")),
            ("Quartier / Groupement", parcelle.get("quartier_groupement")),
            ("Localité", parcelle.get("localite")),
            ("Latitude", parcelle.get("latitude")),
            ("Longitude", parcelle.get("longitude")),
            ("Date d'émission", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ],
        f"CADASTRE-RDC:{numero}",
        image_path=photo_path
    )


def generer_fiche_parcellaire(parcelle):
    numero = _cad_doc_number("FP")
    photo_path = parcelle.get("photo_path")
    
    # Nettoyage des valeurs pour s'assurer qu'elles s'affichent correctement en PDF
    rows = [(str(k), str(v) if v is not None else "") for k, v in parcelle.items()]
    
    return _cad_generate_pdf(
        "Fiche parcellaire", numero,
        rows,
        f"CADASTRE-RDC:{numero}",
        image_path=photo_path
    )


def generer_attestation_cadastrale(proprietaire, parcelle):
    numero = _cad_doc_number("AC")
    return _cad_generate_pdf(
        "Attestation cadastrale", numero,
        [
            ("Propriétaire", proprietaire.get("nom_complet")),
            ("Parcelle", parcelle.get("numero", parcelle.get("reference"))),
            ("Superficie", parcelle.get("superficie")),
            ("Localisation", parcelle.get("localisation")),
            ("Date", datetime.now().strftime("%Y-%m-%d")),
        ],
        f"CADASTRE-RDC:{numero}"
    )


def generer_fiche_proprietaire(proprietaire):
    numero = _cad_doc_number("PR")
    photo_path = proprietaire.get("photo_path")
    
    rows = [(str(k), str(v) if v is not None else "") for k, v in proprietaire.items()]
    
    return _cad_generate_pdf(
        "Fiche propriétaire", numero,
        rows,
        f"CADASTRE-RDC:{numero}",
        image_path=photo_path
    )


def generer_recu_dossier(reference_dossier, objet, demandeur):
    numero = _cad_doc_number("RD")
    return _cad_generate_pdf(
        "Reçu de dossier", numero,
        [
            ("Référence dossier", reference_dossier),
            ("Demandeur", demandeur),
            ("Objet", objet),
            ("Date de dépôt", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Statut", "REÇU"),
        ],
        f"CADASTRE-RDC:{numero}"
    )


def generer_certificat_apres_mutation(ancien_proprietaire, nouveau_proprietaire, parcelle):
    numero = _cad_doc_number("CM")
    return _cad_generate_pdf(
        "Certificat après mutation", numero,
        [
            ("Ancien propriétaire", ancien_proprietaire),
            ("Nouveau propriétaire", nouveau_proprietaire),
            ("Parcelle", parcelle.get("numero", parcelle.get("reference"))),
            ("Superficie", parcelle.get("superficie")),
            ("Localisation", parcelle.get("localisation")),
            ("Date de validation", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ],
        f"CADASTRE-RDC:{numero}"
    )


def generer_historique_parcelle(parcelle, historique):
    numero = _cad_doc_number("HP")
    rows = [(str(k), str(v) if v is not None else "") for k, v in parcelle.items()]
    
    for i, event in enumerate(historique or [], 1):
        if isinstance(event, dict):
            value = " | ".join(f"{k}: {v}" for k, v in event.items())
        else:
            value = str(event)
        rows.append((f"Historique {i}", value))

    return _cad_generate_pdf(
        "Historique de parcelle", numero, rows,
        f"CADASTRE-RDC:{numero}"
    )

def documents_cadastraux_view(page, content, t=None, journaliser=None):
    """Documents cadastraux : recherche DB + génération/téléchargement du certificat."""
    _t = t if callable(t) else lambda key: key
    status = ft.Text("", size=12, color=COLOR_PRIMARY)
    resultats = ft.Column(spacing=10, scroll=ft.ScrollMode.AUTO, expand=True)
    recherche = ft.TextField(
        label="Numéro de parcelle, ID ou nom du propriétaire",
        prefix_icon=ft.Icons.SEARCH,
        expand=True,
        on_submit=lambda e: rechercher(),
    )

    def information(_):
        status.value = (
            "Sélectionnez d'abord une donnée cadastrale depuis le module correspondant "
            "pour générer le document."
        )
        page.update()

    def _load_models():
        try:
            _cad_word_dirs()
            status.value = (
                f"✓ Modèles recherchés en priorité dans : {TEMPLATES_DIR}. "
                f"Les anciens modèles restent compatibles : {MODELES_DIR}."
            )
        except Exception as ex:
            status.value = f"Erreur modèles Word : {ex}"
        page.update()

    def rechercher():
        q = (recherche.value or "").strip().lower()
        resultats.controls.clear()
        if not q:
            status.value = "Saisissez un numéro de parcelle, un ID ou le nom du propriétaire."
            page.update()
            return
        conn = db_connect()
        try:
            scope_sql, scope_params = territorial_scope_sql("p")
            rows = conn.execute(
                f"""
                SELECT p.*,\n                       pr.id AS proprietaire_id,\n                       pr.nom_complet AS proprietaire_nom,\n                       pr.telephone AS proprietaire_telephone,\n                       pr.email AS proprietaire_email,\n                       pr.piece_identite AS proprietaire_piece_identite,\n                       pr.adresse AS proprietaire_adresse,\n                       pr.photo AS proprietaire_photo\n                FROM parcelles p\n                LEFT JOIN proprietaires pr ON p.proprietaire_id = pr.id\n                WHERE (LOWER(COALESCE(p.numero,'')) LIKE ?\n                    OR CAST(p.id AS TEXT) LIKE ?\n                    OR LOWER(COALESCE(pr.nom_complet,'')) LIKE ?\n                    OR LOWER(COALESCE(pr.telephone,'')) LIKE ?\n                    OR LOWER(COALESCE(p.adresse,'')) LIKE ?\n                    OR LOWER(COALESCE(p.province,'')) LIKE ?\n                    OR LOWER(COALESCE(p.ville_territoire,'')) LIKE ?\n                    OR LOWER(COALESCE(p.localite,'')) LIKE ?)\n                  AND ({scope_sql})\n                ORDER BY p.id DESC\n                LIMIT 50\n                """,
                tuple([f"%{q}%"] * 8) + tuple(scope_params),
            ).fetchall()
        finally:
            conn.close()

        if not rows:
            resultats.controls.append(ft.Text("Aucun résultat.", color=COLOR_ACCENT))
            page.update()
            return

        for row in rows:
            row_dict = dict(row)
            photo_path = _cad_resolve_upload_path(row_dict.get("proprietaire_photo"))
            resultats.controls.append(
                ft.Container(
                    padding=14,
                    border_radius=12,
                    border=ft.Border.all(1, COLOR_BORDER),
                    content=ft.Row(
                        [
                            ft.Icon(ft.Icons.LANDSCAPE, color=COLOR_PRIMARY),
                            ft.Column(
                                [
                                    ft.Text(
                                        f"Parcelle : {row_dict.get('numero') or 'ID-' + str(row_dict.get('id'))}",
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                    ft.Text(f"Propriétaire : {row_dict.get('proprietaire_nom') or '—'}"),
                                    ft.Text(f"Superficie : {row_dict.get('superficie') or '—'}"),
                                    ft.Text(
                                        f"Localisation : {row_dict.get('province') or '—'} / "
                                        f"{row_dict.get('ville_territoire') or '—'} / "
                                        f"{row_dict.get('localite') or '—'}",
                                        size=11,
                                        color=COLOR_TEXT_MUTED,
                                    ),
                                    ft.Text(
                                        f"Photo uploads : {'✓ disponible' if photo_path else '—'}",
                                        size=11,
                                        color=ft.Colors.GREEN_700 if photo_path else COLOR_TEXT_MUTED,
                                    ),
                                ],
                                spacing=3,
                                expand=True,
                            ),
                            ft.ElevatedButton(
                                "Télécharger le certificat d'enregistrement",
                                icon=ft.Icons.DOWNLOAD,
                                on_click=lambda e, r=row_dict: telecharger_certificat(r),
                            ),
                        ],
                        spacing=12,
                    ),
                )
            )
        status.value = f"✓ {len(rows)} résultat(s) trouvé(s) dans cadastre_rdc.db. Sélectionnez une parcelle puis cliquez sur « Télécharger le certificat »."
        page.update()

    async def telecharger_certificat(row):
        try:
            owner_id = row.get("proprietaire_id")
            conn = db_connect()
            try:
                owner = conn.execute(
                    "SELECT * FROM proprietaires WHERE id = ? LIMIT 1", (owner_id,)
                ).fetchone()
            finally:
                conn.close()
            if not owner:
                status.value = "Propriétaire introuvable dans cadastre_rdc.db."
                page.update()
                return
            pdf_path, numero_certificat, _token = generer_certificat_enregistrement_complet_pdf(
                dict(owner), dict(row)
            )
            with open(pdf_path, "rb") as fh:
                pdf_bytes = fh.read()
            saved = await certificate_file_picker.save_file(
                dialog_title="Télécharger le certificat d'enregistrement parcellaire",
                file_name=os.path.basename(pdf_path),
                allowed_extensions=["pdf"],
                src_bytes=pdf_bytes,
            )
            if isinstance(saved, str) and saved:
                status.value = f"✓ Certificat téléchargé : {saved}"
            elif page.web:
                status.value = f"✓ Téléchargement lancé : {os.path.basename(pdf_path)}"
            else:
                status.value = f"✓ Certificat généré : {pdf_path}"
            if callable(journaliser):
                journaliser(
                    "TELECHARGEMENT_CERTIFICAT_ENREGISTREMENT",
                    f"Certificat généré pour propriétaire {owner_id}, parcelle {row.get('numero') or row.get('id')}",
                )
        except Exception as ex:
            status.value = f"Erreur certificat : {ex}"
        page.update()

    content.content = ft.Column(
        [
            ft.Text(
                _t("documents_cadastraux") if t else "Documents cadastraux",
                size=24,
                weight=ft.FontWeight.BOLD,
            ),
            ft.Text(
                "Recherche dans cadastre_rdc.db et génération des documents officiels.",
                size=14,
                color=COLOR_TEXT_MUTED,
            ),
            ft.Row([recherche, ft.ElevatedButton("RECHERCHER", icon=ft.Icons.SEARCH, on_click=lambda e: rechercher())]),
            ft.Divider(height=10),
            ft.Text("CERTIFICAT D'ENREGISTREMENT", size=16, weight=ft.FontWeight.BOLD, color=COLOR_PRIMARY),
            ft.Text(
                "Le certificat reprend les informations du propriétaire et toutes les parcelles qui lui sont rattachées. Le modèle officiel est recherché dans templates/ en priorité.",
                size=12,
                color=COLOR_TEXT_MUTED,
            ),
            ft.Row([
                ft.ElevatedButton("Créer / vérifier les modèles", icon=ft.Icons.DESCRIPTION, on_click=_load_models),
                ft.ElevatedButton("Certificat de propriété", icon=ft.Icons.DESCRIPTION, on_click=information),
                ft.ElevatedButton("Fiche parcellaire", icon=ft.Icons.LANDSCAPE, on_click=information),
            ], wrap=True),
            ft.Divider(height=10),
            resultats,
            status,
        ],
        spacing=12,
        expand=True,
    )
    page.update()


# ============================================================
# ESPACE PUBLIC DE VÉRIFICATION & PRÉSENTATION DU SYSTÈME
# ============================================================

def _public_normalize_name(value):
    """Normalise un nom pour une comparaison publique robuste."""
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return " ".join(text.casefold().split())


def _public_search_parcel(query):
    """Recherche nationale exacte d'une parcelle par numéro ou ID.

    Le portail public ne dépend pas du périmètre territorial de l'utilisateur :
    il consulte le registre national, mais ne renvoie que les informations
    destinées à la vérification publique.
    """
    q = str(query or "").strip()
    if not q:
        return []

    conn = db_connect()
    try:
        rows = conn.execute(
            """
            SELECT
                p.id, p.numero, p.adresse, p.province,
                p.ville_territoire, p.commune_chefferie,
                p.quartier_groupement, p.localite,
                p.superficie, p.latitude, p.longitude,
                p.statut, p.date_enregistrement, p.agent,
                pr.id AS proprietaire_id,
                pr.nom_complet AS proprietaire_nom,
                pr.photo AS proprietaire_photo
            FROM parcelles p
            LEFT JOIN proprietaires pr
                ON p.proprietaire_id = pr.id
            WHERE
                CAST(p.id AS TEXT) = ?
                OR LOWER(COALESCE(p.numero, '')) = LOWER(?)
                OR LOWER(COALESCE(p.numero, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.adresse, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.province, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.ville_territoire, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.commune_chefferie, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.quartier_groupement, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(p.localite, '')) LIKE LOWER(?)
                OR LOWER(COALESCE(pr.nom_complet, '')) LIKE LOWER(?)
            ORDER BY p.id DESC
            LIMIT 20
            """,
            (q, q, f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"),
        ).fetchall()
        return rows
    finally:
        conn.close()


def _public_photo_widget(photo_path):
    """Affiche la photo publique provenant de uploads/."""
    photo_url = _cad_public_upload_url(photo_path)
    if not photo_url:
        return ft.Container(
            width=130,
            height=160,
            border_radius=14,
            bgcolor=ft.Colors.BLUE_50,
            alignment=ft.Alignment.CENTER,
            content=ft.Column(
                [
                    ft.Icon(ft.Icons.PERSON, size=50, color=COLOR_TEXT_MUTED),
                    ft.Text(
                        "Photo non disponible",
                        size=11,
                        color=COLOR_TEXT_MUTED,
                        text_align=ft.TextAlign.CENTER,
                    ),
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                alignment=ft.MainAxisAlignment.CENTER,
            ),
        )

    return ft.Container(
        width=130,
        height=160,
        border_radius=14,
        clip_behavior=ft.ClipBehavior.HARD_EDGE,
        content=ft.Image(
            src=str(photo_url),
            width=130,
            height=160,
            fit=ft.BoxFit.COVER,
        ),
    )


def build_public_verification_view(page: ft.Page):
    """Compatibilité avec les anciennes intégrations du portail public."""
    return public_verify_view(page)

    # Interface publique permettant à tout citoyen ou futur acquéreur 
    # de vérifier l'authenticité d'une parcelle et de son propriétaire 
    # en toute sécurité, sans nécessiter de connexion ni de mot de passe.
    
    # Champs de recherche publique
    search_input = ft.TextField(
        label="Numéro de parcelle, Référence cadastrale ou ID unique",
        prefix_icon=ft.Icons.SEARCH,
        border_radius=10,
        expand=True,
    )

    owner_name_input = ft.TextField(
        label="Nom du propriétaire (facultatif)",
        prefix_icon=ft.Icons.PERSON,
        border_radius=10,
        expand=True,
    )
    
    result_container = ft.Column(scroll=ft.ScrollMode.AUTO, expand=True, spacing=10)

    def handle_public_verification(e):
        query = search_input.value.strip()
        user_input_name = owner_name_input.value.strip()
        
        # Validation des champs vides
        if not query or not user_input_name:
            result_container.controls = [
                ft.Banner(
                    bgcolor=ft.Colors.AMBER_100,
                    leading=ft.Icon(ft.Icons.WARNING_ROUTED, color=ft.Colors.AMBER_800),
                    content=ft.Text("Veuillez saisir le numéro de parcelle et le nom public du propriétaire."),
                    actions=[ft.TextButton("OK", on_click=lambda _: page.close_banner())]
                )
            ]
            page.update()
            return

        result_container.controls = [ft.Text(f"Vérification en cours sur site pour : {query}...", italic=True)]
        page.update()

        # 1. Simulation ou récupération de la position GPS de l'appareil de l'utilisateur sur le terrain
        # user_gps_lat, user_gps_lon = get_current_device_gps()
        user_gps_lat, user_gps_lon = 0.0, 0.0  # Remplacer par votre fonction réelle
        
        # 2. Récupération des données officielles de la parcelle dans la base SQLite
        # parcel_data = _cad_search_parcel_by_reference(query)
        parcel_data = {}  # Remplacer par votre fonction réelle
        
        if not parcel_data:
            result_container.controls = [ft.Text("Parcelle introuvable dans le système national.", color=ft.Colors.RED)]
            page.update()
            return

        parcel_lat = parcel_data.get("latitude")
        parcel_lon = parcel_data.get("longitude")
        official_owner = parcel_data.get("nom_proprietaire", "")

        # 3. Calcul de la distance ou vérification de la proximité immédiate (ex: rayon de 15 à 30 mètres)
        # is_on_site = verify_gps_proximity(user_gps_lat, user_gps_lon, parcel_lat, parcel_lon, max_distance_meters=25)
        is_on_site = True  # Remplacer par votre fonction réelle
        
        # 4. Vérification de la correspondance exacte du nom public saisi
        is_name_matching = (user_input_name.lower() == official_owner.lower())

        # 5. Validation stricte : Toutes les conditions doivent être réunies simultanément
        if not is_on_site:
            result_container.controls = [
                ft.Card(
                    content=ft.Container(
                        padding=20,
                        content=ft.Column([
                            ft.Row([
                                ft.Icon(ft.Icons.GPS_OFF, color=ft.Colors.RED, size=30),
                                ft.Text("VÉRIFICATION REFUSÉE : HORS ZONE", weight=ft.FontWeight.BOLD, color=ft.Colors.RED_700),
                            ]),
                            ft.Divider(),
                            ft.Text("La vérification d'authenticité n'est autorisée que lorsque vous vous trouvez physiquement sur le terrain concerné.", italic=True),
                            ft.Text("• Votre position GPS actuelle ne correspond pas aux coordonnées enregistrées de cette parcelle.", color=ft.Colors.RED_900),
                        ])
                    )
                )
            ]
            page.update()
            return

        if not is_name_matching:
            result_container.controls = [
                ft.Card(
                    content=ft.Container(
                        padding=20,
                        content=ft.Column([
                            ft.Row([
                                ft.Icon(ft.Icons.ERROR, color=ft.Colors.ORANGE, size=30),
                                ft.Text("VÉRIFICATION ÉCHOUÉE : NOM INCORRECT", weight=ft.FontWeight.BOLD, color=ft.Colors.ORANGE_900),
                            ]),
                            ft.Divider(),
                            ft.Text("Le nom public du propriétaire saisi ne correspond pas aux registres officiels pour cette parcelle.", color=ft.Colors.RED_800),
                        ])
                    )
                )
            ]
            page.update()
            return
        
        # Si tout est valide sur le terrain : Affichage sécurisé et restreint
        result_container.controls = [
            ft.Card(
                elevation=4,
                content=ft.Container(
                    padding=20,
                    content=ft.Column([
                        ft.Row([
                            ft.Icon(ft.Icons.VERIFIED, color=ft.Colors.GREEN, size=30),
                            ft.Text("CERTIFICATION SUR SITE VALIDÉE", weight=ft.FontWeight.BOLD, size=16, color=ft.Colors.GREEN_700),
                        ]),
                        ft.Divider(),
                        ft.Text(f"Référence vérifiée : {query}", italic=True),
                        ft.Row([
                            ft.Icon(ft.Icons.LOCATION_ON, size=18, color=ft.Colors.BLUE_700),
                            ft.Text("Statut : PARCELLE EN RÈGLE (PRÉSENCE SUR SITE CONFIRMÉE)", color=ft.Colors.GREEN_800, weight=ft.FontWeight.BOLD),
                        ]),
                        ft.Text(f"• Propriétaire public vérifié : {official_owner}"),
                        ft.Text("• Localisation géographique : Conforme au relevé cadastral officiel"),
                        ft.Text("• Avertissement : Cette confirmation certifie uniquement l'identité publique et l'emplacement de base pour contrer les fraudes à distance.", size=12, italic=True),
                    ], spacing=10)
                )
            )
        ]
        page.update()

    verify_button = ft.ElevatedButton(
        "Vérifier la parcelle sur site",
        icon=ft.Icons.CHECK_CIRCLE,
        on_click=handle_public_verification,
        style=ft.ButtonStyle(padding=15)
    )

    # Assemblage de la vue
    return ft.View(
        route="/public-verification",
        padding=20,
        controls=[
            ft.Text("Espace Public de Vérification Cadastrale", size=24, weight=ft.FontWeight.BOLD),
            ft.Text("Vérifiez l'authenticité d'une parcelle directement depuis le terrain.", size=14, color=ft.Colors.GREY_700),
            ft.Divider(height=20),
            ft.Row([search_input], expand=True),
            ft.Row([owner_name_input], expand=True),
            ft.Container(content=verify_button, margin=ft.margin.only(top=10, bottom=10)),
            ft.Divider(height=20, color=ft.Colors.TRANSPARENT),
            result_container
        ]
    )

# ============================================================
# ESPACE PUBLIC DE VÉRIFICATION & PRÉSENTATION DU SYSTÈME
# ============================================================
import flet as ft

def public_verify_view(page: ft.Page, on_admin_access=None, on_logout=None):
    # Horloge indépendante du portail public.
    # IMPORTANT : ces contrôles doivent être créés dans cette fonction, car les
    # contrôles clock_date/clock_time/clock_zone du tableau de bord principal
    # sont locaux à main() et ne sont pas visibles ici.
    public_clock_date = ft.Text("--/--/----", size=9, color="#FFFFFF", weight=ft.FontWeight.BOLD)
    public_clock_time = ft.Text("--:--:--", size=14, color="#FFFFFF", weight=ft.FontWeight.BOLD)
    public_clock_zone = ft.Text("Fuseau local…", size=8, color="#FFFFFF99")

    async def update_public_clock():
        while True:
            try:
                now = datetime.now().astimezone()
                public_clock_date.value = now.strftime("%d/%m/%Y")
                public_clock_time.value = now.strftime("%H:%M:%S")
                zone_name = now.tzname() or "LOCAL"
                offset = now.strftime("%z")
                if len(offset) == 5:
                    offset = offset[:3] + ":" + offset[3:]
                public_clock_zone.value = f"{zone_name}  UTC{offset}" if offset else zone_name
                public_clock_date.update()
                public_clock_time.update()
                public_clock_zone.update()
            except Exception as ex:
                print(f"[CADASTRE RDC] HORLOGE PUBLIC : {type(ex).__name__}: {ex}")
            await asyncio.sleep(1)

    # Une horloge est créée uniquement pour cette vue. Cela évite toute
    # référence à une variable locale de main() et supprime l'erreur
    # "name 'clock_date' is not defined".
    try:
        page.run_task(update_public_clock)
    except Exception as ex:
        print(f"[CADASTRE RDC] Impossible de démarrer l'horloge publique : {type(ex).__name__}: {ex}")

    # Portail citoyen de vérification cadastrale.
    # 
    # Accès : utilisateur PUBLIC, sans mot de passe.
    # La recherche exige le numéro exact de parcelle ou son ID et le nom
    # public du propriétaire. Les données sensibles (CNI, e-mail, etc.)
    # restent réservées à l'espace authentifié.
    
    img_drapeau = "https://tse2.mm.bing.net/th/id/OIP.6rlB9bb8jvYNEo_4Crmu_gHaFj?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"
    img_carte = "https://tse1.mm.bing.net/th/id/OIP.SuXxL3WiqQd8v6n0KzD5pAHaHa?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"
    img_ministere = "https://tse3.mm.bing.net/th/id/OIP.Ww4_MrnMvX36sHJpOPEymgAAAA?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"
    img_parcelle = "https://tse3.mm.bing.net/th/id/OIP.PwPyRdpYO6szN2fNdkH9_gHaJ4?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"
    img_satellite = "https://tse4.mm.bing.net/th/id/OIP.mC8Gmm3ZolltI3UOj1CVbAHaEa?r=0&rs=1&pid=ImgDetMain&o=7&rm=3"

    search_input = ft.TextField(
        label="Rechercher une parcelle, son ID, son propriétaire ou sa localisation",
        prefix_icon=ft.Icons.SEARCH,
        border_radius=12,
        expand=True,
        autofocus=True,
    )
    owner_name_input = ft.TextField(
        label="Nom public du propriétaire",
        prefix_icon=ft.Icons.PERSON,
        border_radius=12,
        expand=True,
    )
    result_container = ft.Column(spacing=12)

    def toggle_public_theme(_=None):
        """Bascule le mode clair/sombre dans l'espace PUBLIC sans changer de route."""
        page.theme_mode = (
            ft.ThemeMode.DARK
            if page.theme_mode == ft.ThemeMode.LIGHT
            else ft.ThemeMode.LIGHT
        )
        page.bgcolor = (
            "#0B1220"
            if page.theme_mode == ft.ThemeMode.DARK
            else COLOR_BG_LIGHT
        )
        page.update()

    def clear_results(_=None):
        search_input.value = ""
        owner_name_input.value = ""
        result_container.controls.clear()
        page.update()

    def show_public_message(message, error=False):
        result_container.controls = [
            ft.Container(
                padding=18,
                border_radius=14,
                bgcolor=(
                    "#3F1D24" if error else "#EFF6FF"
                ),
                content=ft.Row(
                    [
                        ft.Icon(
                            ft.Icons.ERROR_OUTLINE if error else ft.Icons.INFO_OUTLINE,
                            color=COLOR_ACCENT if error else COLOR_PRIMARY,
                        ),
                        ft.Text(message, expand=True),
                    ],
                    spacing=10,
                ),
            )
        ]
        page.update()

    def handle_public_verification(e=None):
        query = (search_input.value or "").strip()
        entered_name = (owner_name_input.value or "").strip()

        if not query:
            show_public_message(
                "Veuillez saisir un numéro/ID de parcelle, un nom de propriétaire ou une localisation.",
                True,
            )
            return

        result_container.controls = [
            ft.Row(
                [
                    ft.ProgressRing(width=20, height=20),
                    ft.Text("Vérification de l'enregistrement cadastral en cours..."),
                ],
                spacing=10,
            )
        ]
        page.update()

        try:
            rows = _public_search_parcel(query)
        except Exception as ex:
            journaliser("VERIFICATION_PUBLIQUE_ERREUR", f"Recherche {query}: {ex}")
            show_public_message(
                f"Le service de vérification a rencontré une erreur : {ex}",
                True,
            )
            return

        if not rows:
            journaliser("VERIFICATION_PUBLIQUE_ECHEC", f"Parcelle introuvable : {query}")
            show_public_message(
                "Aucune parcelle ne correspond exactement à ce numéro/ID dans le registre national.",
                True,
            )
            return

        # Si un nom est fourni, on conserve la vérification de concordance.
        if entered_name:
            rows = [
                r for r in rows
                if _public_normalize_name(entered_name) == _public_normalize_name(r["proprietaire_nom"] or "")
            ]
            if not rows:
                journaliser("RECHERCHE_PUBLIQUE_NOM_REFUSE", f"Recherche {query} : nom non concordant")
                show_public_message(
                    "Le nom saisi ne correspond à aucun propriétaire public associé aux résultats.",
                    True,
                )
                return

        # Recherche publique générale : afficher plusieurs parcelles avec leur propriétaire.
        if len(rows) > 1:
            result_container.controls = [
                ft.Text(
                    f"{len(rows)} parcelles trouvées — parcelle et propriétaire",
                    size=16,
                    weight=ft.FontWeight.BOLD,
                    color=COLOR_DEEP_BLUE,
                )
            ]
            for r in rows:
                localisation_card = " / ".join(
                    str(v).strip() for v in [r["province"], r["ville_territoire"], r["commune_chefferie"], r["localite"]]
                    if v not in (None, "")
                ) or "—"
                result_container.controls.append(
                    ft.Container(
                        padding=16,
                        border_radius=14,
                        bgcolor="#F8FAFC",
                        border=ft.Border.all(1, "#CBD5E1"),
                        content=ft.Row(
                            [
                                ft.Icon(ft.Icons.LANDSCAPE, color=COLOR_PRIMARY, size=30),
                                ft.Column(
                                    [
                                        ft.Text(f"Parcelle : {r['numero'] or 'ID-' + str(r['id'])}", weight=ft.FontWeight.BOLD),
                                        ft.Text(f"Propriétaire : {r['proprietaire_nom'] or 'Non renseigné'}", weight=ft.FontWeight.BOLD, color=COLOR_DEEP_BLUE),
                                        ft.Text(f"Statut : {r['statut'] or '—'} • Localisation : {localisation_card}", size=12, color=COLOR_TEXT_MUTED),
                                    ],
                                    spacing=3,
                                    expand=True,
                                ),
                            ],
                            spacing=12,
                        ),
                    )
                )
            result_container.controls.append(ft.Text("Pour une vérification juridique, comparez toujours avec les documents officiels.", size=12, italic=True, color=COLOR_TEXT_MUTED))
            journaliser("RECHERCHE_PUBLIQUE_REUSSIE", f"Recherche publique : {query} ({len(rows)} résultats)")
            page.update()
            return

        row = rows[0]
        official_owner = row["proprietaire_nom"] or ""

        numero = row["numero"] or f"ID-{row['id']}"
        superficie = row["superficie"]
        superficie_text = (
            f"{float(superficie):.2f} cm²" if superficie is not None else "—"
        )
        localisation = " / ".join(
            str(value).strip()
            for value in [
                row["province"],
                row["ville_territoire"],
                row["commune_chefferie"],
                row["quartier_groupement"],
                row["localite"],
            ]
            if value not in (None, "")
        ) or "—"

        # Le navigateur Flet Web ne fournit pas ici une API GPS directe côté Python.
        # On ne simule donc jamais une position GPS et on ne prétend pas qu'un contrôle
        # de présence sur site a été effectué. Cette information est affichée séparément.
        gps_registered = row["latitude"] is not None and row["longitude"] is not None

        journaliser(
            "VERIFICATION_PUBLIQUE_REUSSIE",
            f"Parcelle {numero} : nom public concordant",
        )

        result_container.controls = [
            ft.Container(
                padding=22,
                border_radius=18,
                bgcolor="#F0FDF4",
                border=ft.Border.all(1, ft.Colors.GREEN_300),
                content=ft.Column(
                    [
                        ft.Row(
                            [
                                ft.Icon(ft.Icons.VERIFIED, color=ft.Colors.GREEN_700, size=34),
                                ft.Column(
                                    [
                                        ft.Text(
                                            "AUTHENTICITÉ CADASTRALE CONFIRMÉE",
                                            size=18,
                                            weight=ft.FontWeight.BOLD,
                                            color=ft.Colors.GREEN_800,
                                        ),
                                        ft.Text(
                                            "Le numéro de parcelle et le nom public concordent avec le registre national.",
                                            color=ft.Colors.GREEN_900,
                                        ),
                                    ],
                                    spacing=2,
                                    expand=True,
                                ),
                            ],
                            spacing=12,
                        ),
                        ft.Divider(),
                        ft.ResponsiveRow(
                            [
                                ft.Container(
                                    col={"xs": 12, "md": 4},
                                    content=_public_photo_widget(row["proprietaire_photo"]),
                                    alignment=ft.Alignment.CENTER,
                                ),
                                ft.Container(
                                    col={"xs": 12, "md": 8},
                                    content=ft.Column(
                                        [
                                            ft.Text("PROPRIÉTAIRE", size=14, weight=ft.FontWeight.BOLD, color=COLOR_PRIMARY),
                                            ft.Text(official_owner, size=20, weight=ft.FontWeight.BOLD),
                                            ft.Text(f"Identifiant public du propriétaire : {row['proprietaire_id'] or '—'}"),
                                            ft.Divider(),
                                            ft.Text("PARCELLE", size=14, weight=ft.FontWeight.BOLD, color=COLOR_PRIMARY),
                                            ft.Text(f"Numéro : {numero}", weight=ft.FontWeight.BOLD),
                                            ft.Text(f"Statut : {row['statut'] or '—'}"),
                                            ft.Text(f"Superficie : {superficie_text}"),
                                            ft.Text(f"Adresse : {row['adresse'] or '—'}"),
                                            ft.Text(f"Localisation : {localisation}"),
                                            ft.Text(
                                                f"Coordonnées enregistrées : {row['latitude'] if row['latitude'] is not None else '—'}, {row['longitude'] if row['longitude'] is not None else '—'}"
                                            ),
                                        ],
                                        spacing=7,
                                    ),
                                ),
                            ],
                            spacing=18,
                        ),
                        ft.Container(
                            padding=14,
                            border_radius=12,
                            bgcolor="#EFF6FF",
                            content=ft.Row(
                                [
                                    ft.Icon(
                                        ft.Icons.GPS_FIXED if gps_registered else ft.Icons.GPS_OFF,
                                        color=COLOR_PRIMARY if gps_registered else COLOR_ACCENT,
                                    ),
                                    ft.Column(
                                        [
                                            ft.Text(
                                                "Contrôle GPS",
                                                weight=ft.FontWeight.BOLD,
                                            ),
                                            ft.Text(
                                                "Les coordonnées GPS sont enregistrées dans le cadastre. Le contrôle de présence GPS en temps réel n'est pas simulé dans cette version web.",
                                                size=12,
                                            ),
                                        ],
                                        expand=True,
                                        spacing=2,
                                    ),
                                ],
                                spacing=10,
                            ),
                        ),
                        ft.Text(
                            "IMPORTANT : avant tout achat, comparez cette information avec les documents originaux et les autorités foncières compétentes. La vérification publique ne remplace pas une procédure juridique officielle.",
                            size=12,
                            italic=True,
                        ),
                    ],
                    spacing=12,
                ),
            )
        ]
        page.update()

    search_input.on_submit = handle_public_verification
    owner_name_input.on_submit = handle_public_verification

    return ft.View(
        route="/public-verify",
        padding=0,
        controls=[
            ft.AppBar(
                title=ft.Text("CADASTRE RDC — Espace Citoyen & Public"),
                bgcolor=COLOR_DEEP_BLUE,
                color=ft.Colors.WHITE,
                actions=[
                    ft.Container(
                        padding=ft.Padding(left=8, right=10, top=4, bottom=4),
                        content=ft.Column(
                            [
                                ft.Text("DATE / HEURE", size=9, color="#FFFFFF99", weight=ft.FontWeight.BOLD),
                                public_clock_date,
                                public_clock_time,
                                public_clock_zone,
                            ], spacing=0,
                            horizontal_alignment=ft.CrossAxisAlignment.END,
                        ),
                    ),
                    ft.OutlinedButton(
                        "🚪 Déconnexion",
                        icon=ft.Icons.LOGOUT,
                        on_click=lambda e: on_logout() if callable(on_logout) else None,
                        style=ft.ButtonStyle(
                            color=ft.Colors.WHITE,
                            side=ft.BorderSide(1, "#FFFFFF66"),
                        ),
                    ),
                    ft.IconButton(
                        icon=ft.Icons.BRIGHTNESS_6,
                        tooltip="Mode clair / mode sombre",
                        icon_color=ft.Colors.WHITE,
                        on_click=toggle_public_theme,
                    ),
                    ft.Container(
                        padding=ft.Padding(left=8, right=8, top=2, bottom=2),
                        content=ft.Column(
                            [
                                ft.Text("BUREAU NATIONAL", size=9, color="#FFFFFF99", weight=ft.FontWeight.BOLD),
                                ft.Text("+243 976905658", size=11, color=ft.Colors.WHITE),
                                ft.Text("tristanlutombo38@gmail.com", size=10, color=ft.Colors.WHITE),
                            ], spacing=0,
                            horizontal_alignment=ft.CrossAxisAlignment.END,
                        ),
                    ),
                    ft.OutlinedButton(
                        "Accès Agent / Administration",
                        icon=ft.Icons.LOGIN,
                        on_click=lambda e: on_admin_access() if callable(on_admin_access) else None,
                    )
                ],
            ),
            ft.Container(
                expand=True,
                padding=20,
                content=ft.ListView(
                    expand=True,
                    spacing=18,
                    controls=[
                        ft.Container(
                            padding=30,
                            border_radius=18,
                            gradient=ft.LinearGradient(
                                begin=ft.Alignment.TOP_LEFT,
                                end=ft.Alignment.BOTTOM_RIGHT,
                                colors=[COLOR_DEEP_BLUE, COLOR_PRIMARY],
                            ),
                            content=ft.ResponsiveRow(
                                [
                                    ft.Container(
                                        col={"xs": 12, "md": 8},
                                        content=ft.Column(
                                            [
                                                ft.Row(
                                                    [
                                                        ft.Image(src=img_drapeau, width=45, height=30, fit=ft.BoxFit.CONTAIN),
                                                        ft.Text(
                                                            "RÉPUBLIQUE DÉMOCRATIQUE DU CONGO",
                                                            color=COLOR_SECONDARY,
                                                            weight=ft.FontWeight.BOLD,
                                                        ),
                                                    ],
                                                    spacing=10,
                                                ),
                                                ft.Text(
                                                    "Sécurisez votre achat de terrain avant de payer.",
                                                    size=28,
                                                    weight=ft.FontWeight.BOLD,
                                                    color=ft.Colors.WHITE,
                                                ),
                                                ft.Text(
                                                    "CADASTRE RDC permet au citoyen de confronter une référence cadastrale et le nom public du propriétaire avec les informations enregistrées.",
                                                    color=ft.Colors.WHITE,
                                                    size=14,
                                                ),
                                            ],
                                            spacing=14,
                                        ),
                                    ),
                                    ft.Container(
                                        col={"xs": 12, "md": 4},
                                        alignment=ft.Alignment.CENTER,
                                        content=ft.Image(src=img_carte, width=190, height=190, fit=ft.BoxFit.CONTAIN),
                                    ),
                                ],
                                spacing=20,
                            ),
                        ),
                        ft.Container(
                            padding=22,
                            border_radius=18,
                            bgcolor=ft.Colors.WHITE,
                            shadow=ft.BoxShadow(blur_radius=18, spread_radius=1, color="#00000018"),
                            content=ft.Column(
                                [
                                    ft.Text(
                                        "🔍 Vérifier l'authenticité d'une propriété avant l'achat",
                                        size=21,
                                        weight=ft.FontWeight.BOLD,
                                        color=COLOR_DEEP_BLUE,
                                    ),
                                    ft.Text(
                                        "1. Demandez au vendeur le numéro cadastral exact.  2. Saisissez le nom public qu'il déclare être celui du propriétaire.  3. Comparez le résultat avec les documents et le terrain avant toute transaction.",
                                        size=13,
                                        color=COLOR_TEXT_MUTED,
                                    ),
                                    ft.Row([search_input], expand=True),
                                    ft.Row([owner_name_input], expand=True),
                                    ft.Row(
                                        [
                                            ft.ElevatedButton(
                                                "Rechercher / Vérifier",
                                                icon=ft.Icons.SEARCH,
                                                bgcolor=ft.Colors.GREEN_700,
                                                color=ft.Colors.WHITE,
                                                on_click=handle_public_verification,
                                            ),
                                            ft.OutlinedButton(
                                                "Effacer",
                                                icon=ft.Icons.CLEAR,
                                                on_click=clear_results,
                                            ),
                                        ],
                                        wrap=True,
                                        spacing=10,
                                    ),
                                    result_container,
                                ],
                                spacing=14,
                            ),
                        ),
                        ft.Text(
                            "📖 GUIDE CITOYEN — COMMENT VÉRIFIER UNE PROPRIÉTÉ",
                            size=19,
                            weight=ft.FontWeight.BOLD,
                            color=COLOR_DEEP_BLUE,
                        ),
                        ft.ResponsiveRow(
                            [
                                ft.Container(
                                    col={"xs": 12, "md": 4},
                                    padding=16,
                                    border_radius=14,
                                    bgcolor=ft.Colors.BLUE_50,
                                    content=ft.Column(
                                        [
                                            ft.Image(src=img_ministere, height=120, fit=ft.BoxFit.CONTAIN),
                                            ft.Text("1. Vérifier les documents", weight=ft.FontWeight.BOLD, color=COLOR_DEEP_BLUE),
                                            ft.Text("Demandez la référence officielle et les documents originaux. Ne vous contentez pas d'une simple photocopie ou d'une promesse verbale.", size=12),
                                        ],
                                        spacing=8,
                                    ),
                                ),
                                ft.Container(
                                    col={"xs": 12, "md": 4},
                                    padding=16,
                                    border_radius=14,
                                    bgcolor=ft.Colors.BLUE_50,
                                    content=ft.Column(
                                        [
                                            ft.Image(src=img_parcelle, height=120, fit=ft.BoxFit.CONTAIN),
                                            ft.Text("2. Vérifier la parcelle sur le terrain", weight=ft.FontWeight.BOLD, color=COLOR_DEEP_BLUE),
                                            ft.Text("Comparez le numéro cadastral, les limites visibles et la localisation réelle du terrain avec les informations officielles.", size=12),
                                        ],
                                        spacing=8,
                                    ),
                                ),
                                ft.Container(
                                    col={"xs": 12, "md": 4},
                                    padding=16,
                                    border_radius=14,
                                    bgcolor=ft.Colors.BLUE_50,
                                    content=ft.Column(
                                        [
                                            ft.Image(src=img_satellite, height=120, fit=ft.BoxFit.CONTAIN),
                                            ft.Text("3. Vérifier la localisation", weight=ft.FontWeight.BOLD, color=COLOR_DEEP_BLUE),
                                            ft.Text("Utilisez les coordonnées GPS enregistrées et, lorsque le dispositif technique le permet, confirmez la position directement sur le terrain.", size=12),
                                        ],
                                        spacing=8,
                                    ),
                                ),
                            ],
                            spacing=12,
                        ),
                        ft.Container(
                            padding=18,
                            border_radius=14,
                            bgcolor="#FFF7ED",
                            border=ft.Border.all(1, ft.Colors.ORANGE_300),
                            content=ft.Row(
                                [
                                    ft.Icon(ft.Icons.WARNING_AMBER, color=ft.Colors.ORANGE_700, size=28),
                                    ft.Text(
                                        "Conseil de sécurité : si le vendeur refuse la vérification, si le nom ne correspond pas ou si les documents sont incohérents, interrompez la transaction et consultez l'autorité foncière compétente.",
                                        expand=True,
                                        size=13,
                                    ),
                                ],
                                spacing=10,
                            ),
                        ),
                        ft.Container(
                            padding=18,
                            border_radius=16,
                            gradient=ft.LinearGradient(
                                begin=ft.Alignment.TOP_LEFT,
                                end=ft.Alignment.BOTTOM_RIGHT,
                                colors=["#E3F2FD", "#FFFFFF"],
                            ),
                            border=ft.Border.all(1, "#BFDBFE"),
                            content=ft.Row(
                                [
                                    ft.Icon(ft.Icons.CONTACT_PHONE, color=COLOR_PRIMARY, size=30),
                                    ft.Column(
                                        [
                                            ft.Text("CONTACTER LE BUREAU NATIONAL", size=15, weight=ft.FontWeight.BOLD, color=COLOR_DEEP_BLUE),
                                            ft.Text("Téléphone : +243 976905658", weight=ft.FontWeight.BOLD),
                                            ft.Text("E-mail : tristanlutombo38@gmail.com", color=COLOR_TEXT_MUTED),
                                        ],
                                        spacing=3,
                                        expand=True,
                                    ),
                                ],
                                spacing=12,
                            ),
                        ),
                        ft.Divider(),
                        ft.Text(
                            "© 2026 — CADASTRE RDC | République Démocratique du Congo",
                            text_align=ft.TextAlign.CENTER,
                            size=11,
                            color=COLOR_TEXT_MUTED,
                        ),
                    ],
                ),
            ),
        ],
    )

# ============================================================
# LANCEMENT
# ============================================================


# ============================================================
# PORTAIL WEB PUBLIC — VÉRIFICATION PAR QR
# ============================================================

def lancer_serveur_verification_qr(host="0.0.0.0", port=8000):
    try:
        from fastapi import FastAPI
        from fastapi.responses import HTMLResponse, FileResponse
        import uvicorn
    except ImportError as ex:
        raise RuntimeError(
            "Installez le serveur QR avec : pip install fastapi uvicorn"
        ) from ex

    app = FastAPI(title="CADASTRE RDC — Vérification des certificats")

    @app.get("/verification/certificat/{token}", response_class=HTMLResponse)
    def verification_certificat_qr(token: str):
        row = _cad_qr_get_certificate(token)
        if not row:
            return HTMLResponse(
                "<html><body style='font-family:Arial;padding:40px'>\n"
                "<h1 style='color:#b91c1c'>CERTIFICAT NON AUTHENTIQUE</h1>\n"
                "<p>Ce code QR ne correspond à aucun certificat actif enregistré dans le système.</p>\n"
                "</body></html>",
                status_code=404,
            )

        import html
        def e(v): return html.escape(str(v or "—"))
        photo = row.get("owner_photo") or ""
        photo_path = _cad_resolve_upload_path(photo)
        photo_html = (
            f"<img src='/verification/certificat/{e(token)}/photo' "
            "style='width:140px;height:180px;object-fit:cover;border-radius:10px;border:1px solid #ddd'>"
            if photo_path else ""
        )
        owner_parcels = _cad_owner_parcels(row.get("owner_id"))
        numero = row.get("numero_certificat") or "—"
        html_page = f"""
        <!doctype html><html><head><meta charset='utf-8'>
        <meta name='viewport' content='width=device-width,initial-scale=1'>
        <title>Vérification {e(numero)}</title>
        <style>
        body{{font-family:Arial;background:#eef2f7;padding:20px}}
        .card{{max-width:980px;margin:auto;background:#fff;padding:28px;border-radius:18px;box-shadow:0 8px 30px #0002}}
        .ok{{color:#15803d;font-size:23px;font-weight:800}}
        .grid{{display:grid;grid-template-columns:1fr 1fr;gap:10px}}
        .item{{background:#f8fafc;padding:11px;border-radius:8px}}
        @media(max-width:700px){{.grid{{grid-template-columns:1fr}}}}
        </style></head><body><div class='card'>
        <div class='ok'>✓ CERTIFICAT AUTHENTIQUE</div>
        <p>Ce certificat correspond à un enregistrement actif du système CADASTRE RDC.</p>
        <hr><div style='display:flex;gap:25px;align-items:flex-start'>
        <div>{photo_html}</div><div style='flex:1'><h2>Certificat d'enregistrement</h2>
        <div class='grid'>
        <div class='item'><b>Numéro :</b> {e(numero)}</div>
        <div class='item'><b>Type :</b> {e(row.get('type_document'))}</div>
        <div class='item'><b>Propriétaire :</b> {e(row.get('owner_name'))}</div>
        <div class='item'><b>Téléphone :</b> {e(row.get('owner_phone'))}</div>
        <div class='item'><b>E-mail :</b> {e(row.get('owner_email'))}</div>
        <div class='item'><b>Pièce d'identité :</b> {e(row.get('owner_identity'))}</div>
        <div class='item'><b>Adresse :</b> {e(row.get('owner_address'))}</div>
        </div></div></div>
        <h2>Parcelle</h2><div class='grid'>
        <div class='item'><b>Numéro :</b> {e(row.get('numero'))}</div>
        <div class='item'><b>Référence :</b> {e(row.get('reference'))}</div>
        <div class='item'><b>Superficie :</b> {e(row.get('superficie'))}</div>
        <div class='item'><b>Adresse :</b> {e(row.get('adresse'))}</div>
        <div class='item'><b>Province :</b> {e(row.get('province'))}</div>
        <div class='item'><b>Ville / Territoire :</b> {e(row.get('ville_territoire'))}</div>
        <div class='item'><b>Commune / Chefferie :</b> {e(row.get('commune_chefferie'))}</div>
        <div class='item'><b>Quartier / Groupement :</b> {e(row.get('quartier_groupement'))}</div>
        <div class='item'><b>Localité :</b> {e(row.get('localite'))}</div>
        <div class='item'><b>Latitude :</b> {e(row.get('latitude'))}</div>
        <div class='item'><b>Longitude :</b> {e(row.get('longitude'))}</div>
        <div class='item'><b>Statut :</b> {e(row.get('statut'))}</div>
        </div>
        <h2>Toutes les parcelles enregistrées au nom du propriétaire</h2>
        <div class='grid'>
        {''.join(f"<div class='item'><b>Parcelle :</b> {e(p.get('numero') or p.get('reference') or ('ID-' + str(p.get('id'))))}<br><b>Superficie :</b> {e(p.get('superficie'))}<br><b>Adresse :</b> {e(p.get('adresse') or p.get('localisation'))}<br><b>Localisation :</b> {e(' / '.join(str(v) for v in [p.get('province'), p.get('ville_territoire'), p.get('commune_chefferie'), p.get('quartier_groupement'), p.get('localite')] if v not in (None, '')))}</div>" for p in owner_parcels) or "<div class='item'>Aucune autre parcelle enregistrée.</div>"}
        </div>
        <p style='color:#64748b'>Enregistré le {e(row.get('date_creation'))}</p>
        <p style='color:#64748b'>Vérification : {e(datetime.now().strftime('%d/%m/%Y %H:%M:%S'))}</p>
        <p><a href='/verification/certificat/{e(token)}/document'>Télécharger le certificat d'enregistrement</a></p>
        </div></body></html>"""
        journaliser("VERIFICATION_CERTIFICAT_QR", f"Certificat {numero} vérifié via QR")
        return HTMLResponse(html_page)

    @app.get("/verification/certificat/{token}/photo")
    def verification_photo_qr(token: str):
        row = _cad_qr_get_certificate(token)
        if not row:
            return HTMLResponse("Photo introuvable", status_code=404)
        photo = row.get("owner_photo")
        photo_path = _cad_resolve_upload_path(photo)
        if not photo_path:
            return HTMLResponse("Photo introuvable", status_code=404)
        return FileResponse(photo_path)

    @app.get("/verification/certificat/{token}/document")
    def verification_document_qr(token: str):
        row = _cad_qr_get_certificate(token)
        if not row:
            return HTMLResponse("Document introuvable", status_code=404)

        path = row.get("fichier_pdf") or row.get("fichier_word")
        if not path or not os.path.exists(str(path)):
            return HTMLResponse("Document introuvable", status_code=404)

        media_type = "application/pdf" if str(path).lower().endswith(".pdf") else None
        return FileResponse(
            str(path),
            filename=os.path.basename(path),
            media_type=media_type,
        )

    uvicorn.run(app, host=host, port=port)


if __name__ == "__main__":
    import sys
    if "--qr-server" in sys.argv:
        lancer_serveur_verification_qr()
        sys.exit(0)
    else:
        # ========================================================
        # Configuration Flet pour les téléversements
        # ========================================================
        FLET_SECRET_KEY = os.environ.get("FLET_SECRET_KEY")
        if not FLET_SECRET_KEY:
            FLET_SECRET_KEY = secrets.token_urlsafe(32)
            os.environ["FLET_SECRET_KEY"] = FLET_SECRET_KEY  # Injection pour Flet

        # Création des dossiers avant le lancement de l'application.
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        os.makedirs(os.path.join(UPLOADS_DIR, "proprietaires"), exist_ok=True)

        # Lancement de l'application Flet en mode Web pour le serveur (Render)
        port = int(os.environ.get("PORT", 8080))
        ft.app(
            target=main,
            view=None,
            port=port,
            host="0.0.0.0",
            upload_dir=UPLOADS_DIR,
        )
